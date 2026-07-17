"""
VivaSense Genetics – Enhanced Word Report Export
================================================
POST /genetics/download-results
POST /genetics/export-word  (alias)

Accepts DownloadReportRequest (UploadAnalysisResponse + optional correlation)
and generates a publication-ready .docx report containing:

  • Title & metadata
  • Trait Summary Table (all traits side-by-side)
  • Trait Correlations (if correlation data supplied)
  • Per-trait sections (new page each):
      – Executive summary
      – Descriptive statistics
      – ANOVA table (significance-starred, narrative)
      – Mean separation (table + bar chart, 300 DPI)
      – Genetic parameters (formulas, GCV/PCV commentary)
      – Interpretation and domain-appropriate recommendations
  • Footer
"""

import io
import logging
import math
import datetime
import re
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import matplotlib
matplotlib.use("Agg")                        # headless – no display needed
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import APIRouter
from fastapi.responses import JSONResponse, Response
from pydantic import BaseModel, Field

from genetics_schemas import AnovaTable, MeanSeparation, InteractionMeans, GeneticsResult, GeneticsResponse
from multitrait_upload_schemas import UploadAnalysisResponse, SummaryTableRow, TraitResult
from module_schemas import PathAnalysisResponse
from trait_relationships_schemas import CorrelationResponse
import result_cache
from genetics_interpretation import (
    generate_genetics_interpretation,
    build_breeding_synthesis,
    _describe_env_effects,
    _describe_gcv_pcv,
    _is_single_environment_analysis,
)
from domain_guard import find_forbidden_breeding_terms, is_plant_breeding_domain
from interpretation import InterpretationEngine

logger = logging.getLogger(__name__)
router = APIRouter(tags=["Export"])

FAILED_TRAIT_CV_MESSAGE = "CV% unavailable due to failed ANOVA estimation."

_REPORT_BLOCKED_PHRASES = [
    "clean genetic signal",
    "top-performing genotype",
    "additive gene effects",
    "non-additive effects",
]

_DUPLICATE_ENV_SENTENCE_1 = (
    "indicating relatively limited environmental variance influence under the evaluated conditions."
)
_DUPLICATE_ENV_SENTENCE_2 = (
    "GCV and PCV values were relatively close, suggesting limited environmental influence on phenotypic expression under the evaluated conditions."
)


# ============================================================================
# REQUEST MODEL
# ============================================================================

class DownloadReportRequest(UploadAnalysisResponse):
    """
    Extends UploadAnalysisResponse with optional correlation and path analysis data.
    The frontend can POST the raw UploadAnalysisResponse and these fields
    will default to None — fully backwards compatible.
    """
    module: str = "genetic_parameters"
    correlation: Optional[CorrelationResponse] = None
    path_analysis: Optional[PathAnalysisResponse] = None


# ─── Tukey-group colour palette ───────────────────────────────────────────────
_GROUP_COLOURS: Dict[str, str] = {
    "a": "#2ecc71",
    "b": "#f39c12",
    "c": "#e74c3c",
    "d": "#95a5a6",
    "e": "#34495e",
}
_DEFAULT_BAR_COLOUR = "#7f8c8d"

# ─── ANOVA source → human label ───────────────────────────────────────────────
_ANOVA_LABELS: Dict[str, str] = {
    "rep": "Replication",
    "genotype": "Genotype",
    "environment": "Environment",
    "environment:rep": "Rep(Environment)",
    "genotype:environment": "G×E Interaction",
    "Residuals": "Error",
}


def _anova_env_effect_stats(at: Optional[AnovaTable]) -> Tuple[Optional[float], Optional[float], Optional[float], Optional[float]]:
    if at is None or not hasattr(at, "source"):
        return None, None, None, None

    def _get_fp(source_name: str) -> Tuple[Optional[float], Optional[float]]:
        try:
            idx = at.source.index(source_name)
            f_val = at.f_value[idx] if idx < len(at.f_value) else None
            p_val = at.p_value[idx] if idx < len(at.p_value) else None
            return (
                float(f_val) if f_val is not None else None,
                float(p_val) if p_val is not None else None,
            )
        except (ValueError, IndexError, TypeError):
            return None, None

    f_env, p_env = _get_fp("environment")
    f_gxe, p_gxe = None, None
    for src in ["genotype:environment", "environment:genotype", "GxE", "gxe"]:
        f_gxe, p_gxe = _get_fp(src)
        if f_gxe is not None or p_gxe is not None:
            break

    return f_env, p_env, f_gxe, p_gxe

_HEADER_BG = "F2F2F2"

def _collect_doc_text(doc: Document) -> str:
    chunks: List[str] = []
    chunks.extend(p.text for p in doc.paragraphs if p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    chunks.append(cell.text)
    return "\n".join(chunks)


# ============================================================================
# NUMBER FORMATTING
# ============================================================================

def _fmt(value, decimals: int = 3, thousands: bool = False) -> str:
    """Format a numeric value for report display.
    Returns em-dash for None/NaN values."""
    if value is None:
        return "—"
    try:
        if isinstance(value, float) and (
            math.isnan(value) or
            math.isinf(value)
        ):
            return "—"
        fmt_spec = f"{',' if thousands else ''}.{decimals}f"
        return format(value, fmt_spec)
    except (TypeError, ValueError):
        return str(value)


def _clean_cv_percent(value: Optional[float]) -> Optional[float]:
    if value is None:
        return None
    try:
        return abs(float(value))
    except (TypeError, ValueError):
        return None


def _fmt_cv(value: Optional[float]) -> str:
    return _fmt(_clean_cv_percent(value), 2)


def _cv_precision_narrative(cv_percent: Optional[float], domain: Optional[str] = "plant_breeding") -> str:
    cv_val = _clean_cv_percent(cv_percent)
    if cv_val is None:
        return ""
    if cv_val < 1:
        return "Residual variability was extremely low relative to the trait mean. Verify raw data consistency and experimental realism."
    if cv_val < 10:
        return "Residual variability was relatively low, suggesting high experimental precision within the scope of this design."
    elif cv_val < 20:
        return (
            "Experimental variability appeared acceptable for treatment comparison under the evaluated conditions."
            if not is_plant_breeding_domain(domain)
            else "Experimental variability appeared acceptable for genotype comparison under the evaluated conditions."
        )
    return "Residual variability was comparatively high, and findings should therefore be interpreted cautiously."


def _find_breeding_governance_hits(text: str, analysis_type: Optional[str]) -> List[str]:
    lower_text = (text or "").lower()
    hits: List[str] = []
    blocked_patterns = [
        ("additive gene effects", "additive gene effects"),
        ("non-additive effects", "non-additive effects"),
        ("clean genetic signal", "clean genetic signal"),
        ("top-performing genotype", "top-performing genotype"),
        ("negligible environmental variance inflation", "negligible environmental variance inflation"),
    ]
    for label, needle in blocked_patterns:
        if needle in lower_text:
            hits.append(label)
    if re.search(r"\benvironmental effects[^.]*\bnon-significant\b|\bnon-significant\b[^.]*\benvironmental effects\b", lower_text):
        hits.append("environmental effects non-significant")

    if _is_single_environment_analysis(analysis_type) and (
        "gxe interaction was non-significant" in lower_text or "gx e interaction was non-significant" in lower_text
    ):
        hits.append("GxE non-significant (single-environment)")
    return hits


def _cv_from_anova_error_ms(at: Optional[AnovaTable], grand_mean: Optional[float]) -> Optional[float]:
    if at is None or grand_mean is None:
        return None
    try:
        gm = float(grand_mean)
    except (TypeError, ValueError):
        return None
    if gm == 0:
        return None
    if not hasattr(at, "source") or not hasattr(at, "ms"):
        return None

    error_terms = ["Error B", "Residuals", "residuals", "Residual", "residual", "error", "Error"]
    for term in error_terms:
        try:
            idx = at.source.index(term)
            mse = at.ms[idx]
            if mse is None:
                continue
            mse_val = float(mse)
            if mse_val < 0:
                continue
            return abs((mse_val ** 0.5) / gm * 100.0)
        except (ValueError, IndexError, TypeError):
            continue
    return None


def _resolve_cv_percent(result: GeneticsResult) -> Optional[float]:
    ds = result.descriptive_stats or {}
    ds_cv = _clean_cv_percent(ds.get("cv_percent")) if isinstance(ds, dict) else None
    if ds_cv is not None:
        return ds_cv
    return _cv_from_anova_error_ms(result.anova_table, result.grand_mean)


def _map_precision_label(cv_percent: Optional[float]) -> Optional[str]:
    cv = _clean_cv_percent(cv_percent)
    if cv is None:
        return None
    if cv < 10:
        return "high"
    if cv < 20:
        return "acceptable"
    return "caution"


def _find_export_quality_hits(report_text: str, data: "DownloadReportRequest") -> List[str]:
    hits: List[str] = []
    lower_text = report_text.lower()

    has_success_trait = any(
        (tr is not None)
        and (tr.status == "success")
        and (tr.analysis_result is not None)
        and (tr.analysis_result.result is not None)
        for tr in (data.trait_results or {}).values()
    )
    if has_success_trait and "cv (%): unavailable" in lower_text:
        hits.append("CV (%): Unavailable present despite successful ANOVA trait(s)")

    s1 = _DUPLICATE_ENV_SENTENCE_1.lower()
    s2 = _DUPLICATE_ENV_SENTENCE_2.lower()
    if s1 in lower_text and s2 in lower_text:
        hits.append("Duplicate environmental variance interpretation sentences detected")

    for phrase in _REPORT_BLOCKED_PHRASES:
        if phrase in lower_text:
            hits.append(f"Forbidden phrase detected: {phrase}")

    return hits


def _fmt_p(p: Optional[float]) -> str:
    if p is None:
        return "—"
    if p < 0.001:
        return "< 0.001 ***"
    if p < 0.01:
        return f"{p:.4f} **"
    if p < 0.05:
        return f"{p:.4f} *"
    return f"{p:.4f} ns"


def _sig_label(p: Optional[float]) -> str:
    if p is None:
        return ""
    if p < 0.001:
        return "***"
    if p < 0.01:
        return "**"
    if p < 0.05:
        return "*"
    return "ns"


def _p_for_sentence(p: Optional[float]) -> str:
    if p is None:
        return "p = —"
    if p < 0.001:
        return "p < 0.001"
    return f"p = {p:.3f}"


def _extract_source_stats(at: Optional[AnovaTable], source_name: str) -> Tuple[Optional[float], Optional[float], Optional[float]]:
    if at is None or not at.source:
        return None, None, None
    try:
        idx = at.source.index(source_name)
    except ValueError:
        return None, None, None

    f_val = at.f_value[idx] if idx < len(at.f_value) else None
    p_val = at.p_value[idx] if idx < len(at.p_value) else None
    ss_val = at.ss[idx] if idx < len(at.ss) else None
    return f_val, p_val, ss_val


def _eta_squared_for_source(at: Optional[AnovaTable], source_name: str) -> Optional[float]:
    if at is None or not at.source or not at.ss:
        return None
    _, _, ss_effect = _extract_source_stats(at, source_name)
    if ss_effect is None:
        return None

    # Exclude Intercept row — its large SS would artificially deflate η².
    # η² = SS_source / SS_total where SS_total = Σ of all non-intercept sources
    ss_total = sum(
        float(at.ss[i])
        for i, src in enumerate(at.source)
        if at.ss[i] is not None
        and str(src).strip().lower() not in {"(intercept)", "intercept"}
    )
    if ss_total <= 0:
        return None
    return float(ss_effect) / ss_total


# Authorized selection intensity values from Falconer & Mackay (1996)
SELECTION_INTENSITY_TABLE = {
    0.05: {"i": 2.063, "label": "Top 5%  (i = 2.063)"},
    0.10: {"i": 1.755, "label": "Top 10% (i = 1.755)"},
    0.20: {"i": 1.400, "label": "Top 20% (i = 1.400)"},
    0.25: {"i": 1.271, "label": "Top 25% (i = 1.271)"},
    0.50: {"i": 0.798, "label": "Top 50% (i = 0.798)"},
}

DEFAULT_SELECTION_INTENSITY = {
    "pct": 0.20,
    "i": 1.400,
    "label": "Top 20% (i = 1.400)"
}


def _selection_label_for_intensity(selection_intensity: Optional[float]) -> str:
    if selection_intensity is None:
        return DEFAULT_SELECTION_INTENSITY["label"]
    try:
        si = float(selection_intensity)
    except (TypeError, ValueError):
        return DEFAULT_SELECTION_INTENSITY["label"]

    # Match to the closest 'i' value in the authorized table
    best_pct = min(
        SELECTION_INTENSITY_TABLE.keys(),
        key=lambda pct: abs(SELECTION_INTENSITY_TABLE[pct]["i"] - si)
    )
    return SELECTION_INTENSITY_TABLE[best_pct]["label"]


def _selection_intensity_disclosure(selection_intensity: Optional[float]) -> str:
    label = _selection_label_for_intensity(selection_intensity)
    return f"Genetic Advance estimated using {label} (Falconer & Mackay, 1996)."


# ============================================================================
# TABLE HELPERS (python-docx XML)
# ============================================================================

def _set_cell_bg(cell, hex_colour: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def _set_cell_border(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _cell_font(cell, size_pt: int = 11, bold: bool = False,
               right_align: bool = False) -> None:
    for para in cell.paragraphs:
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        if right_align:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size_pt)
            run.font.bold = bold


def _bold_row(row, size_pt: int = 12, bg: Optional[str] = None) -> None:
    for cell in row.cells:
        if bg:
            _set_cell_bg(cell, bg)
        _set_cell_border(cell)
        _cell_font(cell, size_pt=size_pt, bold=True)


def _style_data_row(row, numeric_cols: Optional[set] = None) -> None:
    numeric_cols = numeric_cols or set()
    for i, cell in enumerate(row.cells):
        _set_cell_border(cell)
        _cell_font(cell, size_pt=11, right_align=(i in numeric_cols))


def _add_stat_table(
    doc: Document,
    headers: List[str],
    rows: List[List[str]],
    numeric_cols: Optional[set] = None,
) -> None:
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    _bold_row(hdr, size_pt=12, bg=_HEADER_BG)

    for row_data in rows:
        r = table.add_row()
        for i, val in enumerate(row_data):
            r.cells[i].text = str(val)
        _style_data_row(r, numeric_cols=numeric_cols)


# ============================================================================
# VISUALIZATION
# ============================================================================

def _bar_colour(group_letter: str) -> str:
    first = group_letter[0].lower() if group_letter else ""
    return _GROUP_COLOURS.get(first, _DEFAULT_BAR_COLOUR)


def _generate_mean_separation_chart(
    trait_name: str,
    genotypes: List[str],
    means: List[float],
    ses: List[Optional[float]],
    groups: List[str],
    domain: str = "plant_breeding",
    factor_name: Optional[str] = None,
) -> bytes:
    logger.info("Chart generation started for trait: '%s'. Received %d means, %d groups.", trait_name, len(means), len(groups))
    try:
        for style in ("seaborn-v0_8", "ggplot", "default"):
            try:
                plt.style.use(style)
                break
            except OSError:
                continue

        fig, ax = plt.subplots(figsize=(6, 4))
        n = len(genotypes)
        x = np.arange(n)
        colours = [_bar_colour(g) for g in groups]
        err = [s if (s is not None and not math.isnan(s)) else 0.0 for s in ses]

        ax.bar(
            x, means,
            color=colours,
            edgecolor="white",
            linewidth=0.5,
            yerr=err,
            capsize=4,
            error_kw={"elinewidth": 1, "ecolor": "#555555"},
            zorder=3,
        )

        max_val = max(means) if means else 1.0
        label_offset = max_val * 0.03
        for xi, (m, s, g) in enumerate(zip(means, err, groups)):
            ax.text(
                xi, m + s + label_offset,
                g,
                ha="center", va="bottom",
                fontweight="bold", fontsize=10,
                color="#222222",
            )

        ax.set_xticks(x)
        ax.set_xticklabels(
            genotypes,
            rotation=45 if n > 8 else 0,
            ha="right" if n > 8 else "center",
            fontsize=9,
        )
        ax.set_ylabel(trait_name, fontsize=11)
        _xlabel = factor_name if factor_name else ("Treatment" if not is_plant_breeding_domain(domain) else "Genotype")
        ax.set_xlabel(_xlabel, fontsize=11)
        ax.set_title(f"Mean Separation — {trait_name}", fontsize=12, fontweight="bold")
        ax.yaxis.grid(True, color="#cccccc", alpha=0.5, linewidth=0.7, zorder=0)
        ax.set_axisbelow(True)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)

        present_groups = sorted({g[0].lower() for g in groups if g})
        patches = [
            mpatches.Patch(
                color=_GROUP_COLOURS.get(g, _DEFAULT_BAR_COLOUR),
                label=f"Group {g}",
            )
            for g in present_groups
        ]
        if len(patches) > 1:
            ax.legend(handles=patches, loc="upper right", fontsize=8, framealpha=0.7)

        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=300, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        chart_bytes = buf.read()
        logger.info("Chart generation successful for trait: '%s'. Image bytes length: %d", trait_name, len(chart_bytes))
        return chart_bytes

    except Exception as exc:
        logger.error("Chart generation failed for '%s': %s (Exception type: %s)", trait_name, exc, type(exc).__name__, exc_info=True)
        plt.close("all")
        return b""


# ============================================================================
# DOCUMENT HELPERS
# ============================================================================

def _add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(6 if level > 1 else 12)
    h.paragraph_format.space_after = Pt(3)


def _add_body(doc: Document, text: str, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if italic or bold:
        for run in p.runs:
            if italic:
                run.italic = True
            if bold:
                run.bold = True


def _add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    rk = p.add_run(f"{key}: ")
    rk.bold = True
    rk.font.name = "Calibri"
    rk.font.size = Pt(11)
    rv = p.add_run(value)
    rv.font.name = "Calibri"
    rv.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


# ============================================================================
# RECOMMENDATION SUPPORT
# ============================================================================


# ============================================================================
# SECTION: CROSS-TRAIT SUMMARY TABLE
# ============================================================================

def _status_label(raw_status: Any) -> str:
    status = str(raw_status or "").strip().lower()
    mapping = {
        "success": "Analysis complete",
        "warning": "Check assumptions",
        "error": "Analysis failed",
    }
    return mapping.get(status, str(raw_status or "—"))

def _add_summary_table(doc: Document, data: UploadAnalysisResponse, domain: Optional[str] = None) -> None:
    is_anova = getattr(data, "module", "") == "anova"
    is_agronomy = not is_plant_breeding_domain(domain)

    if is_anova:
        headers = ["Trait", "Mean", "CV %", "Status"]
    elif is_agronomy:
        headers = ["Trait", "Mean", "Status"]
    else:
        headers = ["Trait", "Mean", "H²", "GCV %", "PCV %", "GAM %", "Class", "Status"]

    rows_data = []
    for row in data.summary_table:
        if is_anova:
            tr = (data.trait_results or {}).get(row.trait)
            cv_pct = None
            if tr and tr.analysis_result and tr.analysis_result.result:
                cv_pct = _resolve_cv_percent(tr.analysis_result.result)
            rows_data.append([
                row.trait,
                _fmt(row.grand_mean),
                _fmt_cv(cv_pct) if cv_pct is not None else "—",
                _status_label(row.status),
            ])
        elif is_agronomy:
            rows_data.append([
                row.trait,
                _fmt(row.grand_mean),
                _status_label(row.status),
            ])
        else:
            rows_data.append([
                row.trait,
                _fmt(row.grand_mean),
                _fmt(row.h2, 3),
                _fmt(row.gcv, 2),
                _fmt(row.pcv, 2),
                _fmt(row.gam_percent, 2),
                (
                    row.gam_class
                    or (
                        InterpretationEngine.classify_gam(row.gam_percent)
                        if row.gam_percent is not None
                        else None
                    )
                    or "—"
                ),
                _status_label(row.status),
            ])

    if is_anova:
        numeric_cols = {1, 2}
    elif is_agronomy:
        numeric_cols = {1}
    else:
        numeric_cols = {1, 2, 3, 4, 5}
    _add_stat_table(doc, headers, rows_data, numeric_cols=numeric_cols)


def _extract_gxe_stats_from_anova(at: Optional[AnovaTable]) -> Tuple[Optional[float], Optional[float]]:
    if at is None:
        return None, None
    aliases = {"genotype:environment", "environment:genotype", "gxe", "gxe interaction"}
    for idx, src in enumerate(at.source or []):
        src_norm = str(src).strip().lower()
        if src_norm in aliases:
            f_val = at.f_value[idx] if idx < len(at.f_value) else None
            p_val = at.p_value[idx] if idx < len(at.p_value) else None
            return f_val, p_val
    return None, None


def _extract_genotype_stats_from_anova(
    at: Optional[AnovaTable],
) -> Tuple[Optional[float], Optional[float], Optional[bool]]:
    if at is None:
        return None, None, None

    aliases = {"genotype", "genotypes"}

    def _norm(label: Any) -> str:
        return " ".join(str(label).strip().lower().replace("×", "x").split())

    for idx, src in enumerate(at.source or []):
        src_norm = _norm(src)
        if src_norm in aliases:
            f_val = at.f_value[idx] if idx < len(at.f_value) else None
            p_val = at.p_value[idx] if idx < len(at.p_value) else None
            f_num = float(f_val) if f_val is not None else None
            p_num = float(p_val) if p_val is not None else None
            return f_num, p_num, (p_num <= 0.05 if p_num is not None else None)

    return None, None, None


def _build_breeding_input_for_export(data: UploadAnalysisResponse) -> List[Dict[str, Any]]:
    summary_map = {row.trait: row for row in data.summary_table if row.status == "success"}
    synthesis_input: List[Dict[str, Any]] = []

    for trait_name, tr in (data.trait_results or {}).items():
        if tr.status != "success" or tr.analysis_result is None or tr.analysis_result.result is None:
            continue

        result = tr.analysis_result.result
        summary_row = summary_map.get(trait_name)
        ms = result.mean_separation

        genotype_means: List[Dict[str, Any]] = []
        top_genotype: Optional[str] = None
        if ms is not None and ms.genotype and ms.mean:
            means = [float(m) if m is not None else float("-inf") for m in ms.mean]
            order = sorted(range(len(ms.genotype)), key=lambda i: means[i], reverse=True)
            rank_map = {idx: rank + 1 for rank, idx in enumerate(order)}
            if order:
                top_genotype = str(ms.genotype[order[0]])

            for idx, geno in enumerate(ms.genotype):
                mean_val = ms.mean[idx] if idx < len(ms.mean) else None
                grp = ms.group[idx] if idx < len(ms.group) else None
                genotype_means.append({
                    "genotype": str(geno),
                    "mean": float(mean_val) if mean_val is not None else None,
                    "rank": int(rank_map.get(idx, len(ms.genotype))),
                    "group": str(grp) if grp is not None else "",
                })

        f_gxe, p_gxe = _extract_gxe_stats_from_anova(result.anova_table)
        f_genotype, p_genotype, genotype_significant = _extract_genotype_stats_from_anova(result.anova_table)
        synthesis_input.append({
            "trait_name": trait_name,
            "h2": float(summary_row.h2) if summary_row is not None and summary_row.h2 is not None else None,
            "gam_class": summary_row.gam_class if summary_row is not None else None,
            "top_genotype": top_genotype,
            "analysis_type": (
                "multi_environment"
                if (result.n_environments is not None and result.n_environments > 1)
                else "single_environment"
            ),
            "f_gxe": float(f_gxe) if f_gxe is not None else None,
            "p_gxe": float(p_gxe) if p_gxe is not None else None,
            "f_value": f_genotype,
            "p_value": p_genotype,
            "genotype_significant": genotype_significant,
            "genotype_means": genotype_means,
        })

    return synthesis_input


# ============================================================================
# SECTION: TRAIT CORRELATIONS
# ============================================================================

def _add_correlation_section(doc: Document, corr: CorrelationResponse) -> None:
    doc.add_page_break()
    _add_heading(doc, "Trait Correlations", level=1)

    traits = corr.trait_names
    n = len(traits)
    method_label = corr.method.capitalize()

    _add_kv(doc, "Method", f"{method_label} correlation")
    _add_kv(doc, "Phenotypic Observations", str(corr.phenotypic.n_observations))
    _add_kv(doc, "Between-Genotype Means", str(corr.between_genotype.n_observations))
    if corr.genotypic is not None:
        _add_kv(doc, "Genotypic VC Observations", str(corr.genotypic.n_observations))
    doc.add_paragraph()

    # Build mode list: always include phenotypic + between_genotype; add genotypic VC if available
    mode_list = [
        ("Phenotypic (Field-Level)", corr.phenotypic),
        ("Between-Genotype Association", corr.between_genotype),
    ]
    if corr.genotypic is not None:
        mode_list.append(("Genotypic (Variance-Component)", corr.genotypic))

    for mode_name, stats in mode_list:
        _add_heading(doc, f"{mode_name} Correlations", level=2)

        is_vc = getattr(stats, "inference_approximate", False)
        # Distinguish column headers: approximate inference for VC, exact for others
        p_col      = "≈p-value"   if is_vc else "p-value"
        p_adj_col  = "≈p-adj (FDR)" if is_vc else "p-adj (FDR)"
        headers = ["Trait 1", "Trait 2", "rg" if is_vc else "r", p_col, p_adj_col, "Sig."]

        rows_data = []
        for i in range(n):
            for j in range(i + 1, n):
                r_val     = stats.r_matrix[i][j]     if stats.r_matrix     else None
                p_val     = stats.p_matrix[i][j]     if stats.p_matrix     else None
                p_adj_val = stats.p_adj_matrix[i][j] if stats.p_adj_matrix else None
                # For VC mode: if r exists but p is None, inference was suppressed for this pair
                if is_vc and r_val is not None and p_val is None:
                    p_display     = "unavailable"
                    p_adj_display = "unavailable"
                    sig_display   = "—"
                else:
                    p_display     = _fmt_p(p_val)
                    p_adj_display = _fmt_p(p_adj_val)
                    sig_display   = _sig_label(p_val) or "ns"
                rows_data.append([
                    traits[i],
                    traits[j],
                    _fmt(r_val, 3, thousands=False),
                    p_display,
                    p_adj_display,
                    sig_display,
                ])

        if rows_data:
            _add_stat_table(doc, headers, rows_data, numeric_cols={2})
            doc.add_paragraph()

        # Add inference note below the VC table
        if is_vc:
            note = getattr(stats, "inference_note", None) or (
                "Genotypic correlation is estimated from variance components using a mixed model. "
                "Confidence intervals and significance measures are approximate."
            )
            note_para = doc.add_paragraph(f"⚠ Note: {note}")
            note_para.runs[0].italic = True
            doc.add_paragraph()

    # Auto-interpretation
    _add_heading(doc, "Correlation Interpretation", level=2)
    if corr.interpretation:
        _add_body(doc, corr.interpretation)
        doc.add_paragraph()

    # Count strong positive correlations — prefer VC genotypic if available, else between-genotype
    _ref_stats = corr.genotypic if corr.genotypic is not None else corr.between_genotype
    strong_pos = [
        (traits[i], traits[j])
        for i in range(n)
        for j in range(i + 1, n)
        if _ref_stats.r_matrix
        and _ref_stats.r_matrix[i][j] is not None
        and _ref_stats.r_matrix[i][j] >= 0.70
        and _ref_stats.p_matrix
        and _ref_stats.p_matrix[i][j] is not None
        and _ref_stats.p_matrix[i][j] < 0.05
    ]
    if strong_pos:
        pairs_str = ", ".join(f"{a} & {b}" for a, b in strong_pos[:5])
        _add_body(
            doc,
            f"Strong positive correlations (r ≥ 0.70, p < 0.05) were detected "
            f"between: {pairs_str}. "
            "A co-selection strategy is recommended — improving one of these "
            "traits through selection will likely produce concurrent gains in "
            "the correlated traits, improving experimental efficiency.",
        )

    strong_neg = [
        (traits[i], traits[j])
        for i in range(n)
        for j in range(i + 1, n)
        if _ref_stats.r_matrix
        and _ref_stats.r_matrix[i][j] is not None
        and _ref_stats.r_matrix[i][j] <= -0.70
        and _ref_stats.p_matrix
        and _ref_stats.p_matrix[i][j] is not None
        and _ref_stats.p_matrix[i][j] < 0.05
    ]
    if strong_neg:
        pairs_str = ", ".join(f"{a} & {b}" for a, b in strong_neg[:5])
        _add_body(
            doc,
            f"Strong negative correlations were found between: {pairs_str}. "
            "Simultaneous improvement of these traits may require special "
            "crossing strategies or recurrent selection.",
        )

    if corr.statistical_note:
        _add_body(doc, corr.statistical_note, italic=True)


# ============================================================================
# SECTION: EXECUTIVE SUMMARY (per trait)
# ============================================================================

def _add_executive_summary(
    doc: Document,
    trait: str,
    result: GeneticsResult,
    is_anova: bool = False,
    domain: str = "plant_breeding",
) -> None:
    _add_heading(doc, "Executive Summary", level=2)

    is_agronomy = not is_plant_breeding_domain(domain)
    fields = [
        ("Trait", trait),
        ("Grand Mean", _fmt(result.grand_mean)),
    ]

    # Genetic parameters (H², GCV, PCV, GAM) are not meaningful for agronomy
    # treatment trials — suppress them entirely for non-plant_breeding domains.
    if not is_anova and not is_agronomy:
        hp = result.heritability if isinstance(result.heritability, dict) else {}
        gp = result.genetic_parameters if isinstance(result.genetic_parameters, dict) else {}
        h2 = hp.get("h2_broad_sense")
        gcv = gp.get("GCV")
        pcv = gp.get("PCV")
        gam_pct = gp.get("GAM_percent")
        # Use _resolve_cv_percent so ANOVA-MSE fallback applies when ds.cv_percent is absent
        cv = _resolve_cv_percent(result)
        h2_class = (
            "High" if h2 is not None and h2 >= 0.6
            else "Moderate" if h2 is not None and h2 >= 0.3
            else "Low" if h2 is not None
            else "—"
        )
        fields += [
            ("Heritability, broad-sense (H²)", f"{_fmt(h2, 3)} [{h2_class}]"),
            ("GCV (%)", _fmt(gcv, 2)),
            ("PCV (%)", _fmt(pcv, 2)),
            ("GAM (%)", _fmt(gam_pct, 2)),
        ]
        if cv is not None:
            fields.append(("CV (%)", _fmt_cv(cv)))
        else:
            fields.append(("CV (%)", "Unavailable"))
    elif is_anova and not is_agronomy:
        # ANOVA module + plant_breeding: show CV% (no H²/GAM for ANOVA module)
        # For split-plot, show CV-A (whole-plot) and CV-B (subplot) separately.
        vc = result.variance_components if isinstance(result.variance_components, dict) else {}
        cv_a_val = _clean_cv_percent(float(vc["cv_A"])) if vc.get("cv_A") is not None else None
        cv_b_val = _clean_cv_percent(float(vc["cv_B"])) if vc.get("cv_B") is not None else None
        if cv_a_val is not None and cv_b_val is not None:
            fields.append(("CV-A (Main-plot, %)", _fmt_cv(cv_a_val)))
            fields.append(("CV-B (Sub-plot, %)", _fmt_cv(cv_b_val)))
        else:
            cv = _resolve_cv_percent(result)
            if cv is not None:
                fields.append(("CV (%)", _fmt_cv(cv)))
            else:
                fields.append(("CV (%)", "Unavailable"))
    elif is_agronomy:
        # Show CV% and treatment significance for agronomy executive summary
        vc = result.variance_components if isinstance(result.variance_components, dict) else {}
        cv_a_val = _clean_cv_percent(float(vc["cv_A"])) if vc.get("cv_A") is not None else None
        cv_b_val = _clean_cv_percent(float(vc["cv_B"])) if vc.get("cv_B") is not None else None
        if cv_a_val is not None and cv_b_val is not None:
            fields.append(("CV-A (Main-plot, %)", _fmt_cv(cv_a_val)))
            fields.append(("CV-B (Sub-plot, %)", _fmt_cv(cv_b_val)))
        else:
            cv = _resolve_cv_percent(result)
            if cv is not None:
                fields.append(("CV (%)", _fmt_cv(cv)))
        at = result.anova_table
        if at and hasattr(at, "source") and "genotype" in at.source:
            idx = at.source.index("genotype")
            p_val = at.p_value[idx] if idx < len(at.p_value) else None
            sig = _sig_label(p_val) if p_val is not None else ""
            fields.append(("Treatment Effect", f"Significant {sig}" if sig not in ("", "ns") else "Not significant"))

    for key, val in fields:
        _add_kv(doc, key, val)

    if not is_agronomy:
        # Add CV precision narrative for both genetics and ANOVA modules (plant_breeding domain)
        # _resolve_cv_percent includes ANOVA-MSE fallback so narrative is always available
        cv_for_note = _resolve_cv_percent(result)
        cv_note = _cv_precision_narrative(cv_for_note, domain=domain)
        if cv_note:
            _add_body(doc, cv_note)

    doc.add_paragraph()


# ============================================================================
# SECTION: DESCRIPTIVE STATISTICS (per trait)
# ============================================================================

def _add_descriptive_stats(doc: Document, result: GeneticsResult, domain: str = "plant_breeding") -> None:
    _add_heading(doc, "Descriptive Statistics", level=2)

    _entry_label = "No. Treatments" if not is_plant_breeding_domain(domain) else "No. Genotypes"
    # Use only real, directly-available fields — do not fabricate derived stats
    rows: List[tuple] = [("Grand Mean", _fmt(result.grand_mean))]
    if result.n_genotypes is not None:
        rows.append((_entry_label, str(result.n_genotypes)))
    rows.append(("No. Replications", str(result.n_reps)))
    if result.n_environments:
        rows.append(("No. Environments", str(result.n_environments)))

    # Append extra fields from descriptive_stats dict if the R engine provided it
    if result.descriptive_stats and isinstance(result.descriptive_stats, dict):
        for key, val in result.descriptive_stats.items():
            label = key.replace("_", " ").title()
            if isinstance(val, float):
                rows.append((label, _fmt(val)))
            elif val is not None:
                rows.append((label, str(val)))

    _add_stat_table(doc, ["Parameter", "Value"], rows, numeric_cols={1})


# ============================================================================
# SECTION: DESIGN DECLARATION (per trait, before ANOVA)
# ============================================================================

_DESIGN_LABELS: Dict[str, str] = {
    "split_plot_rcbd":  "Split-Plot Randomised Complete Block Design (RCBD)",
    "factorial_rcbd":   "Factorial Randomised Complete Block Design (RCBD)",
    "factorial_crd":    "Factorial Completely Randomised Design (CRD)",
    "rcbd":             "Randomised Complete Block Design (RCBD)",
    "crd":              "Completely Randomised Design (CRD)",
}


def _add_design_statement(
    doc: Document,
    design_type: Optional[str],
    n_reps: Optional[int],
    trait_name: str,
    mp_label: Optional[str] = None,
    sp_label: Optional[str] = None,
) -> None:
    """Render a one-sentence experimental design declaration before the ANOVA table."""
    label = _DESIGN_LABELS.get(design_type or "", "")
    if not label:
        return

    if design_type == "split_plot_rcbd" and mp_label and sp_label:
        from analysis_anova_routes import _fmt_factor_label
        _mp = _fmt_factor_label(mp_label, "the main-plot factor")
        _sp = _fmt_factor_label(sp_label, "the subplot factor")
        stmt = (
            f"{label} analysis was performed for {trait_name} with {_mp} assigned to main plots "
            f"and {_sp} assigned to subplots within {n_reps or '—'} replication(s)."
        )
    elif design_type in {"factorial_rcbd", "factorial_crd"}:
        stmt = (
            f"{label} analysis was performed for {trait_name} "
            f"with {n_reps or '—'} replication(s)."
        )
    else:
        stmt = (
            f"{label} analysis was performed for {trait_name} "
            f"with {n_reps or '—'} replication(s)."
        )

    p = doc.add_paragraph(stmt)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.italic = True


# ============================================================================
# SECTION: ANOVA (per trait)
# ============================================================================

def _add_anova_section(doc: Document, at: AnovaTable, domain: str = "plant_breeding") -> None:
    _add_heading(doc, "Analysis of Variance (ANOVA)", level=2)

    is_agronomy = not is_plant_breeding_domain(domain)
    _domain_labels: Dict[str, str] = {
        **_ANOVA_LABELS,
        **({"genotype": "Treatment", "genotype:environment": "T×E Interaction"} if is_agronomy else {}),
    }

    # Compute SS-total (exclude intercept and error rows) for η²
    _error_labels = {"error a", "error b", "residuals", "residual", "within", "error"}
    ss_total = sum(
        float(at.ss[i])
        for i, src in enumerate(at.source)
        if at.ss[i] is not None
        and str(src).strip().lower() not in {"(intercept)", "intercept"}
    )

    headers = ["Source", "DF", "SS", "MS", "F-value", "p-value", "η²"]
    rows_data = []
    genotype_idx = ge_idx = None

    for i, src in enumerate(at.source):
        src_label = str(src).strip().lower()
        if src_label in {"(intercept)", "intercept"}:
            continue

        label = _domain_labels.get(src, src)
        df_val = at.df[i] if i < len(at.df) else None
        ss_val = at.ss[i] if i < len(at.ss) else None
        ms_val = at.ms[i] if i < len(at.ms) else None
        f_val  = at.f_value[i] if i < len(at.f_value) else None
        p_val  = at.p_value[i] if i < len(at.p_value) else None

        # η² only for treatment rows, not error/residual rows
        is_error_row = src_label in _error_labels
        if ss_val is not None and ss_total > 0 and not is_error_row:
            eta2 = float(ss_val) / ss_total
            eta2_str = f"{eta2:.3f}"
        else:
            eta2_str = "—"

        rows_data.append([
            label,
            str(int(df_val)) if df_val is not None else "—",
            _fmt(ss_val),
            _fmt(ms_val),
            _fmt(f_val, 3) if f_val is not None else "—",
            _fmt_p(p_val),
            eta2_str,
        ])
        if src == "genotype":
            genotype_idx = i
        if src == "genotype:environment":
            ge_idx = i

    _add_stat_table(doc, headers, rows_data, numeric_cols={1, 2, 3, 4})
    doc.add_paragraph()

    if genotype_idx is not None:
        f_val = at.f_value[genotype_idx] if genotype_idx < len(at.f_value) else None
        p_val = at.p_value[genotype_idx] if genotype_idx < len(at.p_value) else None
        sig = _sig_label(p_val)
        sig_word = (
            "highly significant" if sig in ("***", "**")
            else "significant" if sig == "*"
            else "not significant"
        )
        if is_agronomy:
            _add_body(
                doc,
                f"The treatment effect was {sig_word} "
                f"(F = {_fmt(f_val, 3)}, {_fmt_p(p_val)}), indicating "
                + ("substantial variation among treatments."
                   if sig_word != "not significant"
                   else "limited variation among treatments."),
            )
        else:
            _add_body(
                doc,
                f"The genotype effect was {sig_word} "
                f"(F = {_fmt(f_val, 3)}, {_fmt_p(p_val)}), indicating "
                + ("substantial genetic variation among genotypes."
                   if sig_word != "not significant"
                   else "limited genetic variation among genotypes."),
            )

    if ge_idx is not None:
        p_val = at.p_value[ge_idx] if ge_idx < len(at.p_value) else None
        sig = _sig_label(p_val)
        if is_agronomy:
            if sig in ("***", "**", "*"):
                _add_body(
                    doc,
                    f"The Treatment × Environment interaction was significant ({_fmt_p(p_val)}), "
                    "suggesting treatment performance differs across environments.",
                )
            else:
                _add_body(
                    doc,
                    "The Treatment × Environment interaction was not significant — treatment rankings "
                    "are stable across environments.",
                )
        else:
            if sig in ("***", "**", "*"):
                _add_body(
                    doc,
                    f"The G×E interaction was significant ({_fmt_p(p_val)}), "
                    "suggesting genotype performance differs across environments.",
                )
            else:
                _add_body(
                    doc,
                    "The G×E interaction was not significant — treatment rankings "
                    "are stable across environments.",
                )


# ============================================================================
# SECTION: KEY FINDINGS (split-plot and factorial)
# ============================================================================

def _add_key_findings(
    doc: "Document",
    result: "GeneticsResult",
    trait_name: str,
    mp_label: Optional[str] = None,
    sp_label: Optional[str] = None,
) -> None:
    """One-paragraph summary of two-factor ANOVA outcomes for split-plot or factorial designs."""
    at = result.anova_table
    if at is None:
        return

    # p-value lookup by source label (handles remapped labels like "Irrigation")
    pval: Dict[str, float] = {}
    if hasattr(at, "source") and hasattr(at, "p_value"):
        for s, p in zip(at.source, at.p_value):
            if s is not None and p is not None:
                try:
                    pval[str(s).strip()] = float(p)
                except (TypeError, ValueError):
                    pass

    main_col = mp_label or "main-plot factor"
    sub_col  = sp_label or "subplot factor"
    int_key  = f"{main_col}×{sub_col}"

    def _lookup(primary: str, *fallbacks: str) -> Optional[float]:
        for k in (primary,) + fallbacks:
            if k in pval:
                return pval[k]
        return None

    def _fmt_p(p: Optional[float]) -> Optional[str]:
        if p is None:
            return None
        return "p < 0.001" if p < 0.001 else f"p = {p:.4f}"

    def _sig(p: Optional[float]) -> str:
        return "significantly" if (p is not None and p < 0.05) else "non-significantly"

    mp_p   = _lookup(main_col, "main_plot")
    sp_p   = _lookup(sub_col,  "sub_plot")
    int_p  = _lookup(int_key,  "main_plot:sub_plot", "sub_plot:main_plot")

    mp_p_str  = _fmt_p(mp_p)
    sp_p_str  = _fmt_p(sp_p)
    int_p_str = _fmt_p(int_p)

    parts: List[str] = []
    if mp_p_str:
        parts.append(f"{main_col} {_sig(mp_p)} affected {trait_name} ({mp_p_str}).")
    if sp_p_str:
        parts.append(f"{sub_col} {_sig(sp_p)} affected {trait_name} ({sp_p_str}).")
    if int_p_str:
        int_significant = int_p is not None and int_p < 0.05
        int_desc = "significant" if int_significant else "non-significant"
        tail = (
            f", indicating that {sub_col} response depended on {main_col} level."
            if int_significant else "."
        )
        parts.append(
            f"A {int_desc} {main_col} × {sub_col} interaction was detected ({int_p_str}){tail}"
        )

    # Highest treatment-combination mean from interaction_means
    int_means = getattr(result, "interaction_means", None)
    if isinstance(int_means, dict):
        try:
            cell = int_means.get("cell_means", {})
            tv   = cell.get("trait_value") if isinstance(cell, dict) else None
            mps  = cell.get("main_plot")   if isinstance(cell, dict) else None
            sps  = cell.get("sub_plot")    if isinstance(cell, dict) else None
            if tv and mps and sps and len(tv) == len(mps) == len(sps):
                max_idx = max(range(len(tv)), key=lambda i: float(tv[i]))
                combo   = f"{mps[max_idx]} × {sps[max_idx]}"
                max_val = float(tv[max_idx])
                parts.append(
                    f"The highest treatment-combination mean was {combo} ({max_val:.2f})."
                )
        except Exception:
            pass

    if not parts:
        return

    _add_heading(doc, "Key Findings", level=2)
    _add_body(doc, " ".join(parts))


# ============================================================================
# SECTION: MEAN SEPARATION (per trait)
# ============================================================================

def _add_mean_separation_section(
    doc: Document,
    trait_name: str,
    ms: MeanSeparation,
    domain: str = "plant_breeding",
    factor_name: Optional[str] = None,
    interaction_significant: bool = False,
) -> None:
    if interaction_significant and factor_name:
        heading = (
            f"Marginal Mean Separation — {factor_name} "
            f"(Interpret with Interaction Context) — "
            f"{ms.test} (α = {ms.alpha})"
        )
    elif factor_name:
        heading = f"Mean Separation — Main Effect of {factor_name} — {ms.test} (α = {ms.alpha})"
    else:
        heading = f"Mean Separation — {ms.test} (α = {ms.alpha})"
    _add_heading(doc, heading, level=2)

    is_agronomy = not is_plant_breeding_domain(domain)
    entry_label = factor_name if factor_name else ("Treatment" if is_agronomy else "Genotype")
    headers = ["Rank", entry_label, "Mean", "SE", "Group"]
    rows_data = []
    for i, geno in enumerate(ms.genotype):
        se_val = ms.se[i] if i < len(ms.se) else None
        rows_data.append([
            str(i + 1),
            geno,
            _fmt(ms.mean[i]),
            _fmt(se_val),
            ms.group[i] if i < len(ms.group) else "—",
        ])
    _add_stat_table(doc, headers, rows_data, numeric_cols={0, 2, 3})
    doc.add_paragraph()

    top_letter = ms.group[0] if ms.group else "a"
    top_genos = [ms.genotype[i] for i, g in enumerate(ms.group) if g == top_letter]
    _entry_word = factor_name.capitalize() if factor_name else ("Treatment(s)" if is_agronomy else "Genotype(s)")
    _top_phrase = f"{_entry_word} with the highest observed mean in group '{top_letter}': "
    _add_body(
        doc,
        "Means followed by the same letter are not significantly different "
        f"at α = {ms.alpha} ({ms.test}). "
        + _top_phrase
        + ", ".join(top_genos) + ".",
        italic=True,
    )

    chart_bytes = _generate_mean_separation_chart(
        trait_name=trait_name,
        genotypes=ms.genotype,
        means=ms.mean,
        ses=ms.se,
        groups=ms.group,
        domain=domain,
        factor_name=factor_name,
    )
    if chart_bytes:
        doc.add_paragraph()
        _add_heading(doc, "Mean Separation Chart", level=3)
        doc.add_picture(io.BytesIO(chart_bytes), width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(
            f"Figure: Mean ± SE for {trait_name}. "
            "Bars sharing the same letter are not significantly different "
            f"({ms.test}, α = {ms.alpha})."
        )
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap.runs:
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)
    else:
        _add_body(doc, "(Chart unavailable for this trait.)", italic=True)


def _add_interaction_separation_section(
    doc: Document,
    trait_name: str,
    int_sep: InteractionMeans,
    domain: str = "plant_breeding",
) -> None:
    a_label = int_sep.genotype_label or "Factor A"
    b_label = int_sep.factor_label or "Factor B"
    _add_heading(
        doc,
        f"Mean Separation — {a_label} × {b_label} Interaction — {int_sep.test} (α = {int_sep.alpha})",
        level=2,
    )
    headers = ["Rank", a_label, b_label, "Mean", "SE", "Group"]
    rows_data = []
    for i, (geno, fac) in enumerate(zip(int_sep.genotype, int_sep.factor)):
        se_val = int_sep.se[i] if i < len(int_sep.se) else None
        rows_data.append([
            str(i + 1),
            geno,
            fac,
            _fmt(int_sep.mean[i]),
            _fmt(se_val),
            int_sep.group[i] if i < len(int_sep.group) else "—",
        ])
    _add_stat_table(doc, headers, rows_data, numeric_cols={0, 3, 4})
    doc.add_paragraph()
    note = (
        f"Treatment-combination means sorted by mean descending. "
        f"Cells sharing the same letter are not significantly different at α = {int_sep.alpha} ({int_sep.test})."
        if int_sep.test != "Cell Means"
        else "Interaction HSD could not be computed — treatment-combination means displayed without grouping letters."
    )
    _add_body(doc, note, italic=True)


# ============================================================================
# SECTION: INTERACTION MEANS (split-plot A×B cell means table + line plot)
# ============================================================================

def _add_interaction_means_section(
    doc: Document,
    interaction_means_dict: Dict[str, Any],
    mp_label: str = "Main-Plot Factor",
    sp_label: str = "Subplot Factor",
    trait_name: str = "",
    is_significant: bool = True,
    sp_mean_separation: Optional[Any] = None,
) -> None:
    """Render A×B treatment-combination means table for split-plot designs."""
    _add_heading(doc, f"Treatment-Combination Means — {mp_label} × {sp_label}", level=2)

    if is_significant:
        _add_body(
            doc,
            f"Because the {mp_label} × {sp_label} interaction was statistically significant, "
            "treatment-combination means provide the primary basis for biological and agronomic "
            "interpretation. Marginal means (main effects) should be interpreted with caution.",
        )
        doc.add_paragraph()

    cell_means = interaction_means_dict.get("cell_means") or {}
    mp_vals = cell_means.get("main_plot") or []
    sp_vals = cell_means.get("sub_plot") or []
    tv_vals = cell_means.get("trait_value") or []

    if not mp_vals:
        _add_body(doc, "Treatment-combination means not available.", italic=True)
        return

    cell_se = interaction_means_dict.get("cell_se")
    _has_se = (
        cell_se is not None
        and not (isinstance(cell_se, float) and (cell_se != cell_se))
    )
    headers = [mp_label, sp_label, f"Mean ({trait_name})"] + (["SE"] if _has_se else [])
    rows_data = [
        [str(mp), str(sp), _fmt(tv, 3)] + ([_fmt(cell_se, 4)] if _has_se else [])
        for mp, sp, tv in zip(mp_vals, sp_vals, tv_vals)
    ]
    _add_stat_table(doc, headers, rows_data, numeric_cols={2, 3} if _has_se else {2})
    doc.add_paragraph()

    # Dynamic highest/lowest treatment-combination context sentence
    try:
        numeric_tv = [float(v) for v in tv_vals if v is not None]
        if len(numeric_tv) >= 2:
            max_idx = tv_vals.index(max(numeric_tv))
            min_idx = tv_vals.index(min(numeric_tv))
            max_combo = f"{mp_vals[max_idx]} × {sp_vals[max_idx]}"
            min_combo = f"{mp_vals[min_idx]} × {sp_vals[min_idx]}"
            _add_body(
                doc,
                f"The highest {trait_name} response was observed under "
                f"{max_combo} ({float(tv_vals[max_idx]):.3f}), while the lowest response "
                f"occurred under {min_combo} ({float(tv_vals[min_idx]):.3f}). "
                f"These treatment-combination means are the primary basis for applied "
                f"conclusions when the {mp_label} × {sp_label} interaction is significant.",
                italic=True,
            )
            doc.add_paragraph()
    except (TypeError, ValueError, IndexError):
        pass

    # Interpretation note
    note_parts = [
        "Treatment-combination means are observed averages for each factor-level combination. "
        "Cells in the same row share the same main-plot level; differences within a row "
        f"reflect the {sp_label} effect at that {mp_label} level."
    ]
    if _has_se:
        note_parts.append(
            f"SE = sqrt(MS_Error B / n_reps); applies uniformly to all treatment-combination means."
        )
    elif sp_mean_separation is not None:
        try:
            se_vals = [s for s in (sp_mean_separation.se or []) if s is not None]
            if se_vals:
                mean_se = sum(se_vals) / len(se_vals)
                note_parts.append(
                    f"Indicative standard error (SE) from {sp_label} mean separation "
                    f"(Error B): {_fmt(mean_se, 3)}. "
                    f"Use the {sp_label} LSD ({sp_mean_separation.test}, α = {sp_mean_separation.alpha}) "
                    "for pairwise comparisons among subplot levels within each main-plot level."
                )
        except (AttributeError, TypeError, ZeroDivisionError):
            pass
    _add_body(doc, " ".join(note_parts), italic=True)


def _generate_interaction_line_plot(
    trait_name: str,
    interaction_means_dict: Dict[str, Any],
    mp_label: str = "Main-Plot Factor",
    sp_label: str = "Subplot Factor",
) -> Optional[bytes]:
    """
    Generate a publication-quality interaction line plot.
    Returns PNG bytes, or None on any failure.
    One line per main-plot level; x-axis = sub-plot levels; y-axis = mean response.
    """
    try:
        cell_means = interaction_means_dict.get("cell_means") or {}
        mp_vals = cell_means.get("main_plot") or []
        sp_vals = cell_means.get("sub_plot") or []
        tv_vals = cell_means.get("trait_value") or []
        if not mp_vals or not sp_vals or not tv_vals:
            return None

        # Build per-main-plot level series
        mp_levels_ordered = list(dict.fromkeys(mp_vals))  # preserve insertion order
        sp_levels_ordered = list(dict.fromkeys(sp_vals))

        # means[mp_level][sp_level] = float
        means: Dict[str, Dict[str, float]] = {}
        for mp, sp, tv in zip(mp_vals, sp_vals, tv_vals):
            means.setdefault(mp, {})[sp] = float(tv)

        for style in ("seaborn-v0_8", "ggplot", "default"):
            try:
                plt.style.use(style)
                break
            except OSError:
                continue

        fig, ax = plt.subplots(figsize=(6, 4))
        x = list(range(len(sp_levels_ordered)))
        colours = ["#1D9E75", "#E07B3E", "#4A90D9", "#9B59B6", "#E74C3C", "#2ECC71"]

        for k, mp in enumerate(mp_levels_ordered):
            y = [means.get(mp, {}).get(sp, float("nan")) for sp in sp_levels_ordered]
            colour = colours[k % len(colours)]
            ax.plot(x, y, marker="o", linewidth=2, markersize=6, label=mp, color=colour)

        ax.set_xticks(x)
        ax.set_xticklabels(sp_levels_ordered, rotation=30 if len(sp_levels_ordered) > 5 else 0, ha="right")
        ax.set_xlabel(sp_label, fontsize=11)
        ax.set_ylabel(trait_name, fontsize=11)
        ax.set_title(f"{mp_label} × {sp_label} Interaction", fontsize=12, fontweight="bold")
        ax.legend(title=mp_label, fontsize=9, title_fontsize=9, framealpha=0.7)
        ax.grid(axis="y", linestyle="--", alpha=0.5)
        ax.spines[["top", "right"]].set_visible(False)
        fig.tight_layout()

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception as exc:
        logger.warning("Interaction line plot failed for '%s': %s", trait_name, exc)
        return None


def _add_interaction_plot_to_doc(
    doc: Document,
    interaction_means_dict: Dict[str, Any],
    trait_name: str,
    mp_label: str,
    sp_label: str,
) -> None:
    """Embed interaction line plot into the Word document."""
    plot_bytes = _generate_interaction_line_plot(trait_name, interaction_means_dict, mp_label, sp_label)
    if plot_bytes is None:
        return
    _add_heading(doc, f"{mp_label} × {sp_label} Interaction Plot", level=2)
    buf = io.BytesIO(plot_bytes)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(buf, width=Inches(5.5))
    _add_body(
        doc,
        f"Interaction plot for {trait_name}. Lines represent {mp_label} levels; "
        f"x-axis shows {sp_label} levels. Non-parallel lines indicate a significant interaction.",
        italic=True,
    )


# ============================================================================
# SECTION: GENETIC PARAMETERS (per trait)
# ============================================================================

def _add_genetic_parameters_section(doc: Document, result: GeneticsResult, domain: Optional[str] = None) -> None:
    is_agronomy = not is_plant_breeding_domain(domain)
    if is_agronomy:
        logger.info("Skipping genetic parameters section for non-plant-breeding domain: %s", domain)
        return

    _add_heading(doc, "Genetic Parameters", level=2)

    vc = result.variance_components if isinstance(result.variance_components, dict) else {}
    hp = result.heritability if isinstance(result.heritability, dict) else {}
    gp = result.genetic_parameters if isinstance(result.genetic_parameters, dict) else {}

    sigma2_g  = vc.get("sigma2_genotype")
    sigma2_e  = vc.get("sigma2_error")
    sigma2_ge = vc.get("sigma2_ge")
    sigma2_p  = vc.get("sigma2_phenotypic")
    h2        = hp.get("h2_broad_sense")
    gcv       = gp.get("GCV")
    pcv       = gp.get("PCV")
    gam       = gp.get("GAM")
    gam_pct   = gp.get("GAM_percent")
    sel_i     = gp.get("selection_intensity", 1.4)

    # Variance components table
    _vc_g_label  = "Treatment Variance"    if is_agronomy else "Genotypic Variance"
    _vc_ge_label = "T×E Variance"          if is_agronomy else "G×E Variance"
    vc_rows = []
    if sigma2_g  is not None: vc_rows.append([_vc_g_label,            "σ²g",  _fmt(sigma2_g,  4)])
    if sigma2_e  is not None: vc_rows.append(["Error Variance",       "σ²e",  _fmt(sigma2_e,  4)])
    if sigma2_ge is not None: vc_rows.append([_vc_ge_label,           "σ²ge", _fmt(sigma2_ge, 4)])
    if sigma2_p  is not None: vc_rows.append(["Phenotypic Variance",  "σ²p",  _fmt(sigma2_p,  4)])
    if h2        is not None:
        h2_cls = "High" if h2 >= 0.6 else "Moderate" if h2 >= 0.3 else "Low"
        vc_rows.append([f"Heritability, broad-sense (H²) [{h2_cls}]", "H²", _fmt(h2, 4)])

    if vc_rows:
        _add_stat_table(doc, ["Component", "Symbol", "Value"], vc_rows, numeric_cols={2})
        doc.add_paragraph()

    # Formulas
    _add_heading(doc, "Formulas", level=3)
    n_envs = result.n_environments if isinstance(result.n_environments, int) else None
    is_multi_environment = bool(n_envs and n_envs > 1)
    if is_multi_environment:
        formula_lines = [
            "For multi-environment trials:",
            "σ²p = σ²g + (σ²ge / e) + (σ²e / (r × e))",
            "",
            "where:",
            "σ²g  = genotypic variance",
            "σ²ge = genotype × environment interaction variance",
            "σ²e  = error variance",
            "e    = number of environments",
            "r    = number of replications per environment",
            "",
            "H²   = σ²g / σ²p   (broad-sense heritability of genotype means)",
            "GA   = H² × i × σp  (expected gain from selection based on genotype means)",
            "GAM% = (GA / Grand Mean) × 100",
            "σp   = √σ²p (phenotypic standard deviation of genotype means)",
            "i    = selection intensity (Falconer & Mackay, 1996)",
        ]
    else:
        formula_lines = [
            "For single-environment analyses:",
            "σ²p = σ²g + (σ²e / r)",
            "H²  = σ²g / σ²p",
        ]

    for fml in formula_lines:
        p = doc.add_paragraph(fml, style="No Spacing")
        if p.runs:
            p.runs[0].font.name = "Courier New"
            p.runs[0].font.size = Pt(10)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    _add_body(
        doc,
        "Variance components estimated following Comstock and Robinson (1952) and Singh and Chaudhary (1985).",
        italic=True,
    )
    # Selection intensity text is specific to plant breeding reports.
    if not is_agronomy:
        _add_body(doc, _selection_intensity_disclosure(sel_i), italic=True)
    doc.add_paragraph()

    # Genetic advance table
    _ga_heading = "Advance Estimates" if is_agronomy else "Genetic Advance Estimates"
    _ga_label   = "Advance (GA)"                    if is_agronomy else "Genetic Advance (GA)"
    _gam_label  = "Advance as % of Mean (GAM%)"     if is_agronomy else "Genetic Advance as % of Mean (GAM%)"
    ga_rows = []
    if gcv     is not None: ga_rows.append(["GCV (%)",       _fmt(gcv,     2)])
    if pcv     is not None: ga_rows.append(["PCV (%)",       _fmt(pcv,     2)])
    if gam     is not None: ga_rows.append([_ga_label,       _fmt(gam,     4)])
    if gam_pct is not None: ga_rows.append([_gam_label,      _fmt(gam_pct, 2)])

    if ga_rows:
        _add_heading(doc, _ga_heading, level=3)
        _add_stat_table(doc, ["Parameter", "Value"], ga_rows, numeric_cols={1})
        doc.add_paragraph()


# ============================================================================
# SECTION: PATH ANALYSIS
# ============================================================================

def add_path_analysis_section(doc: Document, path_result: Dict[str, Any]) -> None:
    """Add a Path Analysis section to the Word document.

    Args:
        doc: The python-docx Document object
        path_result: Dictionary containing path analysis results with keys:
            - outcome_trait: str
            - predictor_traits: List[str]
            - n_observations: int
            - path_coefficients: List[Dict] with direct_effect, std_error, t_statistic, p_value, significant
            - correlation_decomposition: List[Dict] with total_correlation, direct_effect, indirect_effect, percent_direct
            - r_squared: float
            - residual_effect: float
            - interpretation: str
    """
    if not path_result or not isinstance(path_result, dict):
        return

    _add_heading(doc, "Path Analysis", level=2)

    outcome_trait = path_result.get("outcome_trait", "")
    predictor_traits = path_result.get("predictor_traits", [])
    n_obs = path_result.get("n_observations", 0)
    path_coefs = path_result.get("path_coefficients", [])
    corr_decomp = path_result.get("correlation_decomposition", [])
    r_squared = path_result.get("r_squared", 0)
    residual_effect = path_result.get("residual_effect", 0)
    interpretation = path_result.get("interpretation", "")

    # ── Path Coefficients Table ────────────────────────────────────────────────
    _add_heading(doc, "Path Coefficients", level=3)

    pc_headers = ["Trait", "Direct Effect (β)", "Std Error", "t-statistic", "p-value", "Significant"]
    pc_rows = []

    for coef_dict in path_coefs:
        predictor = coef_dict.get("predictor", "")
        direct = coef_dict.get("direct_effect")
        std_err = coef_dict.get("std_error")
        t_stat = coef_dict.get("t_statistic")
        p_val = coef_dict.get("p_value")
        significant = coef_dict.get("significant", False)

        pc_rows.append({
            "data": [
                predictor,
                _fmt(direct, 3),
                _fmt(std_err, 3),
                _fmt(t_stat, 3),
                _fmt(p_val, 4),
                "Yes*" if significant else "No",
            ],
            "significant": significant,
        })

    # Build table with highlighting for significant rows
    if pc_rows:
        table = doc.add_table(rows=1, cols=len(pc_headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0]
        for i, h in enumerate(pc_headers):
            hdr.cells[i].text = h
        _bold_row(hdr, size_pt=12, bg=_HEADER_BG)

        for row_dict in pc_rows:
            row_data = row_dict["data"]
            is_sig = row_dict["significant"]
            r = table.add_row()
            for i, val in enumerate(row_data):
                r.cells[i].text = str(val)
            # Bold significant rows
            if is_sig:
                for cell in r.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
            _style_data_row(r, numeric_cols={1, 2, 3, 4})

        doc.add_paragraph()

    # ── Correlation Decomposition Table ────────────────────────────────────────
    _add_heading(doc, "Correlation Decomposition", level=3)

    cd_headers = ["Trait", "Total r", "Direct Effect", "Indirect Effect", "% Direct"]
    cd_rows = []

    for decomp_dict in corr_decomp:
        predictor = decomp_dict.get("predictor", "")
        total_r = decomp_dict.get("total_correlation")
        direct = decomp_dict.get("direct_effect")
        indirect = decomp_dict.get("indirect_effect")
        pct_direct = decomp_dict.get("percent_direct")

        cd_rows.append([
            predictor,
            _fmt(total_r, 4),
            _fmt(direct, 4),
            _fmt(indirect, 4),
            _fmt(pct_direct, 2) if pct_direct is not None else "—",
        ])

    if cd_rows:
        _add_stat_table(doc, cd_headers, cd_rows, numeric_cols={1, 2, 3, 4})
        doc.add_paragraph()

    # ── Model Summary ──────────────────────────────────────────────────────────
    _add_heading(doc, "Model Summary", level=3)

    summary_text = (
        f"R² = {_fmt(r_squared, 4)} ({r_squared * 100:.1f}% variation explained); "
        f"Residual path coefficient = {_fmt(residual_effect, 4)}; "
        f"n = {n_obs} observations."
    )
    _add_body(doc, summary_text)
    doc.add_paragraph()

    # ── Interpretation ─────────────────────────────────────────────────────────
    _add_heading(doc, "Interpretation", level=3)

    if interpretation:
        _add_body(doc, interpretation)
    doc.add_paragraph()


# ============================================================================
# SECTION: INTERPRETATION (per trait)
# ============================================================================

def _add_interpretation_section(
    doc: Document,
    ar: GeneticsResponse,
    result: GeneticsResult,
    trait_name: str,
    tr: TraitResult,
    is_anova: bool = False,
    domain: Optional[str] = None,
) -> None:
    is_agronomy = not is_plant_breeding_domain(domain)
    _add_heading(
        doc,
        "Interpretation & Recommendations",
        level=2,
    )

    gp = result.genetic_parameters if isinstance(result.genetic_parameters, dict) else {}

    def _generate_anova_interp() -> Optional[str]:
        """Generate ANOVA-based interpretation on the fly."""
        try:
            from analysis_anova_routes import (
                generate_anova_interpretation,
                get_cv_interpretation_flag,
                is_genotype_effect_significant,
                is_environment_effect_significant,
                is_gxe_effect_significant,
                is_main_plot_significant,
                is_subplot_significant,
                is_interaction_significant,
            )
            ds = result.descriptive_stats or {}
            cv_pct = _resolve_cv_percent(result)
            summary = {
                "grand_mean":     result.grand_mean,
                "cv_percent":     cv_pct,
                "min":            ds.get("min") if isinstance(ds, dict) else None,
                "max":            ds.get("max") if isinstance(ds, dict) else None,
                "range":          ds.get("range") if isinstance(ds, dict) else None,
                "standard_error": ds.get("standard_error") if isinstance(ds, dict) else None,
            }

            # Detect split-plot and extract all required kwargs from result
            _is_split_plot = result.main_plot_mean_separation is not None
            _design_type = "split_plot_rcbd" if _is_split_plot else None
            if _is_split_plot:
                # Extract labels first — needed to resolve remapped ANOVA source names
                _mp_label = getattr(result.main_plot_mean_separation, "treatment_label", None)
                _sp_label = getattr(result.mean_separation, "treatment_label", None)
                mp_sig  = is_main_plot_significant(result.anova_table, label=_mp_label)
                sp_sig  = is_subplot_significant(result.anova_table, label=_sp_label)
                int_sig = is_interaction_significant(result.anova_table, mp_label=_mp_label, sp_label=_sp_label)
                vc = result.variance_components if isinstance(result.variance_components, dict) else {}
                cv_a = _clean_cv_percent(float(vc["cv_A"])) if vc.get("cv_A") is not None else None
                cv_b = _clean_cv_percent(float(vc["cv_B"])) if vc.get("cv_B") is not None else None
                if cv_b is not None:
                    summary["cv_percent"] = cv_b
            else:
                _mp_label = _sp_label = None
                mp_sig = sp_sig = int_sig = None
                cv_a = cv_b = None

            gxe_sig = is_gxe_effect_significant(result.anova_table)
            interpretation = generate_anova_interpretation(
                trait=trait_name,
                summary=summary,
                precision_level=_map_precision_label(cv_pct),
                cv_interpretation_flag=get_cv_interpretation_flag(cv_pct),
                genotype_significant=is_genotype_effect_significant(result.anova_table),
                environment_significant=is_environment_effect_significant(result.anova_table),
                gxe_significant=gxe_sig,
                ranking_caution=gxe_sig,
                selection_feasible=is_genotype_effect_significant(result.anova_table),
                mean_separation=result.mean_separation,
                n_genotypes=result.n_genotypes,
                n_environments=result.n_environments,
                n_reps=result.n_reps,
                domain=domain or "plant_breeding",
                design_type=_design_type,
                main_plot_significant=mp_sig,
                subplot_significant=sp_sig,
                interaction_significant=int_sig,
                cv_a=cv_a,
                cv_b=cv_b,
                main_plot_mean_separation=result.main_plot_mean_separation if _is_split_plot else None,
                mp_label=_mp_label,
                sp_label=_sp_label,
            )
            if is_plant_breeding_domain(domain or "plant_breeding"):
                interpretation = interpretation.replace(
                    "Experimental variability appeared acceptable for treatment comparison under the evaluated conditions.",
                    "Experimental variability appeared acceptable for genotype comparison under the evaluated conditions.",
                )
            return interpretation
        except Exception as _exc:
            logger.warning("Could not generate ANOVA interpretation on-the-fly: %s", _exc)
            return None

    # ── Agronomy: ANOVA-based interpretation only (no H²/GAM genetics text) ──
    if is_agronomy:
        interpretation_text = _generate_anova_interp()
        if interpretation_text:
            _add_heading(doc, "Statistical Interpretation", level=3)
            _add_body(doc, interpretation_text)
            doc.add_paragraph()
        else:
            logger.warning("Agronomy interpretation could not be generated for trait '%s'", trait_name)

        # Practical implication derived from treatment effect significance
        at = result.anova_table
        practical_text = None
        if at and hasattr(at, "source") and "genotype" in at.source:
            idx = at.source.index("genotype")
            p_val = at.p_value[idx] if idx < len(at.p_value) else None
            sig = _sig_label(p_val) if p_val is not None else ""
            ms = result.mean_separation
            top = ms.genotype[0] if ms and ms.genotype else None
            if sig in ("***", "**"):
                practical_text = (
                    f"Treatment differences were highly significant. "
                    + (
                        f"Treatment '{top}' showed comparatively strong performance under the conditions of this experiment "
                        "and may warrant further evaluation across additional environments and seasons before adoption decisions are made."
                        if top else
                        "The leading treatment showed comparatively strong performance under the conditions of this experiment "
                        "and may warrant further evaluation across additional environments and seasons."
                    )
                )
            elif sig == "*":
                practical_text = (
                    f"Treatment differences were significant. "
                    + (
                        f"Treatment '{top}' showed comparatively higher performance among tested treatments; "
                        "results should be validated across additional environments and seasons before operational decisions."
                        if top else
                        "Validate the leading treatment across additional environments and seasons before operational decisions."
                    )
                )
            else:
                practical_text = (
                    "Treatment differences were not statistically significant. "
                    "No single treatment stands out; agronomic or economic criteria may guide management decisions."
                )
        if practical_text:
            _add_heading(doc, "Practical Implication", level=3)
            _add_body(doc, practical_text)
            doc.add_paragraph()
        return

    # ── Plant breeding / ANOVA module: existing genetics-based flow ───────────
    hp = result.heritability if isinstance(result.heritability, dict) else {}
    h2      = hp.get("h2_broad_sense")
    gam_pct = gp.get("GAM_percent")
    generated_implication = None

    if is_anova:
        interpretation_text = None
        if hasattr(tr, 'interpretation') and tr.interpretation:
            interpretation_text = tr.interpretation
            logger.info("Using ANOVA interpretation from trait_result: %d characters", len(interpretation_text))
        elif ar.interpretation:
            interpretation_text = ar.interpretation
            logger.info("Using ANOVA interpretation from analysis_result: %d characters", len(interpretation_text))
        else:
            logger.info("No ANOVA interpretation cached — generating on-the-fly for trait '%s'", trait_name)
            interpretation_text = _generate_anova_interp()
            if interpretation_text:
                logger.info("Generated ANOVA interpretation on-the-fly: %d characters", len(interpretation_text))

        if interpretation_text:
            _add_heading(doc, "Statistical Interpretation", level=3)
            _add_body(doc, interpretation_text)
            doc.add_paragraph()
        else:
            logger.warning("ANOVA interpretation unavailable for trait '%s'", trait_name)
    else:
        gcv_val = gp.get("GCV")
        pcv_val = gp.get("PCV")
        cv_val = _clean_cv_percent((result.descriptive_stats or {}).get("cv_percent") if isinstance(result.descriptive_stats, dict) else None)
        analysis_type = "multi_environment" if (result.n_environments is not None and result.n_environments > 1) else "single_environment"
        anova_f_env, anova_p_env, anova_f_gxe, anova_p_gxe = _anova_env_effect_stats(result.anova_table)
        interpretation_text, generated_implication = generate_genetics_interpretation(
            trait_name=trait_name,
            h2=h2,
            gam=gam_pct,
            gcv=gcv_val,
            pcv=pcv_val,
            anova_f_env=anova_f_env,
            anova_p_env=anova_p_env,
            anova_f_gxe=anova_f_gxe,
            anova_p_gxe=anova_p_gxe,
            cv_percent=cv_val,
            analysis_type=analysis_type,
            domain=domain or "plant_breeding",
        )
        logger.info("Generated genetics interpretation: %d characters", len(interpretation_text))

        if interpretation_text:
            _add_heading(doc, "Statistical Interpretation", level=3)
            _add_body(doc, interpretation_text)
            doc.add_paragraph()

    # Academic interpretation — only render if AI-generated text
    # exists and differs from the rule-based text already written above.
    academic_text = ar.interpretation if ar.interpretation else None
    if academic_text and academic_text != interpretation_text:
        _add_heading(doc, "Academic Interpretation", level=3)
        _add_body(doc, academic_text)
        doc.add_paragraph()

    # Breeding implication
    if not is_anova:
        breeding_text = generated_implication
    else:
        breeding_text = None
        if hasattr(tr, 'breeding_implication') and tr.breeding_implication:
            breeding_text = tr.breeding_implication
            logger.info("Using breeding implication from trait_result: %d characters", len(breeding_text))
        elif result.breeding_implication:
            breeding_text = result.breeding_implication
            logger.info("Using breeding implication from analysis_result: %d characters", len(breeding_text))

    if breeding_text:
        _add_heading(doc, "Research Implications", level=3)
        _add_body(doc, breeding_text)
        doc.add_paragraph()


# ============================================================================
# SECTION: WRITING SUPPORT GUIDE (once per report)
# ============================================================================

def _add_writing_support_guide(doc: Document, data: DownloadReportRequest) -> None:
    _add_heading(doc, "Writing Support Guide", level=1)
    _add_body(
        doc,
        "Use these sentence starters as a scaffold. Edit wording to match your exact study context and supervisor expectations.",
    )
    doc.add_paragraph()

    _add_heading(doc, "Sentence Starters", level=2)

    _guide_domain = getattr(data, "domain", None) or "plant_breeding"
    _guide_agronomy = not is_plant_breeding_domain(_guide_domain)
    _entry = "treatments" if _guide_agronomy else "genotypes"
    _effect = "treatment effect" if _guide_agronomy else "genotype effect"
    _consider = "implementation" if _guide_agronomy else "selection"

    trait_results = data.trait_results or {}
    for row in data.summary_table:
        tr = trait_results.get(row.trait)
        if tr is None or tr.analysis_result is None or tr.analysis_result.result is None:
            continue

        result = tr.analysis_result.result
        at = result.anova_table

        _is_split_plot = result.main_plot_mean_separation is not None

        if _is_split_plot:
            from analysis_anova_routes import (
                _fmt_factor_label,
                is_main_plot_significant,
                is_subplot_significant,
                is_interaction_significant,
            )
            _mp_raw = getattr(result.main_plot_mean_separation, "treatment_label", None)
            _sp_raw = getattr(result.mean_separation, "treatment_label", None) if result.mean_separation else None
            _MP = _fmt_factor_label(_mp_raw, "main-plot factor")
            _SP = _fmt_factor_label(_sp_raw, "subplot factor")
            _mp_sig  = is_main_plot_significant(at, label=_mp_raw)
            _sp_sig  = is_subplot_significant(at, label=_sp_raw)
            _int_sig = is_interaction_significant(at, mp_label=_mp_raw, sp_label=_sp_raw)
            _n_rep   = result.n_reps or data.dataset_summary.n_reps

            # Design + interaction sentence
            _int_word = "statistically significant" if _int_sig else "not significant"
            _add_body(
                doc,
                f"A split-plot RCBD analysis for {row.trait} was conducted with {_MP} assigned to "
                f"main plots and {_SP} assigned to subplots within {_n_rep} replication(s). "
                f"The {_MP} × {_SP} interaction was {_int_word}.",
            )
            # Main effects sentences
            if _int_sig:
                _add_body(
                    doc,
                    f"Because the interaction was significant, interpret {_MP} and {_SP} effects "
                    f"at each specific level of the other factor — do not rely solely on marginal means. "
                    f"Treatment-combination cell means and the interaction plot are the primary reference.",
                )
            else:
                _mp_word = "significant" if _mp_sig else "not significant"
                _sp_word = "significant" if _sp_sig else "not significant"
                _add_body(
                    doc,
                    f"The effect of {_MP} on {row.trait} was {_mp_word}; "
                    f"the effect of {_SP} was {_sp_word}. "
                    f"In the absence of a significant interaction, each factor can be interpreted independently.",
                )

        if not _is_split_plot:
            n_genotypes = result.n_genotypes or data.dataset_summary.n_genotypes
            n_env = result.n_environments or data.dataset_summary.n_environments
            n_rep = result.n_reps or data.dataset_summary.n_reps

            f_g, p_g, _ = _extract_source_stats(at, "genotype")
            eta_g = _eta_squared_for_source(at, "genotype")
            sig_word = "significant" if (p_g is not None and p_g < 0.05) else "not significant"

            env_phrase = (
                f"across {n_genotypes} {_entry} and {n_env} environments"
                if n_env
                else f"across {n_genotypes} {_entry} and {n_rep} replications"
            )

            starter_anova = (
                f"An analysis of variance for {row.trait} evaluated {env_phrase} showed that "
                f"the {_effect} was {sig_word} "
                f"(F = {_fmt(f_g, 3)}, {_p_for_sentence(p_g)}, η² = {_fmt(eta_g, 2, thousands=False)})."
            )
            _add_body(doc, starter_anova)

        if not _guide_agronomy and not _is_split_plot:
            h2 = row.h2
            gam_pct = row.gam_percent
            starter_genetic = (
                f"For {row.trait}, broad-sense heritability was H² = {_fmt(h2, 3, thousands=False)} "
                f"and genetic advance as percent of mean was GAM% = {_fmt(gam_pct, 2, thousands=False)}, "
                "and this should be interpreted within the conditions evaluated in this study."
            )
            _add_body(doc, starter_genetic)

            h2_class = (row.heritability_class or "").lower() if row.heritability_class else ""
            gam_class = (row.gam_class or "").lower() if row.gam_class else ""
            if h2_class == "high" and gam_class == "high":
                _sel_variants = [
                    "Direct phenotypic selection may be effective under the conditions evaluated in this study.",
                    "The trait may respond favorably to phenotypic selection under the evaluated conditions.",
                    "Selection progress may be achievable through direct phenotypic evaluation.",
                ]
                from genetics_interpretation import _select_narrative_variant as _snv
                starter_selection = _snv(row.trait, _sel_variants)
            elif h2_class == "high" and gam_class in {"medium", "moderate"}:
                _sel_variants = [
                    "Moderate selection progress may be achievable through phenotypic evaluation.",
                    "The trait showed moderate expected response to phenotypic selection.",
                ]
                from genetics_interpretation import _select_narrative_variant as _snv
                starter_selection = _snv(row.trait, _sel_variants)
            elif h2_class == "moderate":
                _sel_variants = [
                    "Selection efficiency may partly depend on environmental conditions.",
                    "Environmental influence may contribute to observed phenotypic expression.",
                ]
                from genetics_interpretation import _select_narrative_variant as _snv
                starter_selection = _snv(row.trait, _sel_variants)
            else:
                starter_selection = (
                    "These results may support further evaluation of the highest-performing treatment combinations under additional testing conditions."
                )
            _add_body(doc, starter_selection)

        if result.mean_separation and result.mean_separation.genotype:
            top_genotype = result.mean_separation.genotype[0]
            top_group = result.mean_separation.group[0] if result.mean_separation.group else "a"
            if _is_split_plot:
                from analysis_anova_routes import _fmt_factor_label as _ffl
                _sp_label_ms = _ffl(getattr(result.mean_separation, "treatment_label", None), "subplot factor")
                _mp_label_ms = _ffl(getattr(result.main_plot_mean_separation, "treatment_label", None) if result.main_plot_mean_separation else None, "main-plot factor")
                starter_means = (
                    f"Mean separation for {_sp_label_ms} ({result.mean_separation.test}, α = {result.mean_separation.alpha}, "
                    f"Error B denominator) placed '{top_genotype}' in the highest mean group ({top_group}). "
                    f"These are marginal means averaged across all {_mp_label_ms} levels."
                )
            elif _guide_agronomy:
                starter_means = (
                    f"Mean separation ({result.mean_separation.test}) placed {top_genotype} in the "
                    f"highest observed mean group ({top_group}), suggesting it warrants further evaluation "
                    "under similar management conditions."
                )
            else:
                starter_means = (
                    f"Mean separation ({result.mean_separation.test}) placed {top_genotype} among genotype(s) "
                    f"with the highest observed mean in group '{top_group}'."
                )
            _add_body(doc, starter_means)

        if not _guide_agronomy:
            cv_value = None
            if isinstance(result.descriptive_stats, dict):
                cv_value = _clean_cv_percent(result.descriptive_stats.get("cv_percent"))
            if cv_value is not None:
                _add_body(
                    doc,
                    f"For {row.trait}, CV% was {_fmt_cv(cv_value)}. {_cv_precision_narrative(cv_value, domain='plant_breeding')}",
                )

        doc.add_paragraph()

    _add_heading(doc, "Pre-submission Checklist", level=2)
    _relevance_item = (
        "Report both statistical significance and practical interpretation (e.g., treatment comparison and operational relevance)."
        if _guide_agronomy
        else "Report both statistical significance and practical interpretation (e.g., H², GAM%, and practical relevance)."
    )
    checklist = [
        "Confirm that all reported trait names and units match your tables exactly.",
        "State design context clearly (single environment or multi-environment) before presenting ANOVA outcomes.",
        _relevance_item,
        "Cross-check mean separation group letters with the narrative conclusions about top-performing treatments.",
        "Ensure every recommendation is traceable to a reported result in the table or figure.",
    ]
    for item in checklist:
        doc.add_paragraph(f"□ {item}", style="List Bullet")
    doc.add_paragraph()


# ============================================================================
# SECTION: ASSUMPTION TESTS (per trait, optional)
# ============================================================================

def _add_assumption_tests_section(doc: Document, assumption_tests: Dict[str, Any]) -> None:
    """Render assumption diagnostics. Handles both new (normality/homogeneity/overall)
    and legacy (shapiro_wilk/bartlett) structures from the R engine."""
    _add_heading(doc, "Model Assumptions and Diagnostic Evaluation", level=2)

    _is_new_structure = "overall" in assumption_tests or (
        "normality" in assumption_tests or "homogeneity" in assumption_tests
    )

    if _is_new_structure:
        # Overall sentence
        overall = assumption_tests.get("overall") or {}
        overall_interp = overall.get("interpretation") if isinstance(overall, dict) else None
        reviewer_summary = overall.get("reviewer_summary") if isinstance(overall, dict) else None
        if reviewer_summary:
            _add_body(doc, str(reviewer_summary), bold=True)
        if overall_interp:
            _add_body(doc, overall_interp)
            doc.add_paragraph()

        # Per-test table rows
        rows_data: List[List[str]] = []
        for key in ("normality", "homogeneity"):
            test_result = assumption_tests.get(key)
            if not isinstance(test_result, dict):
                continue
            test_label = test_result.get("test") or key.replace("_", " ").title()
            stat = test_result.get("statistic")
            p_val = test_result.get("p_value") or test_result.get("p.value")
            passed = test_result.get("passed")
            interp = test_result.get("interpretation") or (
                "Passed" if passed is True else ("Failed" if passed is False else "—")
            )
            rows_data.append([
                key.title(),
                test_label,
                _fmt(stat, 4, thousands=False) if stat is not None else "—",
                _fmt_p(p_val) if p_val is not None else "—",
                "Yes" if passed is True else ("No" if passed is False else "—"),
                interp,
            ])

        if rows_data:
            _add_stat_table(
                doc,
                ["Assumption", "Test", "Statistic", "p-value", "Passed", "Interpretation"],
                rows_data,
                numeric_cols={2},
            )
        else:
            _add_body(doc, "Assumption diagnostics were not available for this trait.", italic=True)

        # Split-plot: disclose that assumption tests use pooled residuals and that
        # the two error strata carry separate distributional assumptions (FIX 5).
        stratification_note = assumption_tests.get("stratification_note")
        if stratification_note:
            _add_body(doc, f"Note: {stratification_note}", italic=True)

        outlier_detection = assumption_tests.get("outlier_detection")
        if isinstance(outlier_detection, dict):
            _add_heading(doc, "Outlier and Influence Screening", level=3)
            _add_stat_table(
                doc,
                ["Metric", "Value"],
                [
                    ["|Standardized residual| threshold", _fmt(outlier_detection.get("standardized_residual_threshold"), 2, thousands=False)],
                    ["Cook's distance threshold", _fmt(outlier_detection.get("cooks_distance_threshold"), 4, thousands=False)],
                    ["Extreme outliers detected", str(outlier_detection.get("n_extreme_outliers", 0))],
                    ["Influential observations detected", str(outlier_detection.get("n_influential_observations", 0))],
                ],
            )

            outlier_interp = outlier_detection.get("interpretation")
            if outlier_interp:
                _add_body(doc, str(outlier_interp))

            flagged = outlier_detection.get("flagged_observations")
            if isinstance(flagged, list) and flagged:
                flagged_rows: List[List[str]] = []
                for row in flagged[:25]:
                    if not isinstance(row, dict):
                        continue
                    flagged_rows.append([
                        str(row.get("observation", "—")),
                        str(row.get("treatment", "—")),
                        str(row.get("block", "—")),
                        _fmt(row.get("standardized_residual"), 3, thousands=False),
                        _fmt(row.get("cooks_distance"), 4, thousands=False),
                        "Yes" if bool(row.get("influential")) else "No",
                    ])
                if flagged_rows:
                    _add_stat_table(
                        doc,
                        ["Obs", "Treatment", "Block", "Std Resid", "Cook's D", "Influential"],
                        flagged_rows,
                        numeric_cols={3, 4},
                    )

    else:
        # Legacy structure: shapiro_wilk / bartlett keys with conclusion field
        rows_data = []
        for test_name, test_result in assumption_tests.items():
            label = test_name.replace("_", " ").title()
            if isinstance(test_result, dict):
                p_val = (
                    test_result.get("p_value")
                    or test_result.get("p.value")
                    or test_result.get("p")
                )
                stat = (
                    test_result.get("statistic")
                    or test_result.get("test_stat")
                    or test_result.get("W")
                )
                conclusion = test_result.get("conclusion") or test_result.get("verdict")
                if conclusion is None and p_val is not None:
                    conclusion = "Passed (p ≥ 0.05)" if p_val >= 0.05 else "Failed (p < 0.05)"
                rows_data.append([
                    label,
                    _fmt(stat, 3, thousands=False) if stat is not None else "—",
                    _fmt_p(p_val) if p_val is not None else "—",
                    conclusion or "—",
                ])
            else:
                rows_data.append([label, "—", "—", str(test_result)])

        if rows_data:
            _add_stat_table(
                doc,
                ["Test", "Statistic", "p-value", "Result"],
                rows_data,
                numeric_cols={1},
            )
        else:
            _add_body(doc, "Assumption test data present but could not be parsed.", italic=True)


# ============================================================================
# PER-TRAIT SECTION ORCHESTRATOR
# ============================================================================

def _add_trait_section(
    doc: Document,
    trait_name: str,
    tr: TraitResult,
    row: Optional[SummaryTableRow],
) -> None:
    doc.add_page_break()
    _add_heading(doc, f"Trait: {trait_name}", level=1)

    # Gate only on analysis_result — do NOT trust status field, which can be
    # inferred as "failed" in Pydantic v1 when status is absent in the payload
    # and the validator runs before analysis_result is populated.
    if tr.analysis_result is None:
        error_msg = tr.error or (row.error if row else None) or "No analysis result available"
        _add_body(doc, f"Analysis failed: {error_msg}")
        _add_body(doc, FAILED_TRAIT_CV_MESSAGE)
        logger.warning("Trait '%s' skipped — analysis_result is None, error=%s", trait_name, error_msg)
        return

    ar = tr.analysis_result
    result = ar.result

    if result is None:
        _add_body(doc, "Analysis result object is empty.")
        logger.warning("Trait '%s': analysis_result.result is None", trait_name)
        return

    logger.info(
        "Rendering trait '%s' | anova=%s | mean_sep=%s | h2=%s",
        trait_name,
        result.anova_table is not None,
        result.mean_separation is not None,
        (result.heritability or {}).get("h2_broad_sense"),
    )

    # Data warnings
    if tr.data_warnings:
        _add_heading(doc, "Data Warnings", level=3)
        for w in tr.data_warnings:
            doc.add_paragraph(f"• {w}", style="List Bullet")
        doc.add_paragraph()

    _add_executive_summary(doc, trait_name, result)
    doc.add_paragraph()

    _add_descriptive_stats(doc, result)
    doc.add_paragraph()

    # Determine split-plot status, factor labels, and interaction significance for this trait
    _is_split_plot_ts = getattr(result, "main_plot_mean_separation", None) is not None
    _mp_label_ts = getattr(result.main_plot_mean_separation, "treatment_label", None) if _is_split_plot_ts else None
    _sp_label_ts = getattr(result.mean_separation, "treatment_label", None) if result.mean_separation else None
    _design_type_ts = "split_plot_rcbd" if _is_split_plot_ts else None
    from analysis_anova_routes import is_interaction_significant as _is_int_sig
    _int_sig_ts = bool(_is_int_sig(result.anova_table, mp_label=_mp_label_ts, sp_label=_sp_label_ts)) if result.anova_table else False

    _add_design_statement(
        doc,
        design_type=_design_type_ts,
        n_reps=result.n_reps,
        trait_name=trait_name,
        mp_label=_mp_label_ts,
        sp_label=_sp_label_ts,
    )

    if result.anova_table:
        _add_anova_section(doc, result.anova_table)
    else:
        _add_heading(doc, "Analysis of Variance (ANOVA)", level=2)
        _add_body(doc, "ANOVA table not available for this trait.", italic=True)
    doc.add_paragraph()

    if _is_split_plot_ts or getattr(result, "mean_separation_b", None) is not None:
        _add_key_findings(doc, result, trait_name, mp_label=_mp_label_ts, sp_label=_sp_label_ts)
        doc.add_paragraph()

    if result.mean_separation:
        _add_mean_separation_section(
            doc, trait_name, result.mean_separation,
            factor_name=result.mean_separation.treatment_label,
            interaction_significant=_int_sig_ts,
        )
    else:
        _add_heading(doc, "Mean Separation", level=2)
        _add_body(
            doc,
            "Mean separation (LSD) not available — "
            "insufficient degrees of freedom or singular model.",
            italic=True,
        )
    if result.mean_separation_b:
        doc.add_paragraph()
        _add_mean_separation_section(
            doc, trait_name, result.mean_separation_b,
            factor_name=result.mean_separation_b.treatment_label,
            interaction_significant=_int_sig_ts,
        )
    if getattr(result, "main_plot_mean_separation", None):
        doc.add_paragraph()
        _add_mean_separation_section(
            doc, trait_name, result.main_plot_mean_separation,
            factor_name=getattr(result.main_plot_mean_separation, "treatment_label", None) or "Main-Plot Factor",
            interaction_significant=_int_sig_ts,
        )
    if result.interaction_separation:
        doc.add_paragraph()
        _add_interaction_separation_section(doc, trait_name, result.interaction_separation)
    doc.add_paragraph()

    _add_genetic_parameters_section(doc, result)
    doc.add_paragraph()

    # Log interpretation details before rendering
    tr_interp_len = len(getattr(tr, 'interpretation', '') or '')
    ar_interp_len = len(getattr(ar, 'interpretation', '') or '')
    tr_breeding_len = len(getattr(tr, 'breeding_implication', '') or '')
    logger.info(
        "Interpretation logging for trait '%s': tr.interpretation=%d chars, ar.interpretation=%d chars, tr.breeding_implication=%d chars",
        trait_name, tr_interp_len, ar_interp_len, tr_breeding_len
    )

    _add_interpretation_section(doc, ar, result, trait_name, tr)


# ============================================================================
# GENERIC TABLE HELPER  (DataFrame or list-of-dicts → Word table)
# ============================================================================

def _add_table_to_doc(doc: Document, data: Any) -> None:
    """
    Dynamically build a Word table from a Pandas DataFrame, a list of dicts,
    or a pre-built list of header + row pairs.

    Applies the same grey-header / bordered styling as _add_stat_table so
    the two helpers produce visually consistent output.
    """
    if isinstance(data, pd.DataFrame):
        headers = data.columns.tolist()
        rows = [[str(v) for v in row] for row in data.values.tolist()]
    elif isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
        headers = list(data[0].keys())
        rows = [[str(row.get(h, "")) for h in headers] for row in data]
    else:
        # Fallback: unsupported format — render as plain text
        doc.add_paragraph(str(data))
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
    _bold_row(table.rows[0], size_pt=12, bg=_HEADER_BG)

    # Data rows
    for row_vals in rows:
        r = table.add_row()
        for i, val in enumerate(row_vals):
            r.cells[i].text = val
        _style_data_row(r)


# ============================================================================
# PYDANTIC MODEL → list-of-dicts CONVERTERS
# ============================================================================

def _anova_to_records(at: AnovaTable) -> List[Dict[str, str]]:
    """Convert AnovaTable (parallel arrays) to a list of dicts for _add_table_to_doc."""
    records = []
    for i, src in enumerate(at.source):
        # Suppress Intercept row — not meaningful for report interpretation.
        src_label = str(src).strip().lower()
        if src_label in {"(intercept)", "intercept"}:
            continue

        df_val = at.df[i] if i < len(at.df) else None
        ss_val = at.ss[i] if i < len(at.ss) else None
        ms_val = at.ms[i] if i < len(at.ms) else None
        f_val  = at.f_value[i] if i < len(at.f_value) else None
        p_val  = at.p_value[i] if i < len(at.p_value) else None
        records.append({
            "Source":  _ANOVA_LABELS.get(src, src),
            "DF":      str(int(df_val)) if df_val is not None else "—",
            "SS":      _fmt(ss_val),
            "MS":      _fmt(ms_val),
            "F-value": _fmt(f_val, 3) if f_val is not None else "—",
            "p-value": _fmt_p(p_val),
        })
    return records


def _mean_sep_to_records(ms: MeanSeparation) -> List[Dict[str, str]]:
    """Convert MeanSeparation (parallel arrays) to a list of dicts for _add_table_to_doc."""
    records = []
    for i, geno in enumerate(ms.genotype):
        mean_v = ms.mean[i] if i < len(ms.mean) else None
        se_v   = ms.se[i]   if i < len(ms.se)   else None
        grp    = ms.group[i] if i < len(ms.group) else "—"
        records.append({
            "Rank":     str(i + 1),
            "Genotype": geno,
            "Mean":     _fmt(mean_v),
            "SE":       _fmt(se_v),
            "Group":    grp,
        })
    return records


# ============================================================================
# EXPORT TRAITS TO WORD  (public entry point for per-trait sections)
# ============================================================================

def export_traits_to_word(
    results: DownloadReportRequest,
    doc: Document,
) -> Document:
    """
    Iterate through trait_results and append a complete per-trait section to doc.

    Uses trait_results[trait].status as the primary gate:
      • status == "success" AND analysis_result is not None → full export
      • otherwise → error section with reason

        Per-trait section order:
      1. Executive Summary (grand mean, H², GCV, PCV, GAM)
      2. Descriptive Statistics (real fields only — no fabrication)
      3. ANOVA Table (if anova_table present in result)
      4. Mean Separation table + bar chart (if mean_separation present)
      5. Assumption Tests (if assumption_tests present in result)
      6. Genetic Parameters (variance components, formulas, GA estimates)
            7. Interpretation and domain-appropriate implication/recommendation text
    """
    trait_results = results.trait_results or {}
    domain = getattr(results, "domain", "plant_breeding")

    if not trait_results:
        # Cache recovery found nothing (server restart / token expired).
        # Render what we do know from summary_table so the document is not empty.
        if results.summary_table:
            doc.add_paragraph(
                "Detailed trait sections could not be recovered — the server may have "
                "restarted since the analysis was run. Please re-run the analysis and "
                "download the report immediately afterwards.",
                style="No Spacing",
            )
            doc.add_paragraph()
            for row in results.summary_table:
                status_line = f"{row.trait}: {row.status.upper()}"
                if row.status == "failed" and row.error:
                    status_line += f" — {row.error}"
                doc.add_paragraph(f"• {status_line}", style="List Bullet")
        else:
            doc.add_paragraph("No trait data found in the results.")
        return doc

    for trait, tr in trait_results.items():
        doc.add_page_break()
        
        is_anova = getattr(results, "module", "") == "anova"
        _add_heading(doc, f"Trait: {trait}", level=1)

        # ── Primary gate: status field (inferred from analysis_result when absent)
        if tr.status != "success" or tr.analysis_result is None:
            _add_kv(doc, "Status", "Failed")
            _add_kv(doc, "CV (%)", FAILED_TRAIT_CV_MESSAGE)
            error_msg = tr.error or "No analysis result available."
            # Provide informative diagnostic messages instead of raw error text
            err_lower = (error_msg or "").lower()
            if "residual sum of squares is 0" in err_lower or "ss is 0" in err_lower:
                display_msg = (
                    "Residual variance approached zero, which may indicate highly uniform values, "
                    "duplicated observations, or insufficient variability for ANOVA estimation."
                )
            else:
                display_msg = error_msg
            _add_kv(doc, "Reason", display_msg)
            logger.warning(
                "export_traits_to_word: trait '%s' failed — status=%s error=%s",
                trait, tr.status, tr.error,
            )
            continue

        ar     = tr.analysis_result          # GeneticsResponse
        result = ar.result                   # GeneticsResult

        if result is None:
            # analysis_result parsed but result sub-object is absent — log and skip
            _add_body(doc, "Analysis result structure is incomplete (result object missing).")
            logger.warning(
                "export_traits_to_word: trait '%s' — analysis_result.result is None "
                "despite success status",
                trait,
            )
            continue

        logger.info(
            "export_traits_to_word: rendering trait='%s' | status='%s'\n"
            "  - analysis_result exists: True\n"
            "  - descriptive_stats exists: %s\n"
            "  - assumption_tests exists: %s\n"
            "  - anova_table exists: %s\n"
            "  - mean_separation exists: %s",
            trait, tr.status,
            result.descriptive_stats is not None,
            result.assumption_tests is not None,
            result.anova_table is not None,
            result.mean_separation is not None,
        )

        # ── Data warnings (non-fatal structural notes from the engine) ───────
        if tr.data_warnings:
            _add_heading(doc, "Data Warnings", level=3)
            for w in tr.data_warnings:
                doc.add_paragraph(f"• {w}", style="List Bullet")
            doc.add_paragraph()

        try:
            # ── 1. Executive Summary ─────────────────────────────────────────
            _add_executive_summary(doc, trait, result, is_anova=is_anova, domain=domain)
            doc.add_paragraph()

            # ── 2. Descriptive Statistics ────────────────────────────────────
            _add_descriptive_stats(doc, result, domain=domain)
            doc.add_paragraph()

            # ── 3. ANOVA ─────────────────────────────────────────────────────
            if result.anova_table:
                _add_anova_section(doc, result.anova_table, domain=domain)
            else:
                _add_heading(doc, "Analysis of Variance (ANOVA)", level=2)
                _add_body(doc, "ANOVA table not available for this trait.", italic=True)
            doc.add_paragraph()

            # ── 3b. Key Findings (split-plot / factorial only) ────────────────
            _is_split_plot = getattr(result, "main_plot_mean_separation", None) is not None
            _is_factorial   = getattr(result, "mean_separation_b", None) is not None
            if _is_split_plot or _is_factorial:
                _mp_lbl_kf = getattr(
                    getattr(result, "main_plot_mean_separation", None), "treatment_label", None
                )
                _sp_lbl_kf = getattr(result.mean_separation, "treatment_label", None) if result.mean_separation else None
                _add_key_findings(doc, result, trait, mp_label=_mp_lbl_kf, sp_label=_sp_lbl_kf)
                doc.add_paragraph()

            # ── 4. Mean Separation (table + bar chart) ────────────────────────
            from analysis_anova_routes import is_interaction_significant
            _int_sig = bool(is_interaction_significant(result.anova_table)) if result.anova_table else False
            if result.mean_separation:
                _add_mean_separation_section(
                    doc, trait, result.mean_separation, domain=domain,
                    factor_name=result.mean_separation.treatment_label,
                    interaction_significant=_int_sig,
                )
            elif not _is_split_plot:
                _add_heading(doc, "Mean Separation", level=2)
                _add_body(
                    doc,
                    "Mean separation not available — insufficient degrees of "
                    "freedom or singular model.",
                    italic=True,
                )
            if result.mean_separation_b:
                doc.add_paragraph()
                _add_mean_separation_section(
                    doc, trait, result.mean_separation_b, domain=domain,
                    factor_name=result.mean_separation_b.treatment_label,
                    interaction_significant=_int_sig,
                )
            if getattr(result, "main_plot_mean_separation", None):
                doc.add_paragraph()
                _add_mean_separation_section(
                    doc, trait, result.main_plot_mean_separation, domain=domain,
                    factor_name=getattr(result.main_plot_mean_separation, "treatment_label", None) or "Main-Plot Factor",
                    interaction_significant=_int_sig,
                )
            if result.interaction_separation:
                doc.add_paragraph()
                _add_interaction_separation_section(doc, trait, result.interaction_separation, domain=domain)

            # ── 4b. Split-plot interaction means table + line plot ─────────────
            _int_means = getattr(result, "interaction_means", None)
            if _is_split_plot and _int_means:
                _mp_lbl = getattr(result.main_plot_mean_separation, "treatment_label", None) or "Main-Plot Factor" if getattr(result, "main_plot_mean_separation", None) else "Main-Plot Factor"
                _sp_lbl = getattr(result.mean_separation, "treatment_label", None) or "Subplot Factor" if result.mean_separation else "Subplot Factor"
                doc.add_paragraph()
                _add_interaction_means_section(
                    doc, _int_means,
                    mp_label=_mp_lbl, sp_label=_sp_lbl,
                    trait_name=trait, is_significant=bool(_int_sig),
                    sp_mean_separation=result.mean_separation,
                )
                doc.add_paragraph()
                _add_interaction_plot_to_doc(doc, _int_means, trait, _mp_lbl, _sp_lbl)
            doc.add_paragraph()

            # ── 5. Assumption Tests (optional) ────────────────────────────────
            if result.assumption_tests:
                _add_assumption_tests_section(doc, result.assumption_tests)
                doc.add_paragraph()

            # ── 6. Genetic Parameters (skipped for ANOVA module and agronomy domain) ──
            _is_agronomy_domain = not is_plant_breeding_domain(domain)
            if not is_anova and not _is_agronomy_domain:
                _add_genetic_parameters_section(
                    doc,
                    result,
                    domain=domain,
                )
                doc.add_paragraph()

            # ── 7. Interpretation and domain-appropriate recommendations ─────
            # Log interpretation details before rendering
            tr_interp_len = len(getattr(tr, 'interpretation', '') or '')
            ar_interp_len = len(getattr(ar, 'interpretation', '') or '')
            tr_breeding_len = len(getattr(tr, 'breeding_implication', '') or '')
            logger.info(
                "Interpretation logging for trait '%s' (ANOVA): tr.interpretation=%d chars, ar.interpretation=%d chars, tr.breeding_implication=%d chars",
                trait, tr_interp_len, ar_interp_len, tr_breeding_len
            )
            _add_interpretation_section(
                doc,
                ar,
                result,
                trait,
                tr,
                is_anova=is_anova,
                 domain=domain,
            )

        except Exception as exc:
            logger.error(
                "export_traits_to_word: error in section for trait '%s': %s",
                trait, exc, exc_info=True,
            )
            doc.add_paragraph(
                f"Section rendering error: {type(exc).__name__}: {exc}"
            )

    return doc


# ============================================================================
# FOOTER
# ============================================================================

# Platform citation — printed as the final paragraph of EVERY VivaSense report.
_VIVASENSE_CITATION = (
    "Suggested citation: Fayeun, L. S. (2026). VivaSense Stat (Version 1.0.0) "
    "[Computer software]. Field to Insight Academy. Zenodo. "
    "https://doi.org/10.5281/zenodo.20328141"
)


def _detect_mean_sep_method(data: Any) -> Optional[str]:
    """Return the mean-separation method(s) actually used across all trait results.

    Reads the authoritative ``test`` label emitted by the R engine (e.g.
    "Fisher LSD", "Tukey HSD") rather than assuming one method.  Handles both
    report shapes: ANOVA-export trait results expose the separation objects
    directly, genetics-export trait results nest them under
    ``analysis_result.result``.

    Returns "Fisher LSD" / "Tukey HSD" when a single method was used, the
    generic "mean separation" when several were mixed, or None when no mean
    separation appears in the report (e.g. correlation/regression exports).
    """
    attrs = (
        "mean_separation",
        "main_plot_mean_separation",
        "mean_separation_b",
        "interaction_separation",
    )
    methods: set = set()
    for tr in (getattr(data, "trait_results", None) or {}).values():
        candidates = [tr]
        ar = getattr(tr, "analysis_result", None)
        if ar is not None:
            res = getattr(ar, "result", None)
            if res is not None:
                candidates.append(res)
        for obj in candidates:
            for attr in attrs:
                ms = getattr(obj, attr, None)
                test = getattr(ms, "test", None) if ms is not None else None
                if test:
                    methods.add(str(test))
    # Only mean-separation methods drive the label; ignore descriptive markers.
    methods.discard("Cell Means")
    if not methods:
        return None
    if methods == {"Fisher LSD"}:
        return "Fisher LSD"
    if methods == {"Tukey HSD"}:
        return "Tukey HSD"
    return "mean separation"


def _add_footer(doc: Document, mean_sep_method: Optional[str] = None) -> None:
    # Platform citation as the final paragraph in the document body (every report).
    cite = doc.add_paragraph()
    cite.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cite_run = cite.add_run(_VIVASENSE_CITATION)
    cite_run.font.size = Pt(8)
    cite_run.font.italic = True
    cite_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    cite_run.font.name = "Calibri"

    # Page footer — name the actual mean-separation method (never hardcode Tukey).
    engine_line = (
        f"VivaSense Analysis Engine v1.0  ·  "
        f"Generated {datetime.date.today().strftime('%d %B %Y')}  ·  "
        "Statistical analysis powered by R"
    )
    if mean_sep_method:
        engine_line += f" (ANOVA + {mean_sep_method})"
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(engine_line)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.name = "Calibri"


# ============================================================================
# DOCUMENT BUILDER
# ============================================================================

def _pl(n: Optional[int], singular: str, plural: Optional[str] = None) -> str:
    """Pluralise a count: _pl(1,'block')->'1 block', _pl(3,'block')->'3 blocks'.

    Resolves the '(s)' template markers that were leaking into reports.
    """
    plural = plural or (singular + "s")
    return f"{n} {singular}" if n == 1 else f"{n} {plural}"


def _split_plot_structure(data: Any) -> Optional[Dict[str, Any]]:
    """Return split-plot factor structure for the report header, or None.

    Derived from the per-trait results (the DatasetSummary does not carry factor
    structure).  A report is split-plot when a trait exposes a main-plot mean
    separation.  Labels come from the treatment_label the R engine remapped to
    the user's actual column names; level counts from the separation tables.
    """
    for tr in (getattr(data, "trait_results", None) or {}).values():
        ar = getattr(tr, "analysis_result", None)
        result = getattr(ar, "result", None) if ar is not None else None
        if result is None:
            result = getattr(tr, "result", None) or tr
        mp = getattr(result, "main_plot_mean_separation", None)
        if mp is None:
            continue  # not a split-plot trait
        sp = getattr(result, "mean_separation", None)
        return {
            "mp_label": getattr(mp, "treatment_label", None) or "Main-plot factor",
            "sp_label": getattr(sp, "treatment_label", None) or "Sub-plot factor",
            "n_main": len(getattr(mp, "genotype", None) or []),
            "n_sub": len(getattr(sp, "genotype", None) or []),
        }
    return None


def _build_document(data: DownloadReportRequest) -> Document:
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    # ── Title ─────────────────────────────────────────────────────────────────
    report_domain = getattr(data, "domain", None) or "plant_breeding"
    _is_anova_module = getattr(data, "module", "") == "anova"
    _is_agronomy = not is_plant_breeding_domain(report_domain)
    if _is_anova_module:
        title_text = "VivaSense ANOVA Analysis Report"
    elif _is_agronomy:
        title_text = "VivaSense Agronomy Analysis Report"
    else:
        title_text = "VivaSense Genetics Analysis Report"
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ds = data.dataset_summary
    mode_label = "Multi-environment" if ds.mode == "multi" else "Single-environment"
    env_part = f"  ·  {_pl(ds.n_environments, 'environment')}" if ds.n_environments else ""

    # Split-plot: describe the factor structure (main-plot × sub-plot) instead of
    # a genotype count — the design has no single 'genotype' factor, and the
    # main-plot factor count was previously mislabelled as 'genotypes'.
    _sp = _split_plot_structure(data)
    if _sp:
        _entry_part = (
            f"  ·  {_sp['mp_label']} ({_sp['n_main']} levels)"
            f" × {_sp['sp_label']} ({_sp['n_sub']} levels)"
        )
    else:
        _unit = "treatments" if _is_agronomy else "genotypes"
        _entry_part = (
            f"  ·  {_pl(ds.n_genotypes, _unit[:-1], _unit)}"
            if ds.n_genotypes is not None else ""
        )

    sub = doc.add_paragraph(
        f"{mode_label}{_entry_part}  ·  "
        f"{_pl(ds.n_reps, 'replication')}{env_part}  ·  {_pl(ds.n_traits, 'trait')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if sub.runs:
        sub.runs[0].font.size = Pt(11)

    date_p = doc.add_paragraph(
        f"Report generated: {datetime.date.today().strftime('%d %B %Y')}"
    )
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if date_p.runs:
        date_p.runs[0].font.size = Pt(10)
        date_p.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # ── Cross-trait summary ───────────────────────────────────────────────────
    _add_heading(doc, "Trait Summary", level=1)
    _add_summary_table(doc, data, domain=report_domain)
    doc.add_paragraph()

    all_trait_results = _build_breeding_input_for_export(data)
    synthesis_text = build_breeding_synthesis(
        all_trait_results,
        analysis_type="multi_environment" if ds.mode == "multi" else "single_environment",
    )
    if is_plant_breeding_domain(report_domain) and synthesis_text:
        _add_heading(doc, "Experimental Interpretation Summary", level=1)
        for para in synthesis_text.split("\n\n"):
            _add_body(doc, para)
        doc.add_paragraph()

    if data.failed_traits:
        _add_body(
            doc,
            "Failed traits (excluded from detailed sections): "
            + ", ".join(data.failed_traits),
            italic=True,
        )

    # ── Correlation section (optional) ───────────────────────────────────────
    if data.correlation is not None:
        _add_correlation_section(doc, data.correlation)

    # ── Path Analysis section (optional) ──────────────────────────────────────
    if data.path_analysis is not None:
        add_path_analysis_section(doc, data.path_analysis.model_dump())

    # ── Per-trait sections ────────────────────────────────────────────────────
    logger.info(
        "Building per-trait sections. summary_table=%s | trait_results keys=%s",
        [r.trait for r in data.summary_table],
        list((data.trait_results or {}).keys()),
    )
    export_traits_to_word(data, doc)

    # ── One-time writing support section (all traits) ───────────────────────
    _add_writing_support_guide(doc, data)

    _add_footer(doc, _detect_mean_sep_method(data))
    return doc


# ============================================================================
# CACHE RECOVERY
# ============================================================================

def _recover_from_cache(data: DownloadReportRequest) -> DownloadReportRequest:
    """
    Recover analysis_result objects that the frontend did not include in the
    export POST body.

    Strategy (in order):
      1. If export_token is present → look up by exact token.
      2. Otherwise → find the most-recently-cached response whose trait_results
         keys are a superset of the requested trait names.

    This makes recovery token-independent: even if the frontend never echoes
    the token back (the common case), the cache lookup still succeeds because
    the trait names act as a natural key.

    Returns the (possibly patched) DownloadReportRequest.
    """
    trait_names = list((data.trait_results or {}).keys())

    # Fallback: if the frontend sent an empty trait_results dict, derive the
    # trait names from summary_table (which is always populated).
    if not trait_names and data.summary_table:
        trait_names = [row.trait for row in data.summary_table]
        logger.info(
            "_recover_from_cache: trait_results empty — derived %d trait name(s) "
            "from summary_table for cache lookup",
            len(trait_names),
        )

    # ── Step 1: token lookup ─────────────────────────────────────────────────
    cached = None
    if data.export_token:
        cached = result_cache.get(data.export_token)
        if cached is None:
            logger.warning(
                "_recover_from_cache: token %s not in cache (server restart?), "
                "falling back to trait-name lookup",
                data.export_token,
            )

    # ── Step 2: trait-name fallback ──────────────────────────────────────────
    if cached is None and trait_names:
        cached = result_cache.get_matching_traits(trait_names)

    if cached is None:
        logger.warning(
            "_recover_from_cache: no cached entry found — "
            "per-trait sections will be empty; ask user to re-run analysis"
        )
        return data

    # ── Patch missing analysis_result objects ────────────────────────────────
    patched_trait_results = dict(data.trait_results or {})
    patched_count = 0
    for trait, cached_tr in (cached.trait_results or {}).items():
        current_tr = patched_trait_results.get(trait)
        needs_patch = (
            current_tr is None
            or current_tr.analysis_result is None
        ) and cached_tr.analysis_result is not None

        if needs_patch:
            patched_trait_results[trait] = cached_tr
            patched_count += 1

    if patched_count:
        logger.info(
            "_recover_from_cache: patched analysis_result for %d trait(s) "
            "from cache",
            patched_count,
        )
        data = data.model_copy(update={"trait_results": patched_trait_results})
    else:
        logger.info(
            "_recover_from_cache: all analysis_results already present — "
            "no patching needed"
        )

    return data


# ============================================================================
# ENDPOINT
# ============================================================================

async def export_word_report(data: DownloadReportRequest) -> Response:
    """
    Generate a publication-ready Word (.docx) report from genetics analysis results.

    The request body is the JSON object returned by POST /genetics/analyze-upload,
    optionally extended with a `correlation` field from POST /genetics/correlation.

    Registered at both /genetics/download-results and /genetics/export-word.
    """
    trait_result_keys = list(data.trait_results.keys()) if data.trait_results else []
    summary_traits    = [r.trait for r in data.summary_table]

    logger.info(
        "Download request | summary_traits=%s | trait_result_keys=%s | "
        "failed=%s | has_correlation=%s | export_token=%s",
        summary_traits,
        trait_result_keys,
        data.failed_traits,
        data.correlation is not None,
        data.export_token,
    )

    # ── Recover analysis_result objects from server-side cache ────────────────
    # The frontend sends analysis_result=null for each trait (it stores only
    # display state, not the full GeneticsResponse blob).  _recover_from_cache
    # looks up the server-side cache by trait-name match (no token required)
    # and patches in the missing analysis_result objects before building the doc.
    incoming_domain = data.domain
    data = _recover_from_cache(data)
    if incoming_domain and incoming_domain != (data.domain or "plant_breeding"):
        data = data.model_copy(update={"domain": incoming_domain})

    # Diagnose key-mismatch upfront
    missing_keys = [t for t in summary_traits if t not in trait_result_keys]
    if missing_keys:
        logger.warning(
            "trait_results missing keys for: %s  (present keys: %s)",
            missing_keys,
            trait_result_keys,
        )

    # Log each trait's analysis_result presence after cache recovery
    for trait, tr in (data.trait_results or {}).items():
        has_ar  = tr.analysis_result is not None
        has_res = has_ar and tr.analysis_result.result is not None
        logger.info(
            "  trait='%s' has_analysis_result=%s has_result=%s error='%s'",
            trait,
            has_ar,
            has_res,
            tr.error,
        )

    try:
        doc = _build_document(data)
        report_text = _collect_doc_text(doc)

        # ── TASK 8: Collect wording/governance issues as WARNINGS (not blocks) ──
        # Hard-block is reserved for structural/rendering failures (caught below).
        all_warnings: list[str] = []

        quality_hits = _find_export_quality_hits(report_text, data)
        for hit in quality_hits:
            logger.debug(
                "[EXPORT-QUALITY] Warning triggered | phrase=%r | severity=wording | decision=warn-only",
                hit,
            )
        all_warnings.extend(quality_hits)

        analysis_type = "multi_environment" if getattr(data.dataset_summary, "mode", "") == "multi" else "single_environment"
        if is_plant_breeding_domain(data.domain):
            governance_hits = _find_breeding_governance_hits(report_text, analysis_type=analysis_type)
            for hit in governance_hits:
                logger.debug(
                    "[EXPORT-GOVERNANCE] Warning triggered | phrase=%r | domain=plant_breeding | severity=wording | decision=warn-only",
                    hit,
                )
            all_warnings.extend(governance_hits)
        else:
            forbidden_hits = find_forbidden_breeding_terms(report_text)
            for hit in forbidden_hits:
                logger.debug(
                    "[EXPORT-DOMAIN-GUARD] Warning triggered | phrase=%r | domain=%s | severity=wording | decision=warn-only",
                    hit,
                    data.domain,
                )
            all_warnings.extend(forbidden_hits)

        if all_warnings:
            logger.warning(
                "Export proceeding with %d quality warning(s): %s",
                len(all_warnings),
                all_warnings,
            )

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        response_headers: dict[str, str] = {
            "Content-Disposition": "attachment; filename=vivasense_genetics_report.docx",
        }
        if all_warnings:
            # Expose warnings to frontend via headers for validation-state sync
            response_headers["X-Export-Quality-Warnings"] = str(len(all_warnings))
            response_headers["X-Export-Warning-Detail"] = "; ".join(all_warnings)[:500]
            response_headers["Access-Control-Expose-Headers"] = (
                "X-Export-Quality-Warnings, X-Export-Warning-Detail"
            )

        return Response(
            content=buf.read(),
            media_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            headers=response_headers,
        )
    except Exception as exc:
        logger.error("Report generation failed: %s", exc, exc_info=True)
        return JSONResponse(
            status_code=500,
            content={"detail": f"Report generation failed: {exc}"},
        )


# Register both URL paths
router.add_api_route(
    "/genetics/download-results",
    export_word_report,
    methods=["POST"],
    summary="Download genetics analysis report as Word document",
    tags=["Export"],
)
router.add_api_route(
    "/genetics/export-word",
    export_word_report,
    methods=["POST"],
    summary="Download genetics analysis report as Word document (alias)",
    tags=["Export"],
)
