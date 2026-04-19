from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def generate_descriptive_stats_word_report(results, df, traits, group_variable, interpretations, anova_readiness, warnings, filepath):
    """
    Generates a publication-quality Word document for descriptive statistics.
    """
    doc = Document()
    
    # 1. Title Page
    title = doc.add_heading('VivaSense Descriptive Statistics Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Date Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Sample Size (N): {len(df)}")
    doc.add_paragraph(f"Analyzed Traits: {', '.join(traits)}")
    if group_variable:
        doc.add_paragraph(f"Grouping Variable: {group_variable}")
    
    doc.add_page_break()
    
    # 2. Executive Summary
    doc.add_heading('Executive Summary', level=1)
    exec_summary = doc.add_paragraph("This report outlines the descriptive statistics, data quality checks, and assumption tests for the selected traits. ")
    if warnings:
        exec_summary.add_run("Several statistical warnings were detected that may affect downstream parametric testing (ANOVA).").bold = True
    else:
        exec_summary.add_run("All traits appear suitable for parametric analysis.").bold = True
        
    if warnings:
        doc.add_heading('Key Warnings', level=2)
        for warning in warnings:
            doc.add_paragraph(warning, style='List Bullet')
            
    # 3. Data Overview (Missing Data)
    doc.add_heading('Data Quality Overview', level=1)
    table_missing = doc.add_table(rows=1, cols=4)
    table_missing.style = 'Light Shading Accent 1'
    hdr_cells = table_missing.rows[0].cells
    hdr_cells[0].text = 'Column'
    hdr_cells[1].text = 'Missing (n)'
    hdr_cells[2].text = 'Missing (%)'
    hdr_cells[3].text = 'Pattern'
    
    for col_info in results.get("missing_data", []):
        if col_info["column"] in traits or col_info["column"] == group_variable:
            row_cells = table_missing.add_row().cells
            row_cells[0].text = str(col_info["column"])
            row_cells[1].text = str(col_info["n_missing"])
            row_cells[2].text = f"{col_info['pct_missing']:.2f}%"
            row_cells[3].text = str(col_info["pattern"])

    doc.add_page_break()

    # 4. Descriptive Statistics Table
    doc.add_heading('Descriptive Statistics (Overall)', level=1)
    stat_table = doc.add_table(rows=1, cols=10)
    stat_table.style = 'Light Shading Accent 1'
    headers = ['Trait', 'n', 'Mean', 'SD', 'Min', 'Median', 'Max', 'CV%', 'Skewness', 'Kurtosis']
    for i, header in enumerate(headers):
        stat_table.rows[0].cells[i].text = header
        
    for trait in traits:
        stats = results["summary_stats"][trait]["overall"]
        row = stat_table.add_row().cells
        row[0].text = trait
        row[1].text = str(stats["n"])
        row[2].text = f"{stats['mean']:.4f}" if stats.get("mean") is not None else "-"
        row[3].text = f"{stats['sd']:.4f}" if stats.get("sd") is not None else "-"
        row[4].text = f"{stats['min']:.4f}" if stats.get("min") is not None else "-"
        row[5].text = f"{stats['median']:.4f}" if stats.get("median") is not None else "-"
        row[6].text = f"{stats['max']:.4f}" if stats.get("max") is not None else "-"
        row[7].text = f"{stats['cv_percent']:.2f}%" if stats.get("cv_percent") is not None else "-"
        row[8].text = f"{stats['skewness']:.4f}" if stats.get("skewness") is not None else "-"
        row[9].text = f"{stats['kurtosis']:.4f}" if stats.get("kurtosis") is not None else "-"

    # 5. By-Group Statistics Table
    if group_variable:
        doc.add_heading(f'Descriptive Statistics (By {group_variable})', level=1)
        for trait in traits:
            doc.add_heading(f'Trait: {trait}', level=2)
            group_stats = results["summary_stats"][trait].get("by_group", [])
            
            if group_stats and len(group_stats) > 0:
                g_table = doc.add_table(rows=1, cols=7)
                g_table.style = 'Light Shading Accent 1'
                g_headers = ['Group', 'n', 'Mean', 'SD', 'Min', 'Max', 'CV%']
                for i, header in enumerate(g_headers):
                    g_table.rows[0].cells[i].text = header
                    
                for g_row in group_stats:
                    row = g_table.add_row().cells
                    row[0].text = str(g_row.get(group_variable, "-"))
                    row[1].text = str(g_row.get("n", "-"))
                    row[2].text = f"{g_row.get('mean', 0):.4f}" if g_row.get("mean") is not None else "-"
                    row[3].text = f"{g_row.get('sd', 0):.4f}" if g_row.get("sd") is not None else "-"
                    row[4].text = f"{g_row.get('min', 0):.4f}" if g_row.get("min") is not None else "-"
                    row[5].text = f"{g_row.get('max', 0):.4f}" if g_row.get("max") is not None else "-"
                    row[6].text = f"{g_row.get('cv_percent', 0):.2f}%" if g_row.get("cv_percent") is not None else "-"

    doc.add_page_break()

    # 6. Assumption Tests
    doc.add_heading('Assumption Tests (ANOVA Readiness)', level=1)
    assump_table = doc.add_table(rows=1, cols=5)
    assump_table.style = 'Light Shading Accent 1'
    a_headers = ['Trait', 'Shapiro-Wilk p', 'Normal?', 'Levene p', 'Homogeneous?']
    for i, header in enumerate(a_headers):
        assump_table.rows[0].cells[i].text = header
        
    for trait in traits:
        assumps = results["assumption_tests"][trait]
        norm = assumps["normality"]
        homog = assumps.get("homogeneity")
        
        row = assump_table.add_row().cells
        row[0].text = trait
        row[1].text = f"{norm.get('p_value', 0):.4f}" if norm.get('p_value') is not None else "-"
        row[2].text = "✓ Yes" if norm.get('normal_05') else "✗ No"
        
        if homog and homog.get('p_value') is not None:
            row[3].text = f"{homog.get('p_value', 0):.4f}"
            row[4].text = "✓ Yes" if homog.get('homogeneous_05') else "✗ No"
        else:
            row[3].text = "-"
            row[4].text = "N/A (No Groups)"

    # 7. Outliers Summary
    doc.add_heading('Outliers Summary (IQR Method)', level=1)
    outliers_found = False
    for trait in traits:
        outliers = results["outliers"].get(trait, [])
        if len(outliers) > 0:
            outliers_found = True
            doc.add_paragraph(f"{trait}: {len(outliers)} outliers detected.", style='List Bullet')
            for out in outliers:
                doc.add_paragraph(f"   Row {out['row_index']}: Value {out['value']:.4f} ({out['outlier_type']})", style='List Continue')
    
    if not outliers_found:
        doc.add_paragraph("No statistical outliers detected outside 1.5 * IQR bounds.")

    doc.add_page_break()

    # 8. Interpretation and Recommendations
    doc.add_heading('Trait-by-Trait Interpretations', level=1)
    for trait in traits:
        doc.add_heading(trait, level=2)
        interp = interpretations.get(trait, {})
        doc.add_paragraph(f"Variability: {interp.get('variability', '')}")
        doc.add_paragraph(f"Distribution: {interp.get('distribution', '')}")
        doc.add_paragraph(f"Data Quality: {interp.get('data_quality', '')}")
        
        decision_p = doc.add_paragraph()
        decision_p.add_run(f"ANOVA Recommendation: {anova_readiness.get(trait, '')}").bold = True
        doc.add_paragraph(interp.get('decision_guidance', ''))

    # Footer
    doc.add_paragraph("\n\n---\nReport Generated by VivaSense Statistics Engine")
    
    doc.save(filepath)