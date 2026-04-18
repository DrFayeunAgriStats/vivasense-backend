"""
VivaSense Regression Module — Full QA Suite
============================================

10 datasets spanning the full R² / significance space.
Each dataset uses an independent fixed seed (101-110) so
changing one case never perturbs another.

Consistency validator checks for every case:
  C1  r² ≡ R²           (no mismatch between correlation and model fit)
  C2  equation format    (' × ' present, starts with y_var =)
  C3  strength vs R²     (<0.25 weak, 0.25-0.49 moderate, >=0.50 strong)
  C4  non-sig overrides  (direction/strength locked when p >= 0.05)
  C5  no placeholders    (no {…} in any text field)
  C6  single p-value     (p_value_slope is a finite float in (0, 1])
  C7  reliability flags  (sample_size_ok XOR small_sample; non_sig flag when expected)
"""

import math
import re
import unittest
import numpy as np
import pandas as pd
import statsmodels.api as sm


# ── Pure regression logic (mirrors analysis_regression_routes exactly) ───────

def run_regression(x_arr, y_arr, x_name="X", y_name="Y"):
    x = pd.Series(x_arr, name=x_name, dtype=float)
    y = pd.Series(y_arr, name=y_name, dtype=float)
    mask = x.notna() & y.notna()
    x, y = x[mask], y[mask]
    n = len(x)

    X = sm.add_constant(x)
    results = sm.OLS(y, X).fit()

    intercept     = float(results.params.get("const", 0.0))
    slope         = float(results.params[x_name])
    p_value_slope = float(results.pvalues[x_name])
    se_slope      = float(results.bse[x_name])
    r_squared     = float(results.rsquared)
    r_coef        = math.copysign(math.sqrt(max(0.0, r_squared)), slope)

    is_sig = p_value_slope < 0.05

    if not is_sig:
        direction = "no_clear_relationship"
        strength  = "negligible_or_unreliable"
        plain_effect = (
            "Because the fitted linear relationship is not statistically reliable, "
            "the estimated change per unit increase should not be interpreted as meaningful."
        )
        summary_interp = (
            f"{x_name} does not appear to be a useful linear predictor of {y_name} "
            "in this dataset. Other factors may be more important in explaining "
            "variation in the response."
        )
    else:
        direction = "positive" if slope > 0 else "negative" if slope < 0 else "none"
        strength  = "weak" if r_squared < 0.25 else "moderate" if r_squared < 0.50 else "strong"
        plain_effect = (
            f"For every 1-unit increase in {x_name}, {y_name} "
            + ("increases" if slope > 0 else "decreases")
            + f" by {abs(slope):.4f} units on average."
        )
        summary_interp = (
            "A statistically detectable relationship is present, but the model explains "
            "only a limited proportion of the variation in the response."
            if r_squared < 0.25 else
            "A statistically reliable linear relationship is present, and the model "
            "captures a meaningful proportion of the variation in the response."
        )

    _sign    = "+" if slope >= 0 else "-"
    equation = f"{y_name} = {intercept:.4f} {_sign} {abs(slope):.4f} \u00d7 {x_name}"

    flags, warnings = [], []
    if n < 10:
        warnings.append("Small sample size — results may be unstable.")
        flags.append("small_sample")
    else:
        flags.append("sample_size_ok")
    if r_squared > 0.9 and n < 15:
        warnings.append("Very high model fit with small sample size — may be misleading.")
        flags.append("high_fit_small_n")
    if not is_sig:
        warnings.append("No statistically reliable linear relationship detected.")
        flags.append("non_significant_slope")

    return {
        "n": n, "equation": equation,
        "intercept": intercept, "slope": slope,
        "r_squared": r_squared, "correlation_coefficient": r_coef,
        "p_value_slope": p_value_slope, "se_slope": se_slope,
        "direction": direction, "strength_class": strength,
        "significance_class": "significant" if is_sig else "not_significant",
        "plain_language_effect": plain_effect,
        "summary_interpretation": summary_interp,
        "reliability_flags": flags, "warnings": warnings,
    }


# ── Dataset factory ──────────────────────────────────────────────────────────

def _make(seed, n, slope, noise_sd, x_lo=0, x_hi=100, x_name="X", y_name="Y"):
    """Independent seed per dataset — changing one case never perturbs another."""
    rng = np.random.default_rng(seed)
    x = rng.uniform(x_lo, x_hi, n)
    y = 5.0 + slope * x + rng.normal(0, noise_sd, n)
    return x, y, x_name, y_name


# Verified R² at their calibrated noise levels (all within correct strength band):
#  seed 101 Perfect    R²=1.000 → strong
#  seed 102 Strong+    R²=0.783 → strong
#  seed 103 Strong-    R²=0.728 → strong
#  seed 104 Moderate+  R²=0.307 → moderate
#  seed 105 Moderate-  R²=0.250 → moderate
#  seed 106 Weak sig   R²=0.123 → weak   (n=100 ensures significance)
#  seed 107 Non-sig    R²=0.018 → weak   (p=0.41 → non-significant)
#  seed 108 Near-zero  R²=0.087 → weak   (p=0.11 → non-significant)
#  seed 109 Small n    R²=1.000 → strong  (flagged: small_sample + high_fit_small_n)
#  seed 110 Boundary   R²=0.254 → moderate

DATASETS = {
    1:  ("Perfect linear (R²=1.00)",
         *_make(101, 20,  2.0,    0.05,  0, 100, "temperature",    "pressure"),
         True,  "strong",                "positive"),
    2:  ("Strong positive (R²~0.78)",
         *_make(102, 30,  1.5,   25.0,   0, 100, "rainfall_mm",    "grain_yield"),
         True,  "strong",                "positive"),
    3:  ("Strong negative (R²~0.73)",
         *_make(103, 30, -1.8,   32.0,   0, 100, "plant_density",  "ear_weight"),
         True,  "strong",                "negative"),
    4:  ("Moderate positive (R²~0.31)",
         *_make(104, 40,  0.6,   21.0,   0, 100, "fertilizer_kg",  "shoot_length"),
         True,  "moderate",              "positive"),
    5:  ("Moderate negative (R²~0.25)",
         *_make(105, 40, -0.55,  20.0,   0, 100, "weed_cover_pct", "pod_count"),
         True,  "moderate",              "negative"),
    6:  ("Weak but significant (R²~0.12, n=100)",
         *_make(106, 100, 0.18,  14.0,   0, 100, "days_to_flower", "seed_weight"),
         True,  "weak",                  "positive"),
    7:  ("Weak, non-significant — vine-length (R²~0.02)",
         *_make(107, 40,  0.022,  3.8,  30, 210, "vine-length",    "branches"),
         False, "negligible_or_unreliable", "no_clear_relationship"),
    8:  ("Near-zero, non-significant (R²~0.09)",
         *_make(108, 30,  0.004,  5.0,   0, 100, "leaf_width",     "root_mass"),
         False, "negligible_or_unreliable", "no_clear_relationship"),
    9:  ("Small sample, high fit (n=5, R²=1.00)",
         *_make(109,  5,  3.0,    0.3,   0, 100, "N_dose",         "biomass"),
         True,  "strong",                "positive"),
    10: ("Boundary moderate (R²~0.25)",
         *_make(110, 60,  0.42,  18.0,   0, 100, "stem_diameter",  "fruit_mass"),
         True,  "moderate",              "positive"),
}


# ── Consistency validator ────────────────────────────────────────────────────

def validate(res, exp_sig, exp_strength, exp_dir):
    failures = []

    # C1: r² ≡ R²
    r2_from_r = res["correlation_coefficient"] ** 2
    if abs(r2_from_r - res["r_squared"]) > 1e-9:
        failures.append(
            f"C1 r-R² mismatch: r={res['correlation_coefficient']:.6f} "
            f"→ r²={r2_from_r:.6f} ≠ R²={res['r_squared']:.6f}"
        )

    # C2: equation format "Y = a ± b × X"
    if " \u00d7 " not in res["equation"]:
        failures.append(f"C2 equation missing ' × ': {res['equation']!r}")
    eq_lhs = res["equation"].split(" = ")[0]
    if not eq_lhs:
        failures.append(f"C2 equation missing y_var: {res['equation']!r}")

    # C3: strength vs R² (only checked when significant)
    r2 = res["r_squared"]
    is_sig = res["significance_class"] == "significant"
    if is_sig:
        expected_from_r2 = (
            "weak"     if r2 < 0.25 else
            "moderate" if r2 < 0.50 else
            "strong"
        )
        if res["strength_class"] != expected_from_r2:
            failures.append(
                f"C3 strength mismatch: R²={r2:.4f} → "
                f"expected {expected_from_r2!r}, got {res['strength_class']!r}"
            )

    # C4: non-significant overrides
    if not is_sig:
        if res["direction"] != "no_clear_relationship":
            failures.append(f"C4 non-sig but direction={res['direction']!r}")
        if res["strength_class"] != "negligible_or_unreliable":
            failures.append(f"C4 non-sig but strength={res['strength_class']!r}")
        if "not statistically reliable" not in res["plain_language_effect"]:
            failures.append("C4 non-sig but plain_effect doesn't state unreliability")

    # C5: no {…} placeholders in any text field
    for field in ("equation", "plain_language_effect", "summary_interpretation"):
        if re.search(r"\{[^}]+\}", res[field]):
            failures.append(f"C5 placeholder in {field}: {res[field]!r}")

    # C6: p-value is a single finite float in (0, 1]
    p = res["p_value_slope"]
    if not math.isfinite(p) or not (0 < p <= 1.0):
        failures.append(f"C6 invalid p_value_slope={p}")

    # C7: reliability flags — exactly one of sample_size_ok / small_sample
    has_ok    = "sample_size_ok" in res["reliability_flags"]
    has_small = "small_sample"   in res["reliability_flags"]
    if has_ok == has_small:
        failures.append(
            f"C7 flag conflict: sample_size_ok={has_ok}, small_sample={has_small}"
        )
    if not is_sig and "non_significant_slope" not in res["reliability_flags"]:
        failures.append("C7 non_significant_slope flag missing")

    # Expected values
    act_sig = "significant" if is_sig else "not_significant"
    exp_sig_str = "significant" if exp_sig else "not_significant"
    if act_sig != exp_sig_str:
        failures.append(
            f"Expected sig={exp_sig_str!r}, got {act_sig!r} (p={res['p_value_slope']:.4f})"
        )
    if res["strength_class"] != exp_strength:
        failures.append(
            f"Expected strength={exp_strength!r}, got {res['strength_class']!r} "
            f"(R²={res['r_squared']:.4f})"
        )
    if res["direction"] != exp_dir:
        failures.append(
            f"Expected direction={exp_dir!r}, got {res['direction']!r}"
        )

    return failures


# ── Test class ───────────────────────────────────────────────────────────────

class TestRegressionQA(unittest.TestCase):

    def _run(self, ds_id):
        label, x, y, x_name, y_name, exp_sig, exp_strength, exp_dir = DATASETS[ds_id]
        res = run_regression(x, y, x_name, y_name)
        failures = validate(res, exp_sig, exp_strength, exp_dir)
        return res, label, failures

    def _assert(self, ds_id):
        res, label, failures = self._run(ds_id)
        self.assertEqual(
            failures, [],
            msg=f"\nDataset {ds_id}: {label}\n" + "\n".join(f"  FAIL {f}" for f in failures),
        )
        return res

    def test_01_perfect_linear(self):          self._assert(1)
    def test_02_strong_positive(self):         self._assert(2)
    def test_03_strong_negative(self):         self._assert(3)
    def test_04_moderate_positive(self):       self._assert(4)
    def test_05_moderate_negative(self):       self._assert(5)
    def test_06_weak_significant(self):        self._assert(6)
    def test_07_vine_length_non_significant(self): self._assert(7)
    def test_08_near_zero_non_significant(self):   self._assert(8)

    def test_09_small_sample_flags(self):
        res = self._assert(9)
        self.assertIn("small_sample",     res["reliability_flags"])
        self.assertIn("high_fit_small_n", res["reliability_flags"])
        self.assertTrue(any("unstable" in w for w in res["warnings"]))
        self.assertTrue(any("misleading" in w for w in res["warnings"]))

    def test_10_boundary_moderate(self):       self._assert(10)

    def test_summary_table(self):
        """Emit human-readable QA table (run with -s to see output)."""
        rows, all_pass = [], True
        for ds_id in sorted(DATASETS):
            res, label, failures = self._run(ds_id)
            if failures:
                all_pass = False
            rows.append((ds_id, label, res["n"],
                          res["r_squared"], res["correlation_coefficient"],
                          res["p_value_slope"], res["strength_class"],
                          res["significance_class"],
                          "PASS" if not failures else "FAIL",
                          "; ".join(failures)))

        print("\n")
        print("=" * 110)
        print("  VIVASENSE REGRESSION QA -- 10-DATASET CONSISTENCY REPORT")
        print("=" * 110)
        print(f"  {'#':>2}  {'Dataset':<42} {'n':>4} {'R2':>6} {'r':>7} {'p':>8}  "
              f"{'strength':<25} {'sig':<16} result")
        print("-" * 110)
        for ds_id, label, n, r2, r, p, strength, sig, status, errs in rows:
            print(f"  {ds_id:>2}  {label:<42} {n:>4} {r2:>6.3f} {r:>+7.3f} {p:>8.4f}  "
                  f"{strength:<25} {sig:<16} {status}")
            if errs:
                print(f"      !! {errs}")
        print("-" * 110)
        print(f"  Result: {'ALL 10 PASSED' if all_pass else 'FAILURES DETECTED -- see above'}")
        print("=" * 110)
        self.assertTrue(all_pass, "QA summary: one or more cases failed — see table above")


# ── Word export QA ───────────────────────────────────────────────────────────

class TestRegressionWordExport(unittest.TestCase):
    """
    Builds a Word document from the vine-length regression result,
    saves it to bytes, reloads it, and verifies content fidelity.
    """

    def _vine_result(self):
        _, x, y, x_name, y_name, *_ = DATASETS[7]
        return run_regression(x, y, x_name, y_name), x_name, y_name

    def _build_doc(self, res, x_name, y_name):
        """Mirrors what /export/regression-word produces."""
        from docx import Document
        doc = Document()
        doc.add_heading("Regression Analysis", 0)
        doc.add_heading(f"{y_name} ~ {x_name}", 1)
        doc.add_heading("Fitted Equation", 1)
        doc.add_paragraph(res["equation"])
        doc.add_heading("Model Statistics", 1)
        for label, val in [
            ("n",              str(res["n"])),
            ("R\u00b2",        f"{res['r_squared']:.4f}"),
            ("r",              f"{res['correlation_coefficient']:.4f}"),
            ("slope",          f"{res['slope']:.6f}"),
            ("p (slope only)", f"{res['p_value_slope']:.4f}"),
            ("strength",       res["strength_class"]),
            ("significance",   res["significance_class"]),
        ]:
            doc.add_paragraph(f"{label}: {val}")
        doc.add_heading("Interpretation", 1)
        doc.add_paragraph(res["plain_language_effect"])
        doc.add_paragraph(res["summary_interpretation"])
        if res["warnings"]:
            doc.add_heading("Warnings", 1)
            for w in res["warnings"]:
                doc.add_paragraph(f"\u26a0 {w}")
        return doc

    def test_equation_unicode_roundtrip(self):
        """U+00D7 (×) survives python-docx save/load without corruption."""
        import io
        from docx import Document
        res, x_name, y_name = self._vine_result()
        doc = self._build_doc(res, x_name, y_name)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)
        full_text = "\n".join(p.text for p in doc2.paragraphs)
        self.assertIn("\u00d7", full_text,  "× lost in Word round-trip")
        self.assertIn(x_name,  full_text,  "Predictor name lost in Word round-trip")
        self.assertIn(y_name,  full_text,  "Outcome name lost in Word round-trip")

    def test_interpretation_matches_statistics(self):
        """For a non-sig result the doc must state non-reliability."""
        import io
        from docx import Document
        res, x_name, y_name = self._vine_result()
        self.assertEqual(res["significance_class"], "not_significant",
                         "vine-length fixture should be non-significant")
        doc = self._build_doc(res, x_name, y_name)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        full_text = "\n".join(p.text for p in Document(buf).paragraphs)
        self.assertIn("not statistically reliable", full_text)
        self.assertIn("does not appear to be a useful linear predictor", full_text)

    def test_r_squared_in_doc(self):
        """R² is stated in the statistics section."""
        import io
        from docx import Document
        res, x_name, y_name = self._vine_result()
        doc = self._build_doc(res, x_name, y_name)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        full_text = "\n".join(p.text for p in Document(buf).paragraphs)
        self.assertIn("R\u00b2", full_text)

    def test_no_encoding_corruption(self):
        """×, R², and warning symbol all survive the round-trip without U+FFFD."""
        import io
        from docx import Document
        res, x_name, y_name = self._vine_result()
        doc = self._build_doc(res, x_name, y_name)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        full_text = "\n".join(p.text for p in Document(buf).paragraphs)
        self.assertNotIn("\ufffd", full_text,
                         "U+FFFD replacement character detected — encoding corruption")

    def test_single_p_value_in_doc(self):
        """Only slope p-value appears; no second p-value source."""
        import io
        from docx import Document
        res, x_name, y_name = self._vine_result()
        doc = self._build_doc(res, x_name, y_name)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        full_text = "\n".join(p.text for p in Document(buf).paragraphs)
        # p-value appears exactly once under "p (slope only)"
        p_occurrences = [l for l in full_text.splitlines() if l.startswith("p (slope only)")]
        self.assertEqual(len(p_occurrences), 1,
                         f"Expected exactly one p-value line, found: {p_occurrences}")


if __name__ == "__main__":
    unittest.main(verbosity=2)
