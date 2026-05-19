"""
VivaSense – Module-specific Word export endpoints (VivaSense standard v2)
=========================================================================

POST /export/anova-word              → ANOVA + mean separation per trait
POST /export/genetic-parameters-word → Variance components + parameters per trait
POST /export/correlation-word        → Pairwise table + co-selection advice
POST /export/heatmap-report          → Heatmap image + numeric matrix
POST /export/regression-word         → Regression equation + statistics + interpretation

Design principles (VivaSense standard):
  • Each report contains ONLY the content for its module — no cross-mixing.
  • Per-trait sections end with an Academic Writing Support block (Layer C).
  • Interpretation text is polished academic prose — no raw engine block text.
  • Mean separation charts are truly embedded as 300 DPI PNG images.
  • The legacy combined export (/genetics/download-results) is preserved
    but is NOT the default export path.
"""

import datetime
import io
import logging
import re
from typing import Any, Dict, List, Optional

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fastapi import APIRouter
from fastapi.responses import JSONResponse, Response

from genetics_export import (
    _add_anova_section,
    _add_assumption_tests_section,
    _add_body,
    _add_correlation_section,
    _add_footer,
    _add_heading,
    _add_kv,
    _add_mean_separation_section,
    _add_stat_table,
    _fmt,
    _fmt_p,
    _sig_label,
    _HEADER_BG,
)
from analysis_regression_routes import RegressionResponse
from genetics_schemas import GeneticsResult
from trait_relationships_schemas import CorrelationResponse
from module_schemas import (
    AnovaExportRequest,
    AnovaTraitResult,
    CorrelationExportRequest,
    GeneticParametersExportRequest,
    GeneticParametersTraitResult,
    HeatmapExportRequest,
    DescriptiveResponse,
)
from guided_writing import build_guided_writing
from academic_schemas import GuidedWritingBlock
from genetics_interpretation import _describe_env_effects, _describe_gcv_pcv
from interpretation import InterpretationEngine
from report_validator import ReportValidator
from domain_guard import is_plant_breeding_domain
from report_terms import TERMS, GOVERNANCE, get_report_title

logger = logging.getLogger(__name__)
router = APIRouter(tags=["Export"])

_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# Use the governance-controlled scope note from the shared terms module.
_SCOPE_NOTE = GOVERNANCE.SCOPE_NOTE


# ============================================================================
# VARIANCE COMPONENT COMPUTATION FROM ANOVA TABLE
# ============================================================================

def _compute_variance_components_from_anova(
    anova_table: Any,
    n_reps: Optional[int],
) -> Dict[str, Optional[float]]:
    """
    Derive simple variance components from an ANOVA table's MS values.

    Used by the ANOVA Word export to add a Treatment Variance section
    (σ²g, σ²e, σ²p, Repeatability) after Tukey HSD — mirroring what the UI
    shows in the Treatment Variance panel.
    """
    if anova_table is None or not hasattr(anova_table, "source"):
        return {}

    ms_treatment: Optional[float] = None
    ms_error:     Optional[float] = None

    for i, src in enumerate(anova_table.source):
        src_l = str(src).strip().lower()
        if src_l in {"genotype", "genotypes", "treatment", "treatments"}:
            val = anova_table.ms[i] if i < len(anova_table.ms) else None
            if val is not None:
                ms_treatment = float(val)
        elif src_l in {"residuals", "residual", "error", "error b"}:
            val = anova_table.ms[i] if i < len(anova_table.ms) else None
            if val is not None:
                ms_error = float(val)

    if ms_error is None:
        return {}

    r = max(int(n_reps or 1), 1)
    sigma2_e = ms_error
    sigma2_g: Optional[float] = None
    sigma2_p: Optional[float] = None
    repeatability: Optional[float] = None

    if ms_treatment is not None:
        raw_g = (ms_treatment - ms_error) / r
        sigma2_g = max(0.0, raw_g)   # clamp negative to 0
        sigma2_p = sigma2_g + sigma2_e / r
        if sigma2_p > 0:
            repeatability = sigma2_g / sigma2_p

    return {
        "sigma2_g":     sigma2_g,
        "sigma2_e":     sigma2_e,
        "sigma2_p":     sigma2_p,
        "repeatability": repeatability,
    }


# ============================================================================
# TREATMENT VARIANCE SECTION (BUG FIX: absent from ANOVA Word export)
# ============================================================================

def _add_treatment_variance_section(
    doc: Any,
    trait: str,
    anova_table: Any,
    n_reps: Optional[int],
    domain: str = "plant_breeding",
) -> None:
    """
    Render variance components after the Tukey HSD section.

    BUG FIX (Part 7): The variance component block visible in the UI was
    absent from the Word ANOVA export. This function synchronises both outputs.
    """
    is_agronomy = not is_plant_breeding_domain(domain)
    heading = TERMS.treatment_variance if is_agronomy else TERMS.variance_components
    _add_heading(doc, heading, level=2)

    vc = _compute_variance_components_from_anova(anova_table, n_reps)

    if not vc:
        _add_body(doc, "Variance components not available for this trait.", italic=True)
        doc.add_paragraph()
        return

    rows: List[List[str]] = []
    g_label = "Treatment Variance" if is_agronomy else "Genotypic Variance"

    if vc.get("sigma2_g") is not None:
        flag = " ⚠" if vc["sigma2_g"] == 0.0 else ""
        rows.append([g_label, "σ²g", _fmt(vc["sigma2_g"], 4) + flag])
    if vc.get("sigma2_e") is not None:
        rows.append(["Error Variance", "σ²e", _fmt(vc["sigma2_e"], 4)])
    if vc.get("sigma2_p") is not None:
        rows.append(["Phenotypic Variance", "σ²p", _fmt(vc["sigma2_p"], 4)])
    if vc.get("repeatability") is not None:
        r_label = "Repeatability"
        rows.append([r_label, "H²", _fmt(vc["repeatability"], 4)])

    if rows:
        _add_stat_table(doc, ["Component", "Symbol", "Value"], rows, numeric_cols={2})

    doc.add_paragraph()


# ============================================================================
# STRUCTURED ACADEMIC INTERPRETATION (BUG FIX: UI sub-sections absent in Word)
# ============================================================================

def _add_structured_academic_interpretation(
    doc: Any,
    tr: Any,
    trait: str,
    domain: str = "plant_breeding",
    analysis_type: Optional[str] = None, # Added for SingleEnvironmentGuard
) -> None:
    """
    Render the 5-panel structured interpretation block that mirrors the UI
    AcademicInterpretationPanel.

    BUG FIX (Part 8): The UI shows Overall Finding / Statistical Evidence /
    Treatment Interpretation / Assumption Check / Examiner Checkpoint as
    distinct sub-sections. The Word export previously rendered a flat block.

    Shapiro-Wilk and Levene absence is explicitly stated per the spec.
    """
    _add_heading(doc, TERMS.academic_interpretation, level=2)
    is_agronomy = not is_plant_breeding_domain(domain)

    # ── 1. Overall Finding ────────────────────────────────────────────────────
    _add_heading(doc, TERMS.overall_finding, level=3)
    interp_text = getattr(tr, "interpretation", None) or ""
    first_para = interp_text.split("\n\n")[0].strip() if interp_text else ""
    if first_para:
        _add_body(doc, first_para)
    else:
        _add_body(doc, "Interpretation not available for this trait.", italic=True)
    doc.add_paragraph()

    # ── 2. Statistical Evidence ───────────────────────────────────────────────
    _add_heading(doc, TERMS.statistical_evidence, level=3)
    genotype_sig = getattr(tr, "genotype_significant", None)
    if genotype_sig is not None:
        effect_term = "treatment" if is_agronomy else "genotype"
        sig_word = "significant" if genotype_sig else "not significant"
        _add_body(
            doc,
            f"The {effect_term} effect was {sig_word} "
            f"{GOVERNANCE.WITHIN_SCOPE}.",
        )
    else:
        _add_body(doc, "Statistical significance information not available.", italic=True)
    doc.add_paragraph()

    # ── 3. Treatment Interpretation ───────────────────────────────────────────
    _add_heading(doc, TERMS.treatment_interpretation, level=3)
    ms = getattr(tr, "mean_separation", None)
    if ms and getattr(ms, "genotype", None):
        top   = ms.genotype[0]
        group = ms.group[0] if ms.group else "a"
        _add_body(
            doc,
            f"{top} was placed in group '{group}' — "
            f"{GOVERNANCE.COMPARATIVELY_HIGHER} "
            f"among tested entries {GOVERNANCE.WITHIN_SCOPE}.",
        )
    else:
        _add_body(doc, "Mean separation data not available.", italic=True)
    doc.add_paragraph()

    # ── 4. Assumption Check ───────────────────────────────────────────────────
    _add_heading(doc, TERMS.assumption_check, level=3)
    assumption_tests = getattr(tr, "assumption_tests", None) or {}

    def _extract_assumption(tests: dict, *keys: str) -> Optional[dict]:
        for k in keys:
            result = tests.get(k)
            if result and isinstance(result, dict):
                return result
        return None

    def _format_assumption(test_dict: Optional[dict], label: str) -> str:
        if test_dict is None:
            return GOVERNANCE.SHAPIRO_NOT_REPORTED
        p = (
            test_dict.get("p_value")
            or test_dict.get("p.value")
            or test_dict.get("p")
        )
        verdict = test_dict.get("conclusion") or test_dict.get("verdict")
        if verdict:
            return verdict
        if p is not None:
            outcome = "Passed (p ≥ 0.05)" if p >= 0.05 else "Failed (p < 0.05)"
            return f"{outcome} — {_fmt_p(p)}"
        return GOVERNANCE.SHAPIRO_NOT_REPORTED

    sw = _extract_assumption(assumption_tests, "shapiro_wilk", "normality", "shapiro")
    lv = _extract_assumption(assumption_tests, "levene", "homogeneity", "bartlett")

    _add_body(doc, f"Normality (Shapiro-Wilk): {_format_assumption(sw, 'Shapiro-Wilk')}")
    _add_body(doc, f"Homogeneity of Variance (Levene): {_format_assumption(lv, 'Levene')}")
    doc.add_paragraph()

    # ── 5. Examiner Checkpoint ────────────────────────────────────────────────
    _add_heading(doc, TERMS.examiner_checkpoint, level=3)
    try:
        gw: GuidedWritingBlock = build_guided_writing("anova", trait, tr.model_dump())
        if gw.examiner_checkpoint:
            for item in gw.examiner_checkpoint:
                item_p = doc.add_paragraph(f"☐  {item}")
                item_p.paragraph_format.left_indent = Inches(0.3)
                item_p.paragraph_format.space_before = Pt(1)
                item_p.paragraph_format.space_after  = Pt(1)
                for run in item_p.runs:
                    run.font.size = Pt(10)
        else:
            _add_body(doc, "Examiner checkpoint not available.", italic=True)
    except Exception as exc:
        logger.warning("Examiner checkpoint failed for %s: %s", trait, exc)
        _add_body(doc, "Examiner checkpoint not available.", italic=True)

    doc.add_paragraph()

    # Single Environment Guard: Force disclaimer
    if _is_single_environment_analysis(analysis_type):
        _add_body(doc, GOVERNANCE.SINGLE_ENVIRONMENT_DISCLAIMER, italic=True)


# ============================================================================
# DOCUMENT HELPERS
# ============================================================================

def _new_document(title: str, subtitle: str = "") -> Document:
    """Create a new Document with standard VivaSense margins and centred title."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sub.runs:
            sub.runs[0].font.size = Pt(11)

    date_p = doc.add_paragraph(
        f"Generated: {datetime.date.today().strftime('%d %B %Y')}"
    )
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if date_p.runs:
        date_p.runs[0].font.size = Pt(10)
        date_p.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()
    return doc


def _docx_response(doc: Document, filename: str) -> Response:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type=_DOCX_MIME,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


def _scope_paragraph(doc: Document) -> None:
    """Add the standard scope / limitation note."""
    p = doc.add_paragraph(_SCOPE_NOTE)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


# ============================================================================
# INTERPRETATION TEXT CLEANER
# ============================================================================

def _clean_interpretation(raw: Optional[str], trait: str = "") -> str:
    """
    Convert raw R-engine block text into polished academic prose.

    Strips:
      • "Trait: (not specified)" / "Trait: Trait"
      • Section headers like "ANOVA Result:", "Interpretation:"
      • Duplicate blank lines
      • Leading/trailing whitespace per line

    Returns a clean paragraph-formatted string, or "" if input is None/empty.
    """
    if not raw or not raw.strip():
        return ""

    text = raw

    # Remove "Trait: (not specified)" or "Trait: Trait" artifacts
    text = re.sub(
        r"^Trait:\s*(not\s+specified|Trait)[\s\n]*",
        "", text, flags=re.IGNORECASE | re.MULTILINE,
    )

    # Remove raw engine section labels (e.g. "ANOVA Result:", "Interpretation:")
    text = re.sub(
        r"^(ANOVA\s+Result|Interpretation|Genetic\s+Parameters?|Heritability"
        r"|Breeding\s+Implication|Statistical\s+Output)\s*:?\s*\n",
        "", text, flags=re.IGNORECASE | re.MULTILINE,
    )

    # Collapse 3+ blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Strip trailing whitespace from each line
    lines = [ln.rstrip() for ln in text.splitlines()]
    text = "\n".join(lines).strip()

    # If after cleaning the text still starts with the trait name verbatim,
    # leave it — it's now context, not a label.
    return text


# ============================================================================
# ACADEMIC WRITING SUPPORT SECTION (Layer C)
# ============================================================================

def _add_writing_support_section(
    doc: Document,
    module_type: str,
    trait: Optional[str],
    result_dict: Dict[str, Any],
    analysis_type: Optional[str] = None, # Added for SingleEnvironmentGuard
) -> None:
    """
    Build and render the Layer-C guided writing block for one trait.

    Adds to *doc*:
      • Section heading "Academic Writing Support"
      • Per-sentence-starter box: purpose label + template + fill hints
      • Examiner checklist
      • Scope and supervisor note
    """
    try:
        gw: GuidedWritingBlock = build_guided_writing(module_type, trait, result_dict)
    except Exception as exc:
        logger.warning("build_guided_writing failed for %s/%s: %s", module_type, trait, exc)
        return

    _add_heading(doc, TERMS.writing_support, level=2)

    # Caution note (e.g. low-rep warning)
    if gw.caution_note:
        p = doc.add_paragraph(f"⚠ {gw.caution_note}")
        p.paragraph_format.space_before = Pt(2)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xB8, 0x68, 0x00)
        doc.add_paragraph()

    # Sentence starters
    if gw.sentence_starters:
        intro = doc.add_paragraph(
            "Complete the sentences below using values from your own analysis. "
            "Every blank (___) must be filled by you — do not submit this text directly."
        )
        intro.paragraph_format.space_after = Pt(6)
        for run in intro.runs:
            run.italic = True
            run.font.size = Pt(10)

        for i, s in enumerate(gw.sentence_starters, 1):
            # Purpose label
            purpose_p = doc.add_paragraph()
            run_label = purpose_p.add_run(f"Sentence {i} — {s.purpose}")
            run_label.bold = True
            run_label.font.size = Pt(11)
            purpose_p.paragraph_format.space_before = Pt(6)
            purpose_p.paragraph_format.space_after  = Pt(2)

            # Template (monospace-style)
            tpl_p = doc.add_paragraph(s.template)
            tpl_p.paragraph_format.left_indent = Inches(0.3)
            tpl_p.paragraph_format.space_before = Pt(1)
            tpl_p.paragraph_format.space_after  = Pt(2)
            for run in tpl_p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(10)

            # Fill hints
            if s.values_to_fill:
                for j, fill_hint in enumerate(s.values_to_fill, 1):
                    hint_p = doc.add_paragraph(f"  Fill {j}: {fill_hint}")
                    hint_p.paragraph_format.left_indent = Inches(0.5)
                    hint_p.paragraph_format.space_before = Pt(0)
                    hint_p.paragraph_format.space_after  = Pt(1)
                    for run in hint_p.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

            if s.hint:
                src_p = doc.add_paragraph(f"  Source: {s.hint}")
                src_p.paragraph_format.left_indent = Inches(0.5)
                src_p.paragraph_format.space_before = Pt(0)
                src_p.paragraph_format.space_after  = Pt(3)
                for run in src_p.runs:
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()

    # Examiner checklist
    if gw.examiner_checkpoint:
        ck_heading = doc.add_paragraph()
        run_ck = ck_heading.add_run("Examiner Checklist — tick each item before submission:")
        run_ck.bold = True
        run_ck.font.size = Pt(11)
        ck_heading.paragraph_format.space_before = Pt(6)

        for item in gw.examiner_checkpoint:
            item_p = doc.add_paragraph(f"☐  {item}")
            item_p.paragraph_format.left_indent = Inches(0.3)
            item_p.paragraph_format.space_before = Pt(1)
            item_p.paragraph_format.space_after  = Pt(1)
            for run in item_p.runs:
                run.font.size = Pt(10)

        doc.add_paragraph()

    # Supervisor prompt
    sup_p = doc.add_paragraph(gw.supervisor_prompt)
    sup_p.paragraph_format.space_before = Pt(4)
    for run in sup_p.runs:
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x55, 0x33)


# ============================================================================
# GENETIC PARAMETERS — focused table builder (standalone, no GeneticsResult)
# ============================================================================

def _add_gp_tables(
    doc: Document,
    tr: GeneticParametersTraitResult,
    domain: str = "plant_breeding",
) -> None:
    """
    Render variance components + heritability + GCV/PCV/GA/GAM tables
    directly from a GeneticParametersTraitResult without the GeneticsResult
    adapter.  This is the VivaSense standard layout for the GP report.

    DOMAIN GUARD: GCV, PCV, Genetic Advance, selection intensity, and
    Falconer & Mackay references are plant_breeding-specific and are
    suppressed for agronomy/general domains.
    """
    is_agronomy = not is_plant_breeding_domain(domain)
    vc = tr.variance_components or {}
    hp = tr.heritability or {}

    h2 = hp.get("h2_broad_sense")
    h2_class = (
        "High"     if h2 is not None and h2 >= 0.60
        else "Moderate" if h2 is not None and h2 >= 0.30
        else "Low"      if h2 is not None
        else "—"
    )

    # ── 1. Variance Components ────────────────────────────────────────────────
    _add_heading(doc, TERMS.variance_components, level=2)
    g_label  = "Treatment Variance" if is_agronomy else "Genotypic Variance"
    ge_label = "T×E Variance"       if is_agronomy else "G×E Variance"
    _vc_keys = [
        ("sigma2_genotype",   g_label,              "σ²g"),
        ("sigma2_error",      "Error Variance",      "σ²e"),
        ("sigma2_ge",         ge_label,              "σ²ge"),
        ("sigma2_phenotypic", "Phenotypic Variance", "σ²p"),
    ]
    vc_rows: List[List[str]] = []
    for key, label, sym in _vc_keys:
        val = vc.get(key)
        if val is not None:
            flag = " ⚠" if (isinstance(val, (int, float)) and val < 0) else ""
            vc_rows.append([label, sym, _fmt(val, 6) + flag])
    if vc_rows:
        _add_stat_table(doc, ["Component", "Symbol", "Value"], vc_rows, numeric_cols={2})
    else:
        _add_body(doc, "Variance components not available.", italic=True)
    doc.add_paragraph()

    # ── 2. Heritability / Repeatability ──────────────────────────────────────
    h2_heading = "Repeatability" if is_agronomy else TERMS.heritability
    _add_heading(doc, h2_heading, level=2)
    if h2 is not None:
        h2_label_row = "Repeatability (H²)" if is_agronomy else "Broad-sense Heritability (H²)"
        h2_rows = [
            [h2_label_row, _fmt(h2, 4)],
            ["Classification", h2_class],
            ["Basis", hp.get("interpretation_basis", "—")],
        ]
        _add_stat_table(doc, ["Parameter", "Value"], h2_rows, numeric_cols={1})
        doc.add_paragraph()

        if is_agronomy:
            # Agronomy narrative — no genetics/selection language
            if h2_class == "High":
                _add_body(
                    doc,
                    f"Repeatability was H² = {_fmt(h2, 3)} (high) in this experiment, "
                    "suggesting that treatment effects were consistently expressed "
                    "across replications under the evaluated conditions.",
                )
            elif h2_class == "Moderate":
                _add_body(
                    doc,
                    f"Repeatability was H² = {_fmt(h2, 3)} (moderate) in this experiment. "
                    "Both treatment and environmental replication effects contributed "
                    "to observed variation.",
                )
            else:
                _add_body(
                    doc,
                    f"Repeatability was H² = {_fmt(h2, 3)} (low) in this experiment, "
                    "indicating that environmental variation within the experiment "
                    "accounted for a large share of observed phenotypic variation.",
                )
        else:
            # Plant breeding narrative
            if h2_class == "High":
                _add_body(
                    doc,
                    f"Broad-sense heritability for this trait was estimated at "
                    f"H² = {_fmt(h2, 3)} (high) in this experiment, suggesting "
                    "that a substantial proportion of observed phenotypic variation "
                    "is attributable to genetic differences among genotypes under "
                    "these experimental conditions.",
                )
            elif h2_class == "Moderate":
                _add_body(
                    doc,
                    f"Broad-sense heritability was H² = {_fmt(h2, 3)} (moderate) "
                    "in this experiment. Both genetic and environmental effects "
                    "contributed to phenotypic variation; multi-environment "
                    "evaluation is advisable before drawing selection conclusions.",
                )
            else:
                _add_body(
                    doc,
                    f"Broad-sense heritability was H² = {_fmt(h2, 3)} (low) "
                    "in this experiment, indicating that environmental effects "
                    "accounted for a large share of observed phenotypic variation. "
                    "Phenotypic selection is unlikely to be efficient under these conditions.",
                )
    else:
        _add_body(doc, "Heritability/repeatability estimate not available.", italic=True)
    doc.add_paragraph()

    # ── 3. GCV and PCV (plant_breeding only) ─────────────────────────────────
    # DOMAIN GUARD: GCV/PCV are genetics-specific — must not appear in agronomy reports.
    if not is_agronomy:
        _add_heading(doc, TERMS.coefficients_variation, level=2)
        gcv, pcv = tr.gcv, tr.pcv
        cv_rows: List[List[str]] = []
        if gcv is not None:
            cv_rows.append(["Genotypic Coefficient of Variation (GCV %)", _fmt(gcv, 2)])
        if pcv is not None:
            cv_rows.append(["Phenotypic Coefficient of Variation (PCV %)", _fmt(pcv, 2)])
        if cv_rows:
            _add_stat_table(doc, ["Parameter", "Value"], cv_rows, numeric_cols={1})
            doc.add_paragraph()
            if gcv is not None and pcv is not None:
                diff = abs(gcv - pcv)
                env_sentence = None
                if diff < 1.0:
                    env_sentence = _describe_env_effects(
                        f_env=float(tr.anova_f_env) if tr.anova_f_env is not None else 0.0,
                        p_env=float(tr.anova_p_env) if tr.anova_p_env is not None else None,
                        f_gxe=float(tr.anova_f_gxe) if tr.anova_f_gxe is not None else 0.0,
                        p_gxe=float(tr.anova_p_gxe) if tr.anova_p_gxe is not None else None,
                        analysis_type=(
                            "multi_environment"
                            if (tr.n_environments is not None and tr.n_environments > 1)
                            else "single_environment"
                        ),
                    )
                    cv_comment = _describe_gcv_pcv(gcv, pcv, tr.trait)
                elif gcv < pcv:
                    cv_comment = (
                        f"GCV ({_fmt(gcv, 2)}%) < PCV ({_fmt(pcv, 2)}%) — "
                        "environmental effects contributed to observed phenotypic variation "
                        f"in this experiment (PCV − GCV = {_fmt(diff, 2)}%)."
                    )
                else:
                    cv_comment = (
                        f"GCV ({_fmt(gcv, 2)}%) > PCV ({_fmt(pcv, 2)}%) — "
                        "verify variance component estimates; this relationship is unusual."
                    )
                _add_body(doc, cv_comment)
                if env_sentence:
                    doc.add_paragraph()
                    _add_body(doc, env_sentence)
        else:
            _add_body(doc, "GCV/PCV not available.", italic=True)
        doc.add_paragraph()

    # ── 4. Genetic Advance (plant_breeding only) ──────────────────────────────
    # DOMAIN GUARD: Genetic Advance, selection intensity, i = 1.40, and
    # Falconer & Mackay must NOT appear in agronomy or general reports.
    if not is_agronomy:
        _add_heading(doc, TERMS.genetic_advance, level=2)
        ga, gam = tr.ga, tr.gam
        ga_rows: List[List[str]] = []
        if ga is not None:
            ga_rows.append(["Genetic Advance (GA, absolute)", _fmt(ga, 4)])
        if gam is not None:
            gam_class = InterpretationEngine.classify_gam(gam)
            gam_class = gam_class.capitalize() if gam_class != "not_computed" else "—"
            ga_rows.append([
                "Genetic Advance as % of Mean (GAM %)",
                f"{_fmt(gam, 2)}  [{gam_class}]",
            ])
        if ga_rows:
            _add_stat_table(doc, ["Parameter", "Value"], ga_rows, numeric_cols={1})
            doc.add_paragraph()
            _add_heading(doc, TERMS.formulas, level=3)
            for fml in [
                "GA  = H² × i × σp",
                "GAM (%) = (GA / Grand Mean) × 100",
                "Where: i = selection intensity (Falconer & Mackay, 1996)",
                "       σp = √σ²p (phenotypic standard deviation)",
            ]:
                fml_p = doc.add_paragraph(fml, style="No Spacing")
                if fml_p.runs:
                    fml_p.runs[0].font.name = "Courier New"
                    fml_p.runs[0].font.size = Pt(10)
        else:
            _add_body(doc, "Genetic advance estimates not available.", italic=True)
        doc.add_paragraph()


# ============================================================================
# POST /export/anova-word
# ============================================================================

@router.post(
    "/export/anova-word",
    summary="Export ANOVA results as a Word document",
    tags=["Export"],
)
async def export_anova_word(data: AnovaExportRequest):
    """
    VivaSense ANOVA report — one section per trait (section order matches UI):
      1. Descriptive Statistics
      2. ANOVA Table (with significance narrative, Intercept row included)
      3. Assumption Tests
      4. Mean Separation (table + embedded bar chart)
      5. Treatment Variance (σ²g, σ²e, σ²p, Repeatability)  ← BUG FIX
      6. Academic Interpretation (5 structured sub-sections)  ← BUG FIX
      7. Academic Writing Support (sentence starters + examiner checklist)
      8. Scope note
    """
    try:
        domain     = getattr(data, "domain", "plant_breeding") or "plant_breeding"
        n_traits   = len(data.trait_results)
        n_success  = sum(1 for tr in data.trait_results.values() if tr.status == "success")
        mode_label = "Multi-environment" if data.mode == "multi" else "Single-environment"
        analysis_type = "multi_environment" if data.mode == "multi" else "single_environment"

        logger.info(
            "ANOVA export started: %d traits, %d successful, domain=%s",
            n_traits, n_success, domain,
        )
        for trait, tr in data.trait_results.items():
            logger.info(
                "  trait '%s': status=%s, has_interpretation=%s, interpretation_len=%d",
                trait, tr.status,
                "YES" if tr.interpretation else "NO",
                len(tr.interpretation) if tr.interpretation else 0,
            )

        # Initialize ReportValidator
        validator = ReportValidator()

        # BUG FIX: domain-aware title (previously hardcoded "VivaSense ANOVA Analysis Report")
        doc = _new_document(
            get_report_title(domain, module="anova"),
            f"{mode_label}  ·  {n_traits} trait(s)  ·  {n_success} analysed successfully",
        )

        if data.failed_traits:
            _add_body(
                doc,
                "Traits excluded due to analysis failure: " + ", ".join(data.failed_traits),
                italic=True,
            )
            doc.add_paragraph()

        for trait, tr in data.trait_results.items():
            doc.add_page_break()
            _add_heading(doc, trait, level=1)

            if tr.status != "success" or tr.anova_table is None:
                _add_body(
                    doc,
                    f"Analysis failed for this trait: {tr.error or 'No ANOVA result available.'}",
                )
                continue
            
            # Validate interpretation before adding to report
            validation_errors = validator.validate_anova_interpretation(
                interpretation_text=tr.interpretation or "",
                analysis_type=analysis_type,
                interaction_significant=tr.interaction_significant # Assuming this field exists or can be derived
            )
            all_warnings.extend(validation_errors)

            # ── Data warnings ──────────────────────────────────────────────────
            if tr.data_warnings:
                _add_heading(doc, TERMS.data_warnings, level=3)
                for w in tr.data_warnings:
                    doc.add_paragraph(f"• {w}", style="List Bullet")
                doc.add_paragraph()

            # ── 1. Descriptive Statistics ──────────────────────────────────────
            _add_heading(doc, TERMS.descriptive_stats, level=2)
            rows = []
            is_agronomy = not is_plant_breeding_domain(domain)
            entry_label = "No. Treatments" if is_agronomy else "No. Genotypes"
            if tr.grand_mean is not None:
                rows.append(("Grand Mean", _fmt(tr.grand_mean)))
            if tr.n_genotypes:
                rows.append((entry_label, str(tr.n_genotypes)))
            if tr.n_reps:
                rows.append(("No. Replications", str(tr.n_reps)))
            if tr.n_environments:
                rows.append(("No. Environments", str(tr.n_environments)))
            if tr.descriptive_stats:
                ds = tr.descriptive_stats
                if isinstance(ds, dict):
                    for k, v in ds.items():
                        label = k.replace("_", " ").title()
                        rows.append((label, _fmt(v) if isinstance(v, float) else str(v)))
                else:
                    if ds.standard_deviation is not None:
                        rows.append(("Standard Deviation", _fmt(ds.standard_deviation)))
                    if ds.standard_error is not None:
                        rows.append(("Standard Error", _fmt(ds.standard_error)))
                    if ds.min is not None:
                        rows.append(("Minimum", _fmt(ds.min)))
                    if ds.max is not None:
                        rows.append(("Maximum", _fmt(ds.max)))
                    if ds.range is not None:
                        rows.append(("Range", _fmt(ds.range)))
                    if ds.cv_percent is not None:
                        rows.append(("Coefficient of Variation (%)", _fmt(ds.cv_percent, 2)))
                    if ds.variance is not None:
                        rows.append(("Variance", _fmt(ds.variance)))
            if rows:
                _add_stat_table(doc, ["Parameter", "Value"], rows, numeric_cols={1})
            doc.add_paragraph()

            # ── 2. ANOVA Table ─────────────────────────────────────────────────
            _add_anova_section(doc, tr.anova_table, domain=domain)
            doc.add_paragraph()

            # ── 3. Assumption Tests ────────────────────────────────────────────
            if tr.assumption_tests:
                _add_assumption_tests_section(doc, tr.assumption_tests)
                doc.add_paragraph()

            # ── 4. Mean Separation (table + embedded chart) ────────────────────
            # Skip entirely for split-plot RCBD — mean separation over a single
            # factor is not meaningful when effects are estimated in two strata.
            if tr.design_type == "split_plot_rcbd":
                pass  # section omitted — not applicable for this design
            elif tr.mean_separation:
                _add_mean_separation_section(doc, trait, tr.mean_separation, domain=domain)
                doc.add_paragraph()
            else:
                _add_heading(doc, TERMS.mean_separation, level=2)
                _add_body(
                    doc,
                    "Mean separation (Tukey HSD) is not available for this trait — "
                    "insufficient degrees of freedom or a singular model.",
                    italic=True,
                )
                doc.add_paragraph()

            # ── 5. Treatment Variance (BUG FIX: was absent from Word export) ──
            _add_treatment_variance_section(
                doc, trait, tr.anova_table, tr.n_reps, domain=domain
            )

            # ── 6. Academic Interpretation (structured 5-panel) ────────────────
            logger.info("ANOVA export: processing interpretation for trait '%s'", trait)
            logger.info("  Raw interpretation length: %d chars", len(tr.interpretation or ""))
            cleaned = _clean_interpretation(tr.interpretation, trait)
            logger.info("  Cleaned interpretation length: %d chars", len(cleaned))
            if cleaned:
                _add_heading(doc, TERMS.interpretation, level=2)
                for para_text in cleaned.split("\n\n"):
                    para_text = para_text.strip()
                    if para_text:
                        _add_body(doc, para_text)
                doc.add_paragraph()
            else:
                logger.warning(
                    "ANOVA export: interpretation is empty for trait '%s' (raw was: %s)",
                    trait, "None" if tr.interpretation is None else "empty string",
                )

            # Structured academic interpretation (mirrors UI AcademicInterpretationPanel)
            _add_structured_academic_interpretation(doc, tr, trait, domain=domain)

            # ── 7. Academic Writing Support ────────────────────────────────────
            _add_writing_support_section(doc, "anova", trait, tr.model_dump(), analysis_type=analysis_type)
            doc.add_paragraph()

            # ── 8. Scope note ──────────────────────────────────────────────────
            _scope_paragraph(doc)

        # This was missing in the original
        _add_footer(doc)
        return _docx_response(doc, "vivasense_anova_report.docx")

    except Exception as exc:
        logger.error("ANOVA export failed: %s", exc, exc_info=True)
        return JSONResponse(
            status_code=500,
            content={"detail": f"ANOVA export failed: {exc}"},
        )


# ============================================================================
# POST /export/genetic-parameters-word
# ============================================================================

@router.post(
    "/export/genetic-parameters-word",
    summary="Export Genetic Parameters results as a Word document",
    tags=["Export"],
)
async def export_genetic_parameters_word(data: GeneticParametersExportRequest):
    """
    VivaSense Genetic Parameters report — one section per trait:
      1. Variance Components
      2. Heritability (with cautious narrative)
      3. GCV and PCV
      4. Genetic Advance (GA / GAM)
      5. Breeding Implication
      6. Interpretation (cleaned prose)
      7. Academic Writing Support
      8. Scope note
    """
    try:
        n_traits   = len(data.trait_results)
        n_success  = sum(1 for tr in data.trait_results.values() if tr.status == "success")
        mode_label = "Multi-environment" if data.mode == "multi" else "Single-environment"
        analysis_type = "multi_environment" if data.mode == "multi" else "single_environment"

        domain = getattr(data, "domain", "plant_breeding") or "plant_breeding"

        # BUG FIX: domain-aware title
        doc = _new_document(
            get_report_title(domain, module="genetic_parameters"),
            f"{mode_label}  ·  {n_traits} trait(s)  ·  {n_success} analysed successfully",
        )

        if data.failed_traits:
            _add_body(
                doc,
                "Traits excluded due to analysis failure: " + ", ".join(data.failed_traits),
                italic=True,
            )
            doc.add_paragraph()

        # Initialize ReportValidator
        validator = ReportValidator()

        for trait, tr in data.trait_results.items():
            doc.add_page_break()
            _add_heading(doc, trait, level=1)

            if tr.status != "success":
                _add_body(
                    doc,
                    f"Analysis failed for this trait: {tr.error or 'No genetic parameters available.'}",
                )
                continue

            # Validate interpretation before adding to report
            validation_errors = validator.validate_genetics_interpretation(
                interpretation_text=tr.interpretation or "",
                breeding_implication_text=tr.breeding_implication or "",
                analysis_type=analysis_type,
                h2_broad_sense=tr.heritability.get("h2_broad_sense") if tr.heritability else None,
                gxe_significant=tr.gxe_significant,
            )
            all_warnings.extend(validation_errors)
            # ── Data warnings ──────────────────────────────────────────────────
            if tr.data_warnings:
                _add_heading(doc, TERMS.data_warnings, level=3)
                for w in tr.data_warnings:
                    doc.add_paragraph(f"• {w}", style="List Bullet")
                doc.add_paragraph()

            # ── Descriptive Statistics ─────────────────────────────────────────
            _add_heading(doc, TERMS.descriptive_stats, level=2)
            ds_rows = []
            if tr.grand_mean is not None:
                ds_rows.append(("Grand Mean", _fmt(tr.grand_mean)))
            if tr.descriptive_stats:
                ds = tr.descriptive_stats
                if ds.standard_deviation is not None:
                    ds_rows.append(("Standard Deviation", _fmt(ds.standard_deviation)))
                if ds.standard_error is not None:
                    ds_rows.append(("Standard Error", _fmt(ds.standard_error)))
                if ds.min is not None:
                    ds_rows.append(("Minimum", _fmt(ds.min)))
                if ds.max is not None:
                    ds_rows.append(("Maximum", _fmt(ds.max)))
                if ds.range is not None:
                    ds_rows.append(("Range", _fmt(ds.range)))
                if ds.cv_percent is not None:
                    ds_rows.append(("Coefficient of Variation (%)", _fmt(ds.cv_percent, 2)))
            if ds_rows:
                _add_stat_table(doc, ["Parameter", "Value"], ds_rows, numeric_cols={1})
            else:
                _add_body(doc, "Descriptive statistics not available.", italic=True)
            doc.add_paragraph()

            # ── 1–4. Variance components, heritability, GCV/PCV, GA/GAM ───────
            # Domain is passed so genetics-specific sections are guarded
            _add_gp_tables(doc, tr, domain=domain)

            # ── 5. Breeding Implication (plant_breeding only) ──────────────────
            # DOMAIN GUARD: "Breeding Implication" must not appear in agronomy reports
            if not is_plant_breeding_domain(domain):
                _add_heading(doc, "Practical Implication", level=2)
            else:
                _add_heading(doc, TERMS.breeding_implication, level=2)
            r_implication = _clean_interpretation(tr.breeding_implication, trait)
            if r_implication:
                _add_body(doc, r_implication)
            doc.add_paragraph()

            # ── 6. Interpretation ──────────────────────────────────────────────
            cleaned = _clean_interpretation(tr.interpretation, trait)
            if cleaned:
                _add_heading(doc, TERMS.statistical_interp, level=2)
                for para_text in cleaned.split("\n\n"):
                    para_text = para_text.strip()
                    if para_text:
                        _add_body(doc, para_text)
                doc.add_paragraph()

            # ── 7. Academic Writing Support ────────────────────────────────────
            _add_writing_support_section(doc, "genetic_parameters", trait, tr.model_dump(), analysis_type=analysis_type
            )
            doc.add_paragraph()

            # ── 8. Scope note ──────────────────────────────────────────────────
            _scope_paragraph(doc)

        _add_footer(doc)
        return _docx_response(doc, "vivasense_genetic_parameters_report.docx")

    except Exception as exc:
        logger.error("Genetic parameters export failed: %s", exc, exc_info=True)
        return JSONResponse(
            status_code=500,
            content={"detail": f"Genetic parameters export failed: {exc}"},
        )


# ============================================================================
# POST /export/correlation-word
# ============================================================================

@router.post(
    "/export/correlation-word",
    summary="Export Trait Correlations as a Word document",
    tags=["Export"],
)
async def export_correlation_word(data: CorrelationExportRequest):
    """
    VivaSense Correlation report:
      1. Pairwise correlation table (r-value, p-value, significance)
      2. Correlation interpretation (cautious, no causal language)
      3. Co-selection advice (positive pairs only)
      4. Scope note
    """
    try:
        n_traits = len(data.trait_names)
        doc = _new_document(
            "VivaSense Trait Correlation Report",
            f"{data.method.capitalize()} correlation  ·  {n_traits} trait(s)  ·  "
            f"{data.phenotypic.n_observations} genotype mean(s)",
        )

        corr = CorrelationResponse(
            dataset_token=data.dataset_token,
            trait_names=data.trait_names,
            method=data.method,
            phenotypic=data.phenotypic,
            between_genotype=data.between_genotype,
            genotypic=data.genotypic,
            interpretation=data.interpretation or "",
            warnings=data.warnings,
            statistical_note=data.statistical_note or "",
        )

        logger.info(
            "Correlation export: processing interpretation (length: %d chars)",
            len(corr.interpretation)
        )

        _add_correlation_section(doc, corr)

        _scope_paragraph(doc)
        _add_footer(doc)
        return _docx_response(doc, "vivasense_correlation_report.docx")

    except Exception as exc:
        logger.error("Correlation export failed: %s", exc, exc_info=True)
        return JSONResponse(
            status_code=500,
            content={"detail": f"Correlation export failed: {exc}"},
        )


# ============================================================================
# POST /export/heatmap-report
# ============================================================================

def _generate_heatmap_image(
    matrix: List[List[Optional[float]]],
    labels: List[str],
    method: str,
) -> Optional[bytes]:
    """Render the correlation matrix as a colour heatmap PNG (300 DPI)."""
    try:
        n = len(labels)
        data = np.array(
            [[v if v is not None else float("nan") for v in row] for row in matrix],
            dtype=float,
        )

        fig, ax = plt.subplots(figsize=(max(6, n * 0.75), max(5, n * 0.70)))
        cmap = plt.get_cmap("RdYlGn")
        im = ax.imshow(data, cmap=cmap, vmin=-1.0, vmax=1.0, aspect="auto")

        ax.set_xticks(range(n))
        ax.set_yticks(range(n))
        ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=8)
        ax.set_yticklabels(labels, fontsize=8)

        for i in range(n):
            for j in range(n):
                val = data[i, j]
                if not np.isnan(val):
                    text_color = "white" if abs(val) > 0.75 else "black"
                    ax.text(
                        j, i, f"{val:.2f}",
                        ha="center", va="center",
                        fontsize=7, color=text_color,
                    )

        plt.colorbar(im, ax=ax, shrink=0.8, label=f"{method.capitalize()} r")
        ax.set_title(
            f"Trait Correlation Heatmap ({method.capitalize()})",
            fontsize=11, fontweight="bold", pad=12,
        )
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format="png", dpi=300, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        img_bytes = buf.read()
        logger.info(
            "Heatmap image generated: %d traits, %d bytes", n, len(img_bytes)
        )
        return img_bytes
    except Exception as exc:
        logger.warning("Heatmap image generation failed: %s", exc)
        return None


@router.post(
    "/export/heatmap-report",
    summary="Export Heatmap as a Word document",
    tags=["Export"],
)
async def export_heatmap_report(data: HeatmapExportRequest):
    """
    VivaSense Heatmap report:
      1. Heatmap image (300 DPI PNG, truly embedded in document)
      2. Numeric r-value matrix table
      3. Cautious correlation note
      4. Scope note
    """
    try:
        n_traits = len(data.labels)
        doc = _new_document(
            "VivaSense Correlation Heatmap Report",
            f"{data.method.capitalize()} correlation heatmap  ·  {n_traits} trait(s)",
        )

        doc.add_page_break()
        _add_heading(doc, "Trait Correlation Heatmap", level=1)

        # ── Heatmap image (embedded) ───────────────────────────────────────────
        img_bytes = _generate_heatmap_image(data.matrix, data.labels, data.method)
        if img_bytes:
            doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            cap = doc.add_paragraph(
                f"Figure 1. {data.method.capitalize()} correlation heatmap for "
                f"{n_traits} traits. "
                "Colour scale: green = positive correlation; red = negative correlation. "
                "Values on the diagonal = 1.00 (self-correlation). "
                "All r-values displayed are Pearson coefficients computed from "
                "genotype-level means."
            )
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cap.runs:
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
        else:
            _add_body(
                doc,
                "Heatmap image could not be generated. "
                "See the numeric matrix table below.",
                italic=True,
            )

        doc.add_paragraph()

        # ── Numeric r-value matrix table ───────────────────────────────────────
        _add_heading(doc, "Correlation Matrix (r-values)", level=2)
        n = len(data.labels)
        if n > 0 and data.matrix:
            headers = ["Trait"] + data.labels
            rows_data = []
            for i, label in enumerate(data.labels):
                row_vals = [label]
                for j in range(n):
                    val = (
                        data.matrix[i][j]
                        if i < len(data.matrix) and j < len(data.matrix[i])
                        else None
                    )
                    row_vals.append(
                        _fmt(val, 3, thousands=False) if val is not None else "—"
                    )
                rows_data.append(row_vals)
            numeric_cols = set(range(1, n + 1))
            _add_stat_table(doc, headers, rows_data, numeric_cols=numeric_cols)
        doc.add_paragraph()

        # ── Cautious correlation note ──────────────────────────────────────────
        _add_heading(doc, "Interpretation Note", level=2)
        _add_body(
            doc,
            "Correlation coefficients describe co-variation among genotype means "
            "in this experiment. They do not establish causal relationships between "
            "traits. Associations observed here may reflect shared genetic control, "
            "environmental co-responses, or experimental artefacts. "
            "Strong correlations (|r| ≥ 0.70, p < 0.05) may indicate co-selection "
            "potential, but this should be evaluated across multiple environments "
            "before being used to guide a breeding programme.",
        )
        if data.warnings:
            doc.add_paragraph()
            _add_heading(doc, "Warnings", level=3)
            for w in data.warnings:
                doc.add_paragraph(f"• {w}", style="List Bullet")

        doc.add_paragraph()
        _scope_paragraph(doc)
        _add_footer(doc)
        return _docx_response(doc, "vivasense_heatmap_report.docx")

    except Exception as exc:
        logger.error("Heatmap export failed: %s", exc, exc_info=True)
        return JSONResponse(
            status_code=500,
            content={"detail": f"Heatmap export failed: {exc}"},
        )


# ============================================================================
# POST /export/regression-word
# ============================================================================

@router.post(
    "/export/regression-word",
    summary="Export Regression results as a Word document",
    tags=["Export"],
)
async def export_regression_word(data: RegressionResponse):
    """
    VivaSense Regression report — one self-contained Word document:
      1. Fitted Equation  (publication-ready, with × operator)
      2. Model Statistics  (n, R², r, slope, intercept, SE, 95% CI, p)
      3. Interpretation    (plain-language effect + summary)
      4. Reliability Notes (flags + warnings)
      5. Academic Writing Support (Layer C guided-writing starters)
      6. Scope note
    """
    try:
        doc = _new_document(
            title="Regression Analysis",
            subtitle=f"{data.y_variable} ~ {data.x_variable}",
        )

        # ── 1. Fitted Equation ─────────────────────────────────────────────────
        _add_heading(doc, "Fitted Equation", level=1)
        eq_para = doc.add_paragraph()
        run = eq_para.add_run(data.equation)
        run.bold = True
        run.font.size = Pt(13)
        doc.add_paragraph()

        # ── 2. Model Statistics ────────────────────────────────────────────────
        _add_heading(doc, "Model Statistics", level=1)
        ci = data.confidence_interval_slope
        stat_rows = [
            ("Sample size (n)",                str(data.n)),
            ("Intercept (a)",                  _fmt(data.intercept, 4)),
            ("Slope (b)",                      _fmt(data.slope, 6)),
            ("Standard Error of slope",        _fmt(data.standard_error_slope, 6)),
            ("95% CI for slope",               f"[{_fmt(ci.lower, 4)}, {_fmt(ci.upper, 4)}]"),
            ("p-value (slope coefficient)",    _fmt_p(data.p_value_slope)),
            ("R\u00b2 (coefficient of determination)", _fmt(data.r_squared, 4)),
            ("r (Pearson, derived from R\u00b2)",       _fmt(data.correlation_coefficient, 4)),
            ("Adjusted R\u00b2",               _fmt(data.adjusted_r_squared, 4)),
            ("Direction",                      data.direction.replace("_", " ")),
            ("Relationship strength",          data.strength_class.replace("_", " ")),
            ("Significance",                   data.significance_class.replace("_", " ")),
        ]
        _add_stat_table(doc, ["Parameter", "Value"], stat_rows, numeric_cols={1})
        doc.add_paragraph()

        # ── 3. Interpretation ──────────────────────────────────────────────────
        _add_heading(doc, "Interpretation", level=1)
        _add_body(doc, data.plain_language_effect)
        doc.add_paragraph()
        _add_body(doc, data.summary_interpretation)
        doc.add_paragraph()

        # ── 4. Reliability Notes ───────────────────────────────────────────────
        if data.warnings or data.reliability_flags:
            _add_heading(doc, "Reliability Notes", level=1)
            if data.reliability_flags:
                flag_text = ", ".join(data.reliability_flags)
                _add_body(doc, f"Flags: {flag_text}", italic=True)
            for w in data.warnings:
                doc.add_paragraph(f"\u26a0\ufe0e  {w}", style="List Bullet")
            doc.add_paragraph()

        # ── 5. Academic Writing Support ────────────────────────────────────────
        result_dict = data.model_dump()
        _add_writing_support_section(doc, "regression", None, result_dict)
        doc.add_paragraph()

        # ── 6. Scope note ──────────────────────────────────────────────────────
        _scope_paragraph(doc)
        _add_footer(doc)
        return _docx_response(doc, "vivasense_regression_report.docx")

    except Exception as exc:
        logger.error("Regression export failed: %s", exc, exc_info=True)
        return JSONResponse(
            status_code=500,
            content={"detail": f"Regression export failed: {exc}"},
        )


# ============================================================================
# POST /export/descriptive-stats-word
# ============================================================================

@router.post(
    "/export/descriptive-stats-word",
    summary="Export Descriptive Statistics results as a Word document",
    tags=["Export"],
)
async def export_descriptive_stats_word(data: DescriptiveResponse):
    logger.info(
        "[DESC_EXPORT] Request received: %d traits, token=%s",
        len(data.summary_table),
        data.dataset_token or "none",
    )
    try:
        doc = _new_document(
            "VivaSense Descriptive Statistics Report",
            f"{len(data.summary_table)} trait(s) analysed",
        )

        # ── 1. Executive Summary ──────────────────────────────────────────────
        _add_heading(doc, "Executive Summary", level=1)
        _add_body(doc, "This report outlines the descriptive statistics, data quality checks, and variability analyses for the selected traits.")
        if data.global_flags:
            _add_body(doc, "Several statistical warnings or data quality issues were detected:", italic=True)
            _add_heading(doc, "Key Warnings", level=2)
            for flag in data.global_flags:
                doc.add_paragraph(f"• {flag}")
        else:
            _add_body(doc, "All traits exhibit good experimental precision with no severe quality flags.", italic=True)

        doc.add_paragraph()
        _add_body(doc, f"Recommendation: {data.recommendation}")
        doc.add_page_break()

        # ── 2. Data Quality / Missing Data ────────────────────────────────────
        _add_heading(doc, "Data Quality Overview", level=1)
        dq_headers = ["Trait", "Missing (n)", "Zero Values (n)", "Precision Class", "Flags"]
        dq_rows = [
            [
                t.trait,
                str(t.missing_count),
                str(t.zero_count),
                t.precision_class.capitalize() if t.precision_class else "—",
                ", ".join(t.flags) if t.flags else "None",
            ]
            for t in data.summary_table
        ]
        _add_stat_table(doc, dq_headers, dq_rows, numeric_cols={1, 2})
        doc.add_paragraph()

        # ── 3. Descriptive Statistics Table ───────────────────────────────────
        _add_heading(doc, "Descriptive Statistics (Overall)", level=1)
        headers = ["Trait", "n", "Mean", "SD", "Min", "Median", "Max", "CV%", "Skewness", "Kurtosis"]
        rows = [
            [
                t.trait,
                str(t.n),
                _fmt(t.mean, 4),
                _fmt(t.sd, 4),
                _fmt(t.minimum, 4),
                _fmt(t.median, 4),
                _fmt(t.maximum, 4),
                f"{_fmt(t.cv_percent, 2)}%" if t.cv_percent is not None else "—",
                _fmt(t.skewness, 4),
                _fmt(t.kurtosis, 4),
            ]
            for t in data.summary_table
        ]
        _add_stat_table(doc, headers, rows, numeric_cols=set(range(1, 10)))
        doc.add_paragraph()

        # ── 4. Trait-by-Trait Interpretations ─────────────────────────────────
        doc.add_page_break()
        _add_heading(doc, "Trait-by-Trait Interpretations", level=1)
        for t in data.summary_table:
            _add_heading(doc, t.trait, level=2)
            _add_body(doc, t.interpretation)
            if t.flags:
                _add_body(doc, f"Flags: {', '.join(t.flags)}", italic=True)
            doc.add_paragraph()

        _scope_paragraph(doc)
        _add_footer(doc)

        logger.info("[DESC_EXPORT] Document built successfully, sending response.")
        return _docx_response(doc, "vivasense_descriptive_stats_report.docx")

    except Exception as exc:
        logger.error("[DESC_EXPORT] Export failed: %s", exc, exc_info=True)
        return JSONResponse(
            status_code=500,
            content={"detail": f"Descriptive Stats export failed: {exc}"},
        )
