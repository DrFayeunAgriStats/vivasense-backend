#!/usr/bin/env python3
"""Regression coverage for two report-rendering defects found in the 4-model
ANOVA audit (CRD / RCBD / split-plot RCBD / factorial).

Both live purely in the Python interpretation & presentation layer — no
statistical computation is exercised or asserted here.

Bug 1 — "Assumption Diagnostics & Data Transformation" crashed with
    TypeError: unsupported format string passed to dict.__format__
because `jsonlite::toJSON(..., auto_unbox = TRUE)` renders an R `NULL` as an
empty JSON *object* (`{}`), not `null`. An absent Box-Cox lambda therefore
arrived as `{}`, survived the `is not None` guard, and blew up inside a
numeric format spec. The crash additionally deleted every *later* section for
that trait, because one try/except wrapped the whole per-trait render.

Bug 2 — numeric genotype/treatment/block identifiers (e.g. `VAR NO` = 13,
with no genotype-name column) displayed as "13.0". `DataFrame.iterrows()`
collapses an all-numeric row to a float64 Series, so `str(row[col])` yielded
"13.0". Datasets with string genotype names keep object dtype and were
unaffected — which is why CRD (named varieties) looked fine and RCBD
(numeric IDs) did not.

Run:  .venv/Scripts/python.exe -m pytest test_report_rendering_regressions.py -v
"""

import os
import re
import shutil
import subprocess
import sys
import textwrap

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "genetics-module"))

import pandas as pd  # noqa: E402
from docx import Document  # noqa: E402

import genetics_export as gx  # noqa: E402
from column_utils import format_label  # noqa: E402
from multitrait_upload_routes import build_observations  # noqa: E402


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def _doc_text(doc: Document) -> str:
    """Every rendered string in the document — body paragraphs *and* table
    cells. Table cells matter: the mean-separation and outlier tables are where
    the ".0" labels showed up."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _ta_strong_nonnormality() -> dict:
    """`transformation_analysis` exactly as the R engine emits it for a trait
    whose values sit inside [0, 100] with assumption-violating skew — the
    DTM_Days / Days_to_flowering / col_100_seed_weight_g shape.

    `boxcox_lambda` and `boxcox_ci` are `{}` — an R NULL through jsonlite, NOT
    a placeholder invented for this test. See `test_r_engine_emits_empty_dict_
    for_null_lambda` which regenerates this from the real engine.
    """
    return {
        "triggered": True,
        "is_proportion": True,
        "boxcox_lambda": {},
        "boxcox_ci": {},
        "recommended_transform": "arcsine",
        "formula_used": "asin(sqrt(x/100))",
        "rationale": "Response is on a percentage (0-100) scale; arcsine square-root applied.",
        "raw_diagnostics": {
            "shapiro": {"statistic": 0.7576, "p_value": 0.0002},
            "levene": {"statistic": 2.11, "p_value": 0.012},
            "assumptions_met": False,
        },
        "transformed_shapiro": {"statistic": 0.9601, "p_value": 0.081},
        "transformed_levene": {"statistic": 1.02, "p_value": 0.44},
        "transformed_assumptions_met": True,
        "means_original_scale": {
            "genotype": [13.0, 7.0, 2.0],
            "mean_transformed": [0.91, 0.88, 0.80],
            "mean_original": [61.2, 59.8, 51.4],
            "ci_lower_original": [58.1, 56.9, 48.0],
            "ci_upper_original": [64.1, 62.5, 54.7],
        },
        "ci_note": "Confidence intervals were back-transformed independently.",
        "disclosure_text": "Shapiro-Wilk indicated non-normal residuals (W = 0.758, p < 0.001).",
        "raw_override_disclosure": "Untransformed results reported at the user's discretion.",
    }


def _ta_mild_nonnormality() -> dict:
    """The confirmed *working* contrast case: RCBD / IFWT_g, W = 0.9220,
    p = 0.0018 — Box-Cox lambda = -1.25 snapped to a reciprocal recommendation,
    with a full back-transformed means table."""
    return {
        "triggered": True,
        "is_proportion": False,
        "boxcox_lambda": -1.25,
        "boxcox_ci": [-1.55, -0.95],
        "recommended_transform": "reciprocal",
        "formula_used": "1/x",
        "rationale": "Box-Cox lambda estimate (-1.25, 95% CI -1.55 to -0.95) closest to reciprocal.",
        "raw_diagnostics": {
            "shapiro": {"statistic": 0.9220, "p_value": 0.0018},
            "levene": {"statistic": 1.44, "p_value": 0.14},
            "assumptions_met": False,
        },
        "transformed_shapiro": {"statistic": 0.9812, "p_value": 0.31},
        "transformed_levene": {"statistic": 0.93, "p_value": 0.55},
        "transformed_assumptions_met": True,
        "means_original_scale": {
            "genotype": ["IT97K-499-35", "IT98K-205-8"],
            "mean_transformed": [0.0071, 0.0080],
            "mean_original": [140.8, 125.0],
            "ci_lower_original": [131.2, 117.4],
            "ci_upper_original": [151.9, 133.6],
        },
        "ci_note": "Confidence intervals were back-transformed independently.",
        "disclosure_text": "A reciprocal transformation (1/x) was applied based on lambda = -1.25.",
        "raw_override_disclosure": "Untransformed results reported at the user's discretion.",
    }


def _download_request():
    """A minimal but realistic single-trait ANOVA report request, carrying the
    numeric-ID mean separation and the crashing transformation payload — so one
    fixture exercises both bugs through the real per-trait render loop."""
    from genetics_schemas import AnovaTable, GeneticsResponse, GeneticsResult, MeanSeparation
    from multitrait_upload_schemas import DatasetSummary, TraitResult

    result = GeneticsResult(
        environment_mode="single", grand_mean=57.4, n_genotypes=20, n_reps=3,
        variance_components={}, heritability={}, genetic_parameters={},
        cv_percent=4.2,
        anova_table=AnovaTable(
            source=["rep", "genotype", "Residuals"],
            df=[2, 19, 38], ss=[10.0, 900.0, 200.0], ms=[5.0, 47.4, 5.26],
            f_value=[0.95, 9.0, None], p_value=[0.39, 0.0001, None],
        ),
        mean_separation=MeanSeparation(
            genotype=["13.0", "7.0", "2.0"], mean=[61.2, 59.8, 51.4],
            se=[1.2, 1.3, 1.1], group=["a", "ab", "b"],
        ),
        assumption_tests={
            "normality": {"test": "Shapiro-Wilk", "statistic": 0.7576, "p_value": 0.0002},
            "homogeneity": {"test": "Levene", "statistic": 2.11, "p_value": 0.012},
        },
        transformation_analysis=_ta_strong_nonnormality(),
    )
    return gx.DownloadReportRequest(
        success=True, module="anova",
        summary_table=[],
        trait_results={"DTM_Days": TraitResult(
            trait="DTM_Days", status="success",
            analysis_result=GeneticsResponse(
                status="success", mode="single", success=True, result=result,
                interpretation="Genotype effects were significant.",
            ),
        )},
        dataset_summary=DatasetSummary(n_reps=3, n_traits=1, mode="single"),
        transformation_choice="transformed",
    )


# ===========================================================================
# Bug 1 — transformation section must never crash, never vanish
# ===========================================================================

class TestTransformationSectionRendering:

    def test_strong_nonnormality_renders_without_crash(self):
        """The reported repro: boxcox_lambda == {} must not reach a format spec."""
        doc = Document()
        gx._add_transformation_section(doc, _ta_strong_nonnormality(), choice="transformed")
        text = _doc_text(doc)

        assert "Assumption Diagnostics & Data Transformation" in text
        # An honest statement of the basis, not a crash and not a raw container.
        assert "arcsine" in text.lower()
        assert "{}" not in text
        assert "dict" not in text.lower()

    def test_mild_nonnormality_renders_full_transformation_table(self):
        """The confirmed-working contrast must keep working: numeric lambda,
        CI, and the back-transformed means table."""
        doc = Document()
        gx._add_transformation_section(doc, _ta_mild_nonnormality(), choice="transformed")
        text = _doc_text(doc)

        assert "Box-Cox λ = -1.25" in text
        assert "95% CI -1.55 to -0.95" in text
        assert "reciprocal" in text.lower()
        # Back-transformed means table with asymmetric CIs.
        assert "Mean (original)" in text
        assert "140.800" in text or "140.8" in text

    def test_absent_lambda_without_proportion_is_stated_explicitly(self):
        """No lambda and not a proportion => say so, rather than silently
        implying an arcsine/proportion basis that was never estimated."""
        ta = _ta_strong_nonnormality()
        ta["is_proportion"] = False
        ta["recommended_transform"] = "log"
        ta["formula_used"] = "log(x)"
        ta["rationale"] = "Box-Cox estimation did not converge; log applied as fallback."

        doc = Document()
        gx._add_transformation_section(doc, ta, choice="transformed")
        text = _doc_text(doc)

        # The headline "Recommended transformation:" line must state the missing
        # lambda honestly rather than claiming a proportion/arcsine basis.
        headline = next(l for l in text.splitlines() if l.startswith("Recommended transformation:"))
        assert "could not be estimated" in headline
        assert "arcsine" not in headline.lower()
        assert "{}" not in text

    @pytest.mark.parametrize("bad_lambda", [{}, {"value": 1.0}, [], None, "n/a", float("nan")])
    def test_no_lambda_shape_can_crash_the_section(self, bad_lambda):
        """Defence in depth: whatever shape the engine sends, render something."""
        ta = _ta_strong_nonnormality()
        ta["boxcox_lambda"] = bad_lambda
        doc = Document()
        gx._add_transformation_section(doc, ta, choice="transformed")
        assert "Assumption Diagnostics & Data Transformation" in _doc_text(doc)

    def test_untriggered_case_is_a_single_clean_sentence(self):
        doc = Document()
        gx._add_transformation_section(
            doc, {"triggered": False, "disclosure_text": "No transformation was required."},
            choice="raw",
        )
        text = _doc_text(doc)
        assert "No transformation was required." in text
        assert "Box-Cox" not in text


class TestRNumCoercion:
    """`_r_num` is the guard that turns the jsonlite NULL -> {} artifact into a
    branchable None."""

    @pytest.mark.parametrize("value,expected", [
        ({}, None),                 # R NULL through jsonlite -- the actual bug
        ({"a": 1}, None),
        ([], None),
        ([1.0], None),
        (None, None),
        ("nope", None),
        (float("nan"), None),
        (float("inf"), None),
        (True, None),               # bool is not a measurement
        (-1.25, -1.25),
        (0, 0.0),
        ("2.5", 2.5),
    ])
    def test_r_num(self, value, expected):
        assert gx._r_num(value) == expected

    @pytest.mark.parametrize("value,expected", [
        ({}, None),
        ([-1.55, -0.95], [-1.55, -0.95]),
        ([1.0], None),
        ([1.0, 2.0, 3.0], None),
        ([None, 2.0], None),
        (None, None),
    ])
    def test_r_num_pair(self, value, expected):
        assert gx._r_num_pair(value) == expected


class TestSectionIsolation:
    """A failing section must not delete its siblings — the RCBD cascade where
    the transformation crash silently removed Interpretation & Recommendations."""

    def test_failing_section_leaves_visible_marker_and_spares_siblings(self):
        doc = Document()

        with gx._section_guard(doc, "DTM_Days", "Assumption Diagnostics & Data Transformation"):
            raise TypeError("unsupported format string passed to dict.__format__")

        # A sibling section rendered after the failure is unaffected.
        with gx._section_guard(doc, "DTM_Days", "Interpretation & Recommendations"):
            doc.add_paragraph("Genotype 13 exhibited the highest DTM_Days performance.")

        text = _doc_text(doc)
        # The failure is disclosed, not silently swallowed.
        assert "Assumption Diagnostics & Data Transformation" in text
        assert "could not be generated" in text
        assert "TypeError" in text
        # ...and the sibling section still made it into the report.
        assert "highest DTM_Days performance" in text

    def test_guard_is_transparent_when_nothing_fails(self):
        doc = Document()
        with gx._section_guard(doc, "T", "Section"):
            doc.add_paragraph("content")
        text = _doc_text(doc)
        assert "content" in text
        assert "could not be generated" not in text

    def test_end_to_end_transformation_failure_spares_interpretation(self, monkeypatch):
        """The reported RCBD cascade, through the real per-trait render loop:
        force the transformation subsection to raise and assert the later
        sections are still in the document."""
        def _boom(*args, **kwargs):
            raise TypeError("unsupported format string passed to dict.__format__")

        monkeypatch.setattr(gx, "_add_transformation_section", _boom)

        doc = Document()
        gx.export_traits_to_word(_download_request(), doc)
        text = _doc_text(doc)

        # The failure is disclosed...
        assert "could not be generated" in text
        assert "Assumption Diagnostics & Data Transformation" in text
        # ...and the sibling sections that used to disappear are still present.
        assert "Mean Separation" in text
        assert "Interpretation" in text

    def test_end_to_end_renders_clean_when_nothing_fails(self):
        doc = Document()
        gx.export_traits_to_word(_download_request(), doc)
        text = _doc_text(doc)

        assert "could not be generated" not in text
        assert "Assumption Diagnostics & Data Transformation" in text
        assert "Interpretation" in text
        # Bug 2, through the whole pipeline: no identifier cell ends in ".0".
        label_cells = [row.cells[0].text for table in doc.tables for row in table.rows]
        offenders = [c for c in label_cells if TRAILING_ZERO.fullmatch(c)]
        assert not offenders, f"trailing .0 identifier cells: {offenders}"


# ===========================================================================
# Bug 2 — numeric identifiers must never render with a trailing ".0"
# ===========================================================================

TRAILING_ZERO = re.compile(r"(?<!\d)(\d+)\.0(?!\d)")


class TestFormatLabel:

    @pytest.mark.parametrize("value,expected", [
        (13.0, "13"),
        (13, "13"),
        ("13.0", "13"),             # already stringified upstream / cached
        ("-4.00", "-4"),
        (0.0, "0"),
        ("IT97K-499-35", "IT97K-499-35"),
        ("2.5", "2.5"),             # genuinely fractional code -- leave alone
        (2.5, "2.5"),
        ("R1", "R1"),
        ("  G07  ", "G07"),
    ])
    def test_format_label(self, value, expected):
        assert format_label(value) == expected

    def test_numpy_scalars(self):
        import numpy as np
        assert format_label(np.float64(13.0)) == "13"
        assert format_label(np.int64(13)) == "13"

    @pytest.mark.parametrize("value", [None, float("nan"), ""])
    def test_missing_values(self, value):
        assert format_label(value) == "—"


class TestBuildObservationsNumericIds:
    """Fixture modelled on the reported RCBD dataset: VAR NO 1-20, numeric REP,
    and no genotype-name column — so every column is numeric and `iterrows()`
    upcasts the row to float64."""

    @staticmethod
    def _numeric_id_frame() -> pd.DataFrame:
        rows = []
        for var in range(1, 21):
            for rep in range(1, 4):
                rows.append({"VAR NO": var, "REP": rep,
                             "DTM_Days": 60.0 + var * 0.5 + rep * 0.3})
        return pd.DataFrame(rows)

    def test_row_series_really_is_upcast(self):
        """Pin the mechanism, so this test still means something if pandas
        changes its dtype behaviour."""
        df = self._numeric_id_frame()
        _, row = next(df.iterrows())
        assert row.dtype == "float64"
        assert str(row["VAR NO"]) == "1.0"   # what the old code produced

    def test_no_trailing_zero_in_observation_records(self):
        obs = build_observations(
            self._numeric_id_frame(),
            genotype_col="VAR NO", rep_col="REP", trait_col="DTM_Days", env_col=None,
        )
        assert len(obs) == 60
        for rec in obs:
            assert not TRAILING_ZERO.search(rec["genotype"]), rec
            assert not TRAILING_ZERO.search(rec["rep"]), rec
        assert {r["genotype"] for r in obs} == {str(i) for i in range(1, 21)}
        assert {r["rep"] for r in obs} == {"1", "2", "3"}

    def test_string_genotype_names_are_untouched(self):
        df = pd.DataFrame({
            "GENOTYPE": ["IT97K-499-35", "IT98K-205-8"] * 3,
            "REP": [1, 1, 2, 2, 3, 3],
            "Yield": [45.2, 50.1, 47.3, 49.9, 46.1, 51.0],
        })
        obs = build_observations(df, genotype_col="GENOTYPE", rep_col="REP",
                                 trait_col="Yield", env_col=None)
        assert {r["genotype"] for r in obs} == {"IT97K-499-35", "IT98K-205-8"}


class TestRenderedReportHasNoFloatLabels:
    """End-to-end on the presentation layer: no label anywhere in the rendered
    output — table cells included — may carry a trailing ".0"."""

    def test_mean_separation_table_and_narrative(self):
        from genetics_schemas import MeanSeparation

        ms = MeanSeparation(
            genotype=["13.0", "7.0", "2.0"],   # as cached/stored before the fix
            mean=[61.2, 59.8, 51.4],
            se=[1.2, 1.3, 1.1],
            group=["a", "ab", "b"],
            test="Tukey HSD",
            alpha=0.05,
        )
        doc = Document()
        gx._add_mean_separation_section(doc, "DTM_Days", ms, domain="plant_breeding")
        text = _doc_text(doc)

        assert "13" in text
        offenders = TRAILING_ZERO.findall(text)
        assert not offenders, f"trailing .0 labels rendered: {offenders}\n{text}"

    def test_outlier_table(self):
        assumption_tests = {
            "normality": {"test": "Shapiro-Wilk", "statistic": 0.7576, "p_value": 0.0002},
            "homogeneity": {"test": "Levene", "statistic": 2.11, "p_value": 0.012},
            "outlier_detection": {
                "standardized_residual_threshold": 3,
                "cooks_distance_threshold": 0.0667,
                "n_extreme_outliers": 1,
                "n_influential_observations": 1,
                "flagged_observations": [
                    {"observation": 17.0, "treatment": 13.0, "block": 2.0,
                     "standardized_residual": 3.42, "cooks_distance": 0.19,
                     "extreme_outlier": True, "influential": True},
                ],
                "interpretation": "Extreme residual outliers were detected.",
            },
        }
        doc = Document()
        gx._add_assumption_tests_section(doc, assumption_tests)
        text = _doc_text(doc)

        offenders = TRAILING_ZERO.findall(text)
        assert not offenders, f"trailing .0 labels rendered: {offenders}\n{text}"

    def test_back_transformed_means_table(self):
        doc = Document()
        gx._add_transformation_section(doc, _ta_strong_nonnormality(), choice="transformed")
        text = _doc_text(doc)
        # Genotype labels arrive as floats 13.0/7.0/2.0 in means_original_scale.
        for label in ("13", "7", "2"):
            assert label in text
        # Assert on the identifier column only: the numeric measurement columns
        # legitimately contain values like "61.200", which are not labels.
        label_cells = [row.cells[0].text for table in doc.tables for row in table.rows]
        offenders = [c for c in label_cells if TRAILING_ZERO.fullmatch(c)]
        assert not offenders, f"trailing .0 labels in identifier column: {offenders}"
        assert "13" in label_cells


# ===========================================================================
# R-engine provenance (skipped when Rscript is unavailable)
# ===========================================================================

def _rscript() -> str | None:
    for candidate in ("Rscript", r"C:\Program Files\R\R-4.6.0\bin\x64\Rscript.exe"):
        found = shutil.which(candidate) or (candidate if os.path.exists(candidate) else None)
        if found:
            return found
    return None


@pytest.mark.skipif(_rscript() is None, reason="Rscript not available on this machine")
def test_r_engine_emits_empty_dict_for_null_lambda(tmp_path):
    """Provenance for `_ta_strong_nonnormality`: prove the engine really does
    emit `"boxcox_lambda":{}` for a [0,100]-range trait with violated
    assumptions, so the fixture above is not an invented shape.

    The dataset is engineered to land in the reported Shapiro-Wilk band
    (W ~ 0.75-0.85) rather than relying on a checked-in fixture.
    """
    engine = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                          "genetics-module", "vivasense_genetics.R").replace("\\", "/")
    script = tmp_path / "probe.R"
    script.write_text(textwrap.dedent(f"""
        suppressWarnings(suppressMessages(source("{engine}")))
        set.seed(11)
        n_geno <- 20; n_rep <- 3; N <- n_geno * n_rep
        geno <- rep(seq_len(n_geno), each = n_rep)
        rp   <- rep(seq_len(n_rep), times = n_geno)
        base <- rep(runif(n_geno, 30, 55), each = n_rep)
        y <- base + rexp(N, 1/3)
        y[sample(N, 6)] <- runif(6, 88, 99)
        y <- pmin(pmax(y, 0), 100)
        d <- data.frame(trait_value = as.numeric(y),
                        genotype = as.character(geno), rep = as.character(rp))
        sw <- shapiro.test(residuals(aov(trait_value ~ factor(rep) + factor(genotype), data = d)))
        ta <- compute_anova_transformation(d, "DTM_Days", crd_mode = FALSE)
        cat("W=", round(sw$statistic, 4), "\\n", sep = "")
        cat("TRIGGERED=", isTRUE(ta$triggered), "\\n", sep = "")
        cat(jsonlite::toJSON(list(boxcox_lambda = ta$boxcox_lambda),
                             auto_unbox = TRUE, na = "null", digits = 10), "\\n")
    """), encoding="utf-8")

    proc = subprocess.run([_rscript(), str(script)], capture_output=True, text=True, timeout=300)
    assert proc.returncode == 0, proc.stderr
    out = proc.stdout

    w = float(re.search(r"W=([\d.]+)", out).group(1))
    assert 0.70 <= w <= 0.90, f"engineered dataset drifted out of the reported band: W={w}"
    assert "TRIGGERED=TRUE" in out, out
    assert '"boxcox_lambda":{}' in out, f"expected the R NULL -> {{}} artifact, got: {out}"


if __name__ == "__main__":
    sys.exit(pytest.main([__file__, "-v"]))
